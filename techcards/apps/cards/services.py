"""
Business logic for tech-card generation:
  - Access control (free quota vs. paid quota)
  - DOCX file creation via python-docx
  - PDF conversion via xhtml2pdf
"""

import io
import os
import logging
from pathlib import Path

from django.conf import settings
from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from apps.accounts.models import FreeCardUsage
from apps.standards.models import NormativeDocument
from .models import TechCard

logger = logging.getLogger(__name__)


def can_create_card(user, doc: NormativeDocument) -> tuple[bool, str, bool]:
    """
    Check whether the user is allowed to create a tech card.

    Returns:
        (allowed: bool, reason: str, is_free: bool)
    """
    if not user.is_authenticated:
        return False, "Для создания техкарты необходима регистрация.", False

    if not FreeCardUsage.has_used_free(user, doc.code):
        return True, "Бесплатная разработка (первая для данного документа).", True

    if user.tech_card_quota > 0:
        return True, "Использован 1 слот из оплаченного лимита.", False

    return (
        False,
        "Исчерпан лимит разработки. Пополните счёт в личном кабинете.",
        False,
    )


def charge_user(user, doc: NormativeDocument, is_free: bool) -> None:
    """
    Deduct quota or record free usage after successful card generation.
    """
    if is_free:
        FreeCardUsage.objects.get_or_create(user=user, normative_doc_code=doc.code)
    else:
        user.consume_quota()


def _get_or_create_docx_template(doc: NormativeDocument) -> Document:
    """
    Load a DOCX template from card_templates/ or create a generic one.
    """
    template_path = Path(settings.CARD_TEMPLATES_DIR) / f"{doc.code.replace('/', '_')}.docx"
    if template_path.exists():
        return Document(str(template_path))
    # Fallback: build a generic template programmatically
    return _build_generic_template()


def _build_generic_template() -> Document:
    document = Document()
    # Page margins
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    return document


def _fill_docx(document: Document, card_data: dict) -> None:
    """Fill the document with computed card data."""
    from docx.shared import RGBColor

    doc_code = card_data.get("document_code", "")
    doc_name = card_data.get("document_name", "")
    method_code = card_data.get("method_code", "")

    # Header
    header_para = document.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("ТЕХНОЛОГИЧЕСКАЯ КАРТА")
    run.bold = True
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"неразрушающего контроля\n({doc_code})")

    document.add_paragraph()

    # Table with all card fields
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Параметр"
    hdr_cells[1].text = "Значение"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    field_labels = {
        "object_type": "Тип объекта",
        "material": "Материал",
        "thickness_mm": "Толщина металла, мм",
        "weld_category": "Категория шва",
        "radiation_source": "Источник излучения",
        "film_type": "Тип плёнки",
        "weld_length_mm": "Длина шва, мм",
        "object_dimensions": "Размеры объекта",
        "sensitivity_class": "Класс чувствительности",
        "sfd_min_mm": "Минимальное РИП, мм",
        "max_pore_diameter_mm": "Допустимый диаметр пор, мм",
        "max_pore_count_per_100mm": "Допустимое кол-во пор/100 мм",
        "acceptance_basis": "Основание для приёмки",
        "nk_specialist": "Специалист НК",
        "document_code": "Нормативный документ",
        "control_stage": "Стадия контроля",
        "welding_method": "Метод сварки",
        "allowable_undercut_depth_mm": "Допустимая глубина подреза, мм",
        "allowable_pore_dia_mm": "Допустимый диаметр поры, мм",
        "allowable_reinforcement_mm": "Допустимая выпуклость шва, мм",
        "scope_of_control": "Объём контроля",
        "penetrant_dwell_time_min": "Время пропитки пенетрантом, мин",
        "developer_dwell_time_min": "Время проявления, мин",
        "test_pressure_mpa": "Испытательное давление, МПа",
        "holding_time_min": "Время выдержки, мин",
        "nk_specialist": "Специалист НК",
    }

    skip_keys = {"document_name", "method_code"}
    for key, value in card_data.items():
        if key in skip_keys or value is None:
            continue
        label = field_labels.get(key, key)
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)

    document.add_paragraph()

    # Signatures section
    sig_table = document.add_table(rows=3, cols=3)
    sig_table.style = "Table Grid"
    labels = ["Разработал", "Проверил", "Утвердил"]
    for i, lbl in enumerate(labels):
        cells = sig_table.rows[i].cells
        cells[0].text = lbl
        cells[1].text = "________________"
        cells[2].text = "«___» ____________ 20___ г."

    # Footer date
    document.add_paragraph()
    footer = document.add_paragraph(f"Дата создания: {timezone.now().strftime('%d.%m.%Y')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def generate_docx(card: TechCard) -> bytes:
    """Generate a DOCX file for the given TechCard and return raw bytes."""
    doc = _get_or_create_docx_template(card.normative_doc)
    _fill_docx(doc, card.generated_data)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_pdf_from_html(html_content: str) -> bytes:
    """Convert HTML to PDF using xhtml2pdf."""
    from xhtml2pdf import pisa

    buffer = io.BytesIO()
    result = pisa.CreatePDF(io.StringIO(html_content), dest=buffer)
    if result.err:
        raise RuntimeError(f"PDF generation error: {result.err}")
    return buffer.getvalue()


def create_tech_card(user, doc: NormativeDocument, input_data: dict) -> TechCard:
    """
    Full pipeline:
    1. Check access
    2. Call ndt_data module to compute card data
    3. Generate DOCX and PDF
    4. Save TechCard record
    5. Charge quota
    """
    allowed, reason, is_free = can_create_card(user, doc)
    if not allowed:
        raise PermissionError(reason)

    data_module = doc.get_data_module()
    if data_module is None:
        raise ValueError(f"Модуль данных для {doc.code} не реализован.")

    generated_data = data_module.generate_card_data(input_data)
    generated_data["method_code"] = doc.method.code

    title = f"Техкарта {doc.code} — {input_data.get('object_type', 'объект')}"
    card = TechCard.objects.create(
        user=user,
        normative_doc=doc,
        title=title,
        input_data=input_data,
        generated_data=generated_data,
        status=TechCard.Status.PENDING,
        is_free=is_free,
    )

    try:
        docx_bytes = generate_docx(card)
        card.docx_file.save(f"{card.id}.docx", ContentFile(docx_bytes), save=False)
        card.status = TechCard.Status.GENERATED
    except Exception as exc:
        logger.exception("DOCX generation failed for card %s", card.id)
        card.status = TechCard.Status.ERROR
        card.error_message = str(exc)

    card.save()

    if card.status == TechCard.Status.GENERATED:
        charge_user(user, doc, is_free)

    return card
