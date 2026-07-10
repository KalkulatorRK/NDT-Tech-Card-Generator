"""
Генератор документа «Техническая справка (расчёты параметров контроля)».

Формирует подробный лог расчётов для технологической карты РГК
со ссылками на нормативные документы, формулы и расшифровкой переменных.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

from techcards.scheme_display import get_scheme_user_label
from normative.gost_50_05_07 import DOCUMENT_CODE, DOCUMENT_FULL_NAME
from normative.np_105_18 import DOCUMENT_CODE as NP105_CODE


# Расшифровка обозначений, используемых в формулах
VARIABLE_GLOSSARY = [
    ('S', 'Номинальная толщина стенки (основного металла) в месте сварки, мм'),
    ('D', 'Наружный диаметр трубопровода, мм'),
    ('d', 'Внутренний диаметр трубопровода, мм (d = D − 2S)'),
    ('g_min', 'Наименьшая допустимая высота усиления (валика) шва, мм'),
    ('g_max', 'Наибольшая допустимая высота усиления (валика) шва, мм'),
    ('S_K', 'Номинальная толщина в месте сварки для поиска K: S + g_min, мм'),
    ('S_рад(K)', 'Радиационная толщина для расчёта чувствительности K, мм'),
    ('S_рад(f)', 'Радиационная толщина для расчёта расстояния f, мм'),
    ('K', 'Требуемая чувствительность контроля (диаметр проволоки ИКИ), мм'),
    ('Φ (d)', 'Размер фокусного пятна источника излучения, мм'),
    ('f (SFD)', 'Расстояние от источника до детектора (плёнки), мм'),
    ('b (OFD)', 'Расстояние от объекта контроля до детектора, мм'),
    ('Ug', 'Геометрическая нерезкость изображения, мм'),
    ('C', 'Вспомогательный коэффициент: C = max(2Φ/K, 4)'),
    ('N', 'Минимальное число экспозиций (снимков) на участок контроля'),
    ('L', 'Длина участка, охватываемого одной экспозицией, мм'),
    ('m', 'Отношение внутреннего диаметра к наружному: m = d/D'),
    ('l', 'Длина радиографического снимка (плёнки), мм'),
]


def build_calculation_log(input_data: dict, params: dict) -> list[dict]:
    """
    Формирует структурированный лог расчётов из исходных и рассчитанных данных.

    Каждый элемент — секция с полями:
      title, normative_ref, logic, variables, formula, steps, result, notes
    """
    rad = params.get('rad_thickness') or {}
    scheme = params.get('exposure_scheme') or {}
    scheme_info = params.get('scheme_info') or {}
    src = params.get('selected_source') or {}
    screens = params.get('screens') or {}
    film = params.get('film_class_info') or {}
    iqi = params.get('recommended_iqi') or {}

    S = params.get('wall_thickness', '')
    D = params.get('outer_diameter', '')
    d_inner = params.get('d_inner_mm', '')
    g_min = params.get('g_min_mm', '')
    g_max = params.get('g_max_mm', '')

    sections: list[dict] = []

    sections.append({
        'title': '1. Исходные данные пользователя',
        'normative_ref': 'Входные параметры формы создания техкарты',
        'logic': (
            'Пользователь вводит идентификационные данные объекта, геометрию, '
            'тип сварного соединения, источник излучения и параметры просвечивания. '
            'Эти значения являются входом для всех последующих расчётов.'
        ),
        'variables': ['S', 'D', 'g_min', 'g_max', 'Φ (d)', 'b (OFD)'],
        'formula': None,
        'steps': [
            f'Организация: {params.get("organization", "—")}',
            f'Объект контроля: {params.get("object_name", "—")}',
            f'Марка стали: {params.get("material", "—")}',
            f'Толщина стенки S = {S} мм',
            f'Наружный диаметр D = {D or "—"} мм',
            f'Условное обозначение шва: {params.get("joint_designation", "—")} '
            f'({params.get("joint_name", "")})',
            f'Категория сварного соединения: {params.get("weld_category", "—")}',
            f'Способ сварки: {params.get("welding_process", "—")}',
            f'Источник излучения: {src.get("name", input_data.get("source_code", "—"))}',
            f'Фокусное пятно Φ = {params.get("source_focal_spot_mm", "—")} мм',
            f'Расстояние b (OFD) = {params.get("ofd_mm", "—")} мм',
            f'Схема просвечивания: {get_scheme_user_label(params.get("scheme_type", input_data.get("scheme_type", "")))}',
        ],
        'result': 'Исходные данные приняты для расчёта.',
        'notes': '',
    })

    sections.append({
        'title': '2. Зоны контроля и размеры шва',
        'normative_ref': 'ГОСТ Р 59023.2-2020; НП-105-18',
        'logic': (
            'По условному обозначению сварного соединения и номинальной толщине S '
            'определяются ширина валика, зона термического влияния (ОШЗ), '
            'контролируемая зона и допустимые значения усиления шва g_min, g_max. '
            'Значения g_min и g_max используются при расчёте радиационной толщины.'
        ),
        'variables': ['g_min', 'g_max', 'S', 'Sпк'],
        'formula': (
            'Ширина валика, ОШЗ и контролируемая зона — по таблицам '
            'ГОСТ Р 59023.2-2020 для кода соединения и толщины S.'
        ),
        'steps': [
            f'Код соединения: {params.get("joint_designation", "—")}',
            f'Толщина S = {S} мм',
            f'Ширина валика: {params.get("weld_bead_width_mm", "—")} мм',
            f'Высота валика (справочно): {params.get("weld_bead_height_mm", "—")} мм',
            f'g_min = {g_min} мм (для расчёта K)',
            f'g_max = {g_max} мм (для расчёта f)',
            f'Толщина подкладки Sпк = {params.get("backing_thickness_mm", 0)} мм',
            f'Ширина ОШЗ: {params.get("haz_width_mm", "—")} мм',
            f'Ширина контролируемой зоны: {params.get("zone_width_mm", "—")} мм',
        ],
        'result': (
            f'g_min = {g_min} мм, g_max = {g_max} мм — приняты для расчёта '
            f'радиационной толщины.'
        ),
        'notes': params.get('zone_note', ''),
    })

    sections.append({
        'title': '3. Категория сварного соединения',
        'normative_ref': 'НП-105-18, Таблица N 4.8',
        'logic': (
            'Объём контроля задаётся пользователем (100, 50, 25, 10 или 5 %). '
            'По НП-105-18, п. 71: для прямолинейных швов и кольцевых швов при D > 250 мм '
            'число экспозиций и контролируемых участков уменьшается пропорционально объёму, '
            'но не менее 1. По п. 72: для кольцевых швов при D ≤ 250 мм контроль '
            'проводится по всей протяжении — расчётные N и N_segments не уменьшаются.'
        ),
        'variables': [],
        'formula': (
            'N_факт = max(1, ⌈N_100% × V / 100⌉); '
            'N_участков аналогично'
        ),
        'steps': [
            f'Объём контроля: {params.get("control_volume_pct", "—")} %',
            f'Режим: {params.get("control_volume_mode", "—")}',
            f'N при 100 %: {params.get("N_calculated_full", params.get("N_calculated", "—"))}',
            f'N в техкарте (6.6): {params.get("N_calculated", "—")}',
            f'Участков в техкарте (6.7): {params.get("N_segments", "—")}',
        ],
        'result': (
            f'Объём: {params.get("control_volume_pct", 100)} %; '
            f'N = {params.get("N_calculated", "—")}, '
            f'участков = {params.get("N_segments", "—")}'
        ),
        'notes': '',
    })

    sk_formula = rad.get('formula_k', '—')
    if rad.get('wall_count') == 2:
        sk_logic = (
            'По НП-105-18, п. 46 / Табл. 4.8, при просвечивании через две стенки '
            'номинальная толщина в месте сварки: S_K = S + S. '
            'Значение K — диаметр проволоки ИКИ в мм.'
        )
        sk_expr = 'S_K = S + S'
    else:
        sk_logic = (
            'По НП-105-18, п. 46 / Табл. 4.8, при просвечивании через одну стенку '
            'номинальная толщина в месте сварки: S_K = S + g_min + Sпк '
            '(Sпк — толщина подкладки/кольца, 0 если не применяется). '
            'Значение K — диаметр проволоки ИКИ в мм.'
        )
        sk_expr = 'S_K = S + g_min + Sпк'

    sections.append({
        'title': '4. Требуемая чувствительность K',
        'normative_ref': 'НП-105-18, п. 46; Таблица N 4.8',
        'logic': sk_logic,
        'variables': ['S', 'g_min', 'Sпк', 'S_K', 'K'],
        'formula': f'{sk_expr}; K = f(S_K, категория) по Табл. 4.8 НП-105-18',
        'steps': [
            f'S = {S} мм, g_min = {g_min} мм, Sпк = {params.get("backing_thickness_mm", 0)} мм',
            f'S_K: {sk_formula}',
            f'Категория сварного соединения: {params.get("weld_category", "—")}',
            f'По Табл. 4.8 НП-105-18: K ≤ {params.get("required_sensitivity_mm", "—")} мм '
            f'({params.get("required_sensitivity_pct", "—")} %)',
        ],
        'result': params.get('sensitivity_desc', '—'),
        'notes': '',
    })

    sections.append({
        'title': '5. Радиационная толщина',
        'normative_ref': f'{DOCUMENT_CODE}, п. 6.3.5; НП-105-18, п. 46',
        'logic': (
            'Для определения K по НП-105-18 используется номинальная толщина S_K '
            '(см. раздел 4). Для расчёта f и подбора источника применяется S_рад(f) '
            'с наибольшим усилением g_max и учётом числа просвечиваемых стенок.'
        ),
        'variables': ['S', 'g_min', 'g_max', 'Sпк', 'S_K', 'S_рад(f)'],
        'formula': (
            'S_K — по п. 46 НП-105-18; '
            'S_рад(f) = S + g_max или 2S + 2g_max (для f)'
        ),
        'steps': [
            f'Схема: {get_scheme_user_label(params.get("scheme_type", ""))} — {rad.get("wall_desc", "")}',
            f'Число стенок (для f): {rad.get("wall_count", "—")}',
            f'S_K = {rad.get("formula_k", "—")}',
            f'S_рад(f) = {rad.get("formula_f", "—")}',
        ],
        'result': (
            f'S_K = {rad.get("s_rad_k_mm", "—")} мм (для K); '
            f'S_рад(f) = {rad.get("s_rad_f_mm", "—")} мм (для f)'
        ),
        'notes': '',
    })

    sections.append({
        'title': '6. Выбор источника излучения',
        'normative_ref': (
            f'{DOCUMENT_CODE}, приложение Б; п. 6.3.2, рис. 6'
        ),
        'logic': (
            'Радионуклидные источники подбираются по табл. Б.1–Б.3 для материала '
            'и радиационной толщины. Дополнительно рекомендуется один рентгеновский '
            'аппарат: максимальное напряжение U на трубке не должно превышать '
            'значения с номограммы рис. 6 (п. 6.3.2) для просвечиваемой толщины w.'
        ),
        'variables': ['S', 'материал'],
        'formula': 'Источники по табл. Б.1–Б.3 для стали / алюминия / титана',
        'steps': [
            f'Толщина S = {S} мм',
            f'Материал: {params.get("material_display", params.get("material", "—"))}',
            f'Радиационная толщина для табл. Б: {params.get("s_rad_f_mm", S)} мм',
            f'Таблица Б.{ {"steel": "1", "aluminum": "2", "titanium": "3"}.get(params.get("material_type", "steel"), "1") } '
            f'ГОСТ Р 50.05.07-2018',
            f'Радионуклидные источники: '
            f'{len([s for s in params.get("suitable_sources", []) if s.get("type") == "isotope"])} шт.',
            f'Рентгеновский аппарат (рис. 6): '
            f'{next((s.get("energy_display") for s in params.get("suitable_sources", []) if s.get("type") == "xray"), "—")}',
            f'Выбран: {src.get("name", "—")} ({src.get("code", "—")})',
            f'Энергия: {src.get("energy_display", "—")}',
        ],
        'result': f'Источник: {src.get("name", "—")}',
        'notes': '',
    })

    scheme_title = get_scheme_user_label(params.get('scheme_type', ''))
    scheme_steps = [
        f'Схема: {scheme_title}',
        f'Описание: {scheme_info.get("description", "—")}',
        f'Внутренний диаметр d = D − 2S = {D} − 2×{S} = {d_inner} мм',
        f'Коэффициент C = {scheme.get("C", params.get("C_coeff", "—"))}',
    ]
    if scheme.get('m') is not None:
        scheme_steps.append(f'Отношение m = d/D = {scheme.get("m")}')
    if params.get('scheme_formula') or scheme.get('formula'):
        scheme_steps.append(f'Формула f: {params.get("scheme_formula") or scheme.get("formula")}')
    if scheme.get('L_formula'):
        scheme_steps.append(scheme['L_formula'])
    elif params.get('L_formula'):
        scheme_steps.append(params['L_formula'])
    scheme_steps.extend([
        f'Минимальное расстояние f = {scheme.get("f_min_mm", params.get("f_calculated_mm", "—"))} мм',
        f'Число экспозиций N ≥ {scheme.get("N", params.get("N_calculated", "—"))}',
        f'Длина участка L = {scheme.get("L_mm", params.get("L_calculated_mm", "—"))} мм',
    ])

    empirical_note = ''
    if params.get('is_empirical') or scheme.get('is_empirical'):
        empirical_note = (
            params.get('empirical_reason')
            or scheme.get('empirical_reason')
            or 'Параметры определяются опытным путём (п. Г.5).'
        )

    sections.append({
        'title': '7. Параметры схемы просвечивания',
        'normative_ref': (
            f'{DOCUMENT_CODE}, раздел 6; ГОСТ 7512-82 (Приложение 4 для чертежа 3б)'
        ),
        'logic': (
            'По выбранной схеме просвечивания рассчитываются минимальное расстояние '
            'источник–детектор f, число экспозиций N и длина участка L. '
            'Коэффициент C = max(2Φ/K, 4). Для чертежа 3в (просветка на эллипс) '
            'L = D×π/4 — на одном снимке одновременно 2 участка шва при N=2. '
            'Для чертежа 3б применяется итерационный алгоритм по ГОСТ 7512-82. '
            'алгоритм по ГОСТ 7512-82. Для чертежа 3ж по п. Г.5 ГОСТ Р 50.05.07-2018 '
            'f и N определяются опытным путём — расчётные значения справочные. '
            'Отрицательное расчётное f в техкарте указывается как 0 мм.'
        ),
        'variables': ['Φ (d)', 'K', 'C', 'f (SFD)', 'N', 'L', 'm', 'D', 'd', 'l'],
        'formula': params.get('scheme_formula') or scheme.get('formula') or 'Зависит от схемы',
        'steps': scheme_steps,
        'result': (
            f'f_min = {scheme.get("f_min_mm", params.get("f_calculated_mm", "—"))} мм, '
            f'N = {scheme.get("N", params.get("N_calculated", "—"))}, '
            f'L = {scheme.get("L_mm", params.get("L_calculated_mm", "—"))} мм'
        ),
        'notes': empirical_note or params.get('scheme_notes') or scheme.get('notes', ''),
    })

    sfd_input = params.get('sfd_mm') or 0
    sfd_used = params.get('sfd_used_mm', '')
    sections.append({
        'title': '8. Геометрическая нерезкость Ug',
        'normative_ref': f'{DOCUMENT_CODE}, п. 6.5; ГОСТ Р 50.05.07-2018',
        'logic': (
            'Геометрическая нерезкость зависит от размера фокусного пятна, '
            'расстояния объект–детектор b и расстояния источник–детектор f. '
            'Если пользователь не задал SFD, используется рассчитанное f_min. '
            'Итоговое SFD = max(SFD_ввод, f_min). Результат сравнивается '
            'с допустимым Ug по требуемой чувствительности K.'
        ),
        'variables': ['Φ (d)', 'b (OFD)', 'f (SFD)', 'Ug'],
        'formula': 'Ug = Φ × b / (f − b)',
        'steps': [
            f'Φ = {params.get("source_focal_spot_mm", "—")} мм',
            f'b (OFD) = {params.get("ofd_mm", "—")} мм',
            f'SFD введённое = {sfd_input or "не задано"} мм',
            f'f_min из схемы = {scheme.get("f_min_mm", params.get("f_calculated_mm", "—"))} мм',
            f'SFD использованное = max(ввод, f_min) = {sfd_used} мм',
            params.get('ug_calculation', ''),
            f'Допустимый Ug_max: {params.get("max_geometric_unsharpness_mm", "—")} мм',
            f'Проверка: {"соответствует" if params.get("geometric_unsharpness_ok") else "ПРЕВЫШЕНО"}',
        ],
        'result': (
            f'Ug = {params.get("geometric_unsharpness_mm", "—")} мм '
            f'(допустимо ≤ {params.get("max_geometric_unsharpness_mm", "—")} мм)'
        ),
        'notes': '',
    })

    placement = params.get('iqi_placement') or {}

    sections.append({
        'title': '9. Плёнка, экраны и ИКИ',
        'normative_ref': (
            f'{DOCUMENT_CODE}, п. 6.1.11; ГОСТ 7512-82, табл. 2; '
            'ГОСТ ИСО 11699-1'
        ),
        'logic': (
            'В техкарте применяется только проволочный ИКИ (ГОСТ 7512-82). '
            'Маркировка: первая цифра — материал (1 — сталь, 2 — Al/Mg, 3 — Ti), '
            'следующие — номер эталона (п. 2.13). '
            'Класс плёнки (ГОСТ ИСО 11699-1) подбирается по категории сварного соединения. '
            'ИКИ по умолчанию — со стороны источника (п. 6.1.11); '
            'при установке со стороны плёнки проволочный эталон подбирается '
            'на одну ступень жёстче относительно требуемой чувствительности K.'
        ),
        'variables': ['K', 'S_рад(f)', 'материал ИКИ'],
        'formula': (
            'Маркировка = код материала + номер эталона (табл. 2); '
            'проволока с d ≤ K'
        ),
        'steps': [
            f'Рекомендуемый класс плёнки: {film.get("class", "—")} — {film.get("description", "")}',
            f'Примеры плёнок: {film.get("examples", "—")}',
            f'Выбранная плёнка: {params.get("film_name", "—")}',
            f'Оптическая плотность: {params.get("optical_density_min", "—")}–'
            f'{params.get("optical_density_max", "—")}',
            f'ИКИ (проволочный): маркировка {params.get("iqi_marking", "—")} '
            f'({params.get("iqi_label", "—")})',
            f'Материал эталона: код {params.get("iqi_material_code", "—")} '
            f'({(params.get("iqi_wire") or {}).get("material_label", "")})',
            f'Сторона установки: {placement.get("side_label", "со стороны источника")}',
            f'Диаметр проволоки №{params.get("iqi_wire_number", "—")}: '
            f'{params.get("iqi_wire_diameter_mm", "—")} мм',
            f'Передний экран: {screens.get("front_mm", "—")} мм ({screens.get("material", "")})',
            f'Задний экран: {screens.get("back_mm", "—")} мм',
        ],
        'result': (
            f'ИКИ: проволочный эталон {params.get("iqi_marking", "—")}, '
            f'{placement.get("side_label", "")}; '
            f'плёнка класса {film.get("class", "—")}'
        ),
        'notes': placement.get('note', '') or screens.get('note', ''),
    })

    sections.append({
        'title': '10. Объём контроля и критерии оценки',
        'normative_ref': f'НП-104-18 (объём); {NP105_CODE} (критерии качества)',
        'logic': (
            'Объём радиографического контроля определяется категорией сварного соединения. '
            'Критерии приёмки дефектов — по НП-105-18 для соответствующей категории.'
        ),
        'variables': [],
        'formula': None,
        'steps': [
            f'Категория: {params.get("weld_category", "—")}',
            f'Объём контроля: {params.get("control_volume_pct", 100)} %',
            f'Норматив оценки: {params.get("quality_normative", NP105_CODE)}',
            params.get('quality_criteria_summary', ''),
        ],
        'result': f'Объём контроля: {params.get("control_volume_pct", 100)} %',
        'notes': '',
    })

    warnings = params.get('warnings') or []
    errors = params.get('errors') or []
    if warnings or errors:
        diag_steps = []
        if errors:
            diag_steps.extend([f'ОШИБКА: {e}' for e in errors])
        if warnings:
            diag_steps.extend([f'ПРЕДУПРЕЖДЕНИЕ: {w}' for w in warnings])
        sections.append({
            'title': '11. Предупреждения и замечания системы',
            'normative_ref': 'Диагностика расчётного модуля',
            'logic': 'Автоматически выявленные несоответствия и замечания при расчёте.',
            'variables': [],
            'formula': None,
            'steps': diag_steps,
            'result': f'{len(errors)} ошибок, {len(warnings)} предупреждений',
            'notes': '',
        })

    return sections


def _add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)


def _add_label_value(doc: Document, label: str, value: str, bold_label: bool = True):
    para = doc.add_paragraph()
    if bold_label:
        run = para.add_run(f'{label}: ')
        run.bold = True
        run.font.size = Pt(9)
    para.add_run(str(value)).font.size = Pt(9)


def generate_calculation_reference_docx(
    input_data: dict,
    params: dict,
    card_number: str = '',
    normative_doc_code: str = DOCUMENT_CODE,
) -> io.BytesIO:
    """
    Генерирует DOCX-документ «Техническая справка (расчёты параметров контроля)».

    :return: BytesIO с содержимым DOCX
    """
    sections = build_calculation_log(input_data, params)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(15)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ТЕХНИЧЕСКАЯ СПРАВКА\n(расчёты параметров контроля)')
    run.bold = True
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f'Технологическая карта № {card_number or "—"}\n'
        f'Методический документ: {normative_doc_code}\n'
        f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}'
    ).font.size = Pt(9)

    doc.add_paragraph()
    _add_heading(doc, 'Расшифровка обозначений в формулах', level=2)
    glossary_table = doc.add_table(rows=1, cols=2)
    glossary_table.style = 'Table Grid'
    hdr = glossary_table.rows[0].cells
    hdr[0].text = 'Обозначение'
    hdr[1].text = 'Описание'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
    for sym, desc in VARIABLE_GLOSSARY:
        row = glossary_table.add_row().cells
        row[0].text = sym
        row[1].text = desc
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph()
    _add_heading(doc, 'Лог расчётов', level=2)

    for sec in sections:
        _add_heading(doc, sec['title'], level=3)
        _add_label_value(doc, 'Нормативная ссылка', sec.get('normative_ref', '—'))
        _add_label_value(doc, 'Логика расчёта', sec.get('logic', '—'))

        if sec.get('variables'):
            _add_label_value(
                doc, 'Переменные',
                ', '.join(sec['variables']),
            )
        if sec.get('formula'):
            _add_label_value(doc, 'Формула', sec['formula'])

        doc.add_paragraph()
        steps_heading = doc.add_paragraph()
        steps_heading.add_run('Ход расчёта:').bold = True
        steps_heading.runs[0].font.size = Pt(9)
        for step in sec.get('steps', []):
            if step:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(str(step)).font.size = Pt(9)

        _add_label_value(doc, 'Результат', sec.get('result', '—'))
        if sec.get('notes'):
            _add_label_value(doc, 'Примечание', sec['notes'])
        doc.add_paragraph()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(
        f'Документ сформирован автоматически системой Карта-НК.\n'
        f'{DOCUMENT_FULL_NAME}'
    ).font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
