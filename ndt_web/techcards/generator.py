"""
Генератор технологических карт радиографического контроля.

Реализует два режима работы:
1. На основе шаблона (DOCX) — заполняет оригинальный бланк данными.
2. Программная генерация — создаёт документ с нуля если шаблон не найден.

Нормативная база: ГОСТ Р 50.05.07-2018, НП-104-18, НП-105-18.
"""

import io
import os
import re
import uuid
import copy
import math
import tempfile
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from common.pdf_fonts import register_cyrillic_fonts, FONT_REGULAR, FONT_BOLD

logger = logging.getLogger(__name__)

from normative.gost_50_05_07 import (
    get_sensitivity, get_sensitivity_mm, get_suitable_sources,
    calc_geometric_unsharpness, calc_min_sfd,
    SCREEN_REQUIREMENTS, FILM_CLASSES, MAX_GEOMETRIC_UNSHARPNESS,
    OPTICAL_DENSITY, PERSONNEL_REQUIREMENTS, SAFETY_REQUIREMENTS,
    FILM_PROCESSING, IQI_TYPES, DOCUMENT_CODE, DOCUMENT_FULL_NAME,
    parse_film_size,
    select_film_size_for_length,
    RADIATION_SOURCES,
)
from normative.gost_59023_2 import (
    resolve_material_type,
    get_material_display_name,
    resolve_welding_material,
    format_dimension_with_tolerance,
)
from normative.np_104_18 import (
    WELD_CATEGORIES,
    DOCUMENT_CODE as NP104_CODE,
    build_titanium_edge_cleaning_requirement,
)
from normative.np_105_18 import (
    DOCUMENT_CODE as NP105_CODE,
    build_acceptance_criteria_docx_data,
    build_root_acceptance_docx_text,
    lookup_root_acceptance_limits,
    resolve_acceptance_table,
)
from normative.calculations import (
    calc_exposure_parameters, recommend_scheme,
    calc_geometric_unsharpness_full, SCHEME_INFO, clamp_f_mm,
    effective_outer_diameter_mm,
    normalize_control_volume_pct,
    requires_full_length_ring_control,
    apply_control_volume_adjustment,
    calc_straight_seam_full_coverage,
)


# ---------------------------------------------------------------
# Вспомогательные функции для работы с таблицами DOCX
# ---------------------------------------------------------------

def _unique_cells(row):
    """
    Возвращает список уникальных ячеек строки (без дублей от слияния).
    В python-docx слитые ячейки возвращаются несколько раз подряд.
    """
    seen = []
    for cell in row.cells:
        if not seen or seen[-1]._tc is not cell._tc:
            seen.append(cell)
    return seen


def _format_date_ddmmyyyy(value) -> str:
    """Форматирует дату для техкарты: дд.мм.гггг (в т.ч. из ISO YYYY-MM-DD)."""
    if value is None or value == '':
        return datetime.now().strftime('%d.%m.%Y')
    if hasattr(value, 'strftime'):
        return value.strftime('%d.%m.%Y')
    s = str(value).strip()
    iso = re.match(r'^(\d{4})-(\d{2})-(\d{2})', s)
    if iso:
        return f'{iso.group(3)}.{iso.group(2)}.{iso.group(1)}'
    return s


def _label_matches_value_key(label_text: str, key: str) -> bool:
    """
    Сопоставляет метку строки шаблона с ключом value_map.
    «1.1» не должно совпадать с «1.10»; «2.2» — с «4.2.2» и т.п.
    """
    if not label_text or not key:
        return False
    label_lower = label_text.lower()
    key_lower = key.lower()
    if key_lower not in label_lower:
        return False
    if not key[0].isdigit():
        return True
    pos = label_lower.find(key_lower)
    if pos > 0 and label_lower[pos - 1] in '.0123456789':
        return False
    tail = label_lower[pos + len(key_lower):]
    if not tail:
        return True
    if tail[0] == '.' and len(tail) > 1 and tail[1].isdigit():
        return False
    return tail[0] in '.\t \xa0'


def _match_value_for_label(label_text: str, value_map: dict):
    """Возвращает значение для метки (приоритет у более длинных ключей)."""
    label_stripped = label_text.strip()
    section_m = re.match(r'^(\d+(?:\.\d+)+)\.?\s', label_stripped)
    if section_m:
        section_key = section_m.group(1).rstrip('.')
        if section_key in value_map:
            return value_map[section_key]

    for key in sorted(value_map.keys(), key=len, reverse=True):
        if not _label_matches_value_key(label_text, key):
            continue
        if key.startswith('4.2.') and key not in ('4.2.3',):
            if any(x in label_text.lower() for x in (
                'внешний диаметр', 'длинна', 'длина', 'наружной поверхности',
                'околошовной', '4.2.6',
            )):
                continue
        if not key[0].isdigit() and section_m:
            continue
        return value_map[key]
    return None


def _reference_run_in_cell(cell):
    """Первый непустой run ячейки — эталон шрифта шаблона."""
    for para in cell.paragraphs:
        for run in para.runs:
            if (run.text or '').strip():
                return run
    for para in cell.paragraphs:
        if para.runs:
            return para.runs[0]
    return None


def _reference_run_in_paragraph(para):
    """Первый непустой run параграфа — эталон шрифта шаблона."""
    for run in para.runs:
        if (run.text or '').strip():
            return run
    return para.runs[0] if para.runs else None


def _copy_run_font(source, target):
    """Копирует оформление run без изменения текста."""
    if source is None or target is None:
        return
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    if source.font.size:
        target.font.size = source.font.size
    if source.font.name:
        target.font.name = source.font.name
    for attr in ('name_ascii', 'name_hansi', 'name_cs', 'name_east_asia'):
        value = getattr(source.font, attr, None)
        if value:
            setattr(target.font, attr, value)


BODY_TABLE_FONT_PT = 12


def _set_cell_text(
    cell,
    text: str,
    bold: bool | None = None,
    font_size: int | None = None,
):
    """
    Устанавливает текст в ячейке, сохраняя шрифт и отступы шаблона.
    По умолчанию — 12 pt (как в эталонной техкарте), а не 9 pt.
    """
    ref = _reference_run_in_cell(cell)
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if para.runs:
        run = para.runs[0]
        run.text = str(text) if text is not None else ''
        for extra in para.runs[1:]:
            extra.text = ''
    else:
        run = para.add_run(str(text) if text is not None else '')
    for extra_para in cell.paragraphs[1:]:
        for extra_run in extra_para.runs:
            extra_run.text = ''
    if bold is not None:
        run.bold = bold
    elif ref is not None:
        run.bold = ref.bold
    if font_size is not None:
        run.font.size = Pt(font_size)
    elif ref is not None and ref.font.size:
        run.font.size = ref.font.size
    else:
        run.font.size = Pt(12)
    if ref is not None:
        _copy_run_font(ref, run)
        if bold is not None:
            run.bold = bold
        if font_size is not None:
            run.font.size = Pt(font_size)
    if font_size is None and (run.font.size is None or run.font.size.pt < BODY_TABLE_FONT_PT):
        run.font.size = Pt(BODY_TABLE_FONT_PT)


def _fill_row_value(table, row_idx: int, value: str):
    """
    Заполняет значение в строке таблицы.
    Предполагает структуру: [метка_ячейки | значение_ячейки].
    Значение записывается в последнюю уникальную ячейку строки.
    """
    if row_idx >= len(table.rows):
        return
    row = table.rows[row_idx]
    ucells = _unique_cells(row)
    if len(ucells) >= 2:
        _set_cell_text(ucells[-1], value)
    elif len(ucells) == 1:
        # Однострочная таблица — добавляем после метки
        pass


def _find_row_by_label(table, label_fragment: str) -> int:
    """Находит индекс строки по фрагменту текста метки."""
    for i, row in enumerate(table.rows):
        ucells = _unique_cells(row)
        if ucells and label_fragment.lower() in ucells[0].text.lower():
            return i
    return -1


def _format_sensitivity_desc(
    display_k_mm: float,
    sk_desc: str,
    weld_category: str,
    *,
    norm_k_mm: float | None = None,
    film_side: bool = False,
) -> str:
    """
    Текст для п. 6.3 техкарты «Требуемая чувствительность K, не более, мм».

    При ИКИ со стороны плёнки указывается K на одну ступень жёстче
    нормативного значения (диаметр проволоки ИКИ).
    """
    if film_side and norm_k_mm is not None:
        return (
            f'K ≤ {display_k_mm:.3f} мм '
            f'(ИКИ со стороны плёнки, на 1 ступень жёстче нормативного '
            f'K ≤ {norm_k_mm:.3f} мм; НП-105-18, Табл. 4.8: {sk_desc}, '
            f'кат. {weld_category})'
        )
    return (
        f'K ≤ {display_k_mm:.3f} мм '
        f'(НП-105-18, Табл. 4.8: {sk_desc}, кат. {weld_category})'
    )


def _format_quality_norm_cite(params: dict) -> str:
    """П. 2.2 — нормативный документ оценки качества."""
    from normative.snip_3_05_05_84 import is_snip_quality_norm

    qn = (params.get('quality_norm_code') or NP105_CODE).strip()
    if is_snip_quality_norm(qn):
        return qn
    if qn.startswith('НП-105'):
        return f'{qn}; {NP104_CODE}'
    return qn


def _format_weld_category_field(params: dict, method_cite: str = '') -> str:
    """П. 3.1 — категория сварного соединения / трубопровода."""
    from normative.snip_3_05_05_84 import (
        is_snip_quality_norm,
        get_pipeline_category_info,
    )

    weld_cat = params.get('weld_category', '')
    qn = params.get('quality_norm_code') or NP105_CODE
    if is_snip_quality_norm(qn):
        info = get_pipeline_category_info(weld_cat)
        return f'{info["name"]} (по СНиП 3.05.05-84)'
    if '7512' in (method_cite or ''):
        return f'{weld_cat} (по ТД / НП-105-18)'
    return f'{weld_cat} (по НП-105-18)'


# ---------------------------------------------------------------
# Расчётное ядро (без изменений)
# ---------------------------------------------------------------

class RadiographicTechCardCalculator:
    """
    Вычислительное ядро техкарты РГК.

    Методика (ГОСТ 7512-82 или ГОСТ Р 50.05.07-2018) задаётся полем
    ``doc_code`` из выбранной кнопки документа. Геометрия схем общая;
    подписи чертежей и ссылки в карте — по выбранной методике.
    """

    def __init__(self, input_data: dict):
        self.data = input_data
        self.params = {}
        self.errors = []
        self.warnings = []
        from techcards.methodology import get_methodology
        self.methodology = get_methodology(input_data.get('doc_code'))

    def calculate(self) -> dict:
        """Выполняет все расчёты и возвращает словарь параметров."""
        self._extract_inputs()
        self.params['doc_code'] = self.methodology.code
        self.params['methodology_code'] = self.methodology.code
        self.params['method_doc_cite'] = self.methodology.method_doc_cite
        qn = (self.data.get('quality_norm_code') or '').strip() or NP105_CODE
        self.params['quality_norm_code'] = qn
        self.params['quality_norm_name'] = (
            self.data.get('quality_norm_name') or qn
        )
        # Сначала получаем g_min/g_max из таблиц сварных соединений
        self._calc_inspection_zones()
        # Нормативное K по НП-105 (S_K), затем ужесточение по п. 6.1.11 при ИКИ
        # со стороны плёнки (ступень = диаметр проволоки ГОСТ 7512).
        self._calc_sensitivity_value()
        self._apply_iqi_side_to_sensitivity()
        self._select_sources()
        self._calc_geometric_params()
        self._calc_exposure_scheme()
        self._apply_control_volume()
        self._select_film()
        self._calc_screens()
        self._calc_iqi()
        self._fill_processing()
        self._fill_personnel()
        self._fill_safety()
        self._fill_acceptance_criteria()
        self._calc_control_volume()
        return self.params

    def _extract_inputs(self):
        d = self.data
        source_code = d.get('source_code', '')
        src_info = next((s for s in RADIATION_SOURCES if s['code'] == source_code), None)
        src_type = src_info.get('type', 'isotope') if src_info else 'isotope'
        focal_raw = d.get('focal_spot_mm')
        if focal_raw in (None, ''):
            focal_spot = 3.0 if src_type == 'isotope' else 2.0
        else:
            focal_spot = float(focal_raw)

        develop_date = _format_date_ddmmyyyy(d.get('develop_date', ''))
        raw_check = d.get('check_date', '')
        if (
            raw_check is None
            or raw_check == ''
            or (isinstance(raw_check, str) and not raw_check.strip())
        ):
            check_date = develop_date
        else:
            check_date = _format_date_ddmmyyyy(raw_check)

        material_display = get_material_display_name(
            d.get('material', ''),
            d.get('material_custom', ''),
        )

        self.params.update({
            'organization': d.get('organization', ''),
            'object_name': d.get('object_name', ''),
            'drawing_number': d.get('drawing_number', ''),
            'weld_number': d.get('weld_number', ''),
            'card_number': d.get('card_number', ''),
            'object_type': d.get('object_type', 'pipe'),
            'material': d.get('material', ''),
            'material_grade': (d.get('material_custom') or '').strip(),
            'material_display': material_display,
            'material_type': resolve_material_type(d.get('material', '')),
            'wall_thickness': float(d.get('wall_thickness', 10)),
            'outer_diameter': float(d.get('outer_diameter', 0) or 0),
            'flat_length_mm': float(d.get('flat_length_mm', 0) or 0),
            'joint_designation': d.get('joint_designation', ''),
            'joint_mobility': d.get('joint_mobility', 'non_rotating'),
            'welding_process': d.get('welding_process', '30'),
            'weld_category': d.get('weld_category', 'II'),
            'control_volume_pct': self._normalize_volume_for_quality_norm(
                d.get('control_volume_pct', 100),
                d.get('quality_norm_code'),
            ),
            'reinforcement_removed': bool(d.get('reinforcement_removed')),
            'has_backing_ring': bool(d.get('has_backing_ring')),
            'backing_ring_thickness_mm': float(d.get('backing_ring_thickness_mm', 0) or 0),
            'weld_material': resolve_welding_material(
                d.get('welding_material', ''),
                d.get('welding_material_custom', ''),
                material_display,
            ),
            'source_code': source_code,
            'source_focal_spot_mm': focal_spot,
            'source_activity': d.get('source_activity', ''),
            'sfd_mm': float(d.get('sfd_mm', 0) or 0),
            'ofd_mm': float(d.get('ofd_mm', 5) or 5),
            'film_name': d.get('film_name', ''),
            'iqi_side': d.get('iqi_side', 'source') or 'source',
            'inspector_name': d.get('inspector_name', ''),
            'develop_date': develop_date,
            'check_date': check_date,
            'developed_by_position': (
                d.get('developed_by_position_resolved')
                or (d.get('developed_by_position_custom', '').strip()
                    if d.get('developed_by_position') == '__custom__'
                    else d.get('developed_by_position', ''))
            ),
            'developed_by_name': d.get('developed_by_name', ''),
            'developed_by_certificate': d.get('developed_by_certificate', ''),
            'checked_by_position': (
                d.get('checked_by_position_resolved')
                or (d.get('checked_by_position_custom', '').strip()
                    if d.get('checked_by_position') == '__custom__'
                    else d.get('checked_by_position', ''))
            ),
            'checked_by_name': d.get('checked_by_name', ''),
            'checked_by_certificate': d.get('checked_by_certificate', ''),
            'gmo_checked_by_position': (
                d.get('gmo_checked_by_position_resolved')
                or (d.get('gmo_checked_by_position_custom', '').strip()
                    if d.get('gmo_checked_by_position') == '__custom__'
                    else d.get('gmo_checked_by_position', ''))
            ),
            'gmo_checked_by_name': d.get('gmo_checked_by_name', ''),
            'gmo_checked_by_certificate': d.get('gmo_checked_by_certificate', ''),
            'gmo_check_date': _format_date_ddmmyyyy(d.get('gmo_check_date', '')),
            'department': (d.get('department') or '').strip(),
            'control_location': (d.get('control_location') or '').strip(),
            'source_name_override': (d.get('source_name_override') or '').strip(),
            'temperature_range': (d.get('temperature_range') or '').strip() or '+5 ÷ +40',
            'tube_voltage_kv': (
                float(d['tube_voltage_kv'])
                if d.get('tube_voltage_kv') not in (None, '')
                else None
            ),
            'assessment_thickness_mm': (
                float(d['assessment_thickness_mm'])
                if d.get('assessment_thickness_mm') not in (None, '')
                else None
            ),
        })

    def _calc_inspection_zones(self):
        """
        Рассчитывает ширину валика шва, ОШЗ и контролируемую зону.
        НП-105 / АЭУ: ГОСТ Р 59023.2-2020; СНиП: ГОСТ 16037-80.
        Сохраняет g_min, g_max для расчёта радиационной толщины в схеме.
        Заполняет поля техкарты 4.2.2, 4.2.4, 4.2.5.
        """
        from normative.snip_3_05_05_84 import is_snip_quality_norm
        from normative.gost_16037_80 import (
            is_gost_16037_joint,
            get_inspection_zone as get_zone_16037,
            get_joint_info as get_joint_info_16037,
            get_joint_image_path as get_image_16037,
        )
        from normative.gost_59023_2 import get_inspection_zone, get_joint_info

        joint_code = self.params.get('joint_designation', '')
        S = self.params['wall_thickness']
        qn = self.params.get('quality_norm_code') or self.data.get('quality_norm_code')
        use_16037 = is_snip_quality_norm(qn) or is_gost_16037_joint(joint_code)
        method = self.params.get('welding_process', 'Р' if use_16037 else '30')
        dn = self.params.get('outer_diameter')
        s1_override = self.params.get('s1_mm')
        if s1_override in (None, ''):
            s1_override = None

        if use_16037:
            zone = get_zone_16037(
                joint_code, S, method or 'Р',
                material_type=self.params.get('material_type', 'steel'),
                outer_diameter_mm=float(dn) if dn not in (None, '') else None,
                s1_override_mm=float(s1_override) if s1_override is not None else None,
                reinforcement_removed=self.params.get('reinforcement_removed', False),
                has_backing_ring=self.params.get('has_backing_ring', False),
                backing_ring_thickness_mm=self.params.get('backing_ring_thickness_mm') or None,
            )
            joint_info = get_joint_info_16037(joint_code)
            self.params['joint_standard'] = 'ГОСТ 16037-80'
        else:
            zone = get_inspection_zone(
                joint_code, S, method,
                material_type=self.params.get('material_type', 'steel'),
                outer_diameter_mm=float(dn) if dn not in (None, '') else None,
                s1_override_mm=float(s1_override) if s1_override is not None else None,
                reinforcement_removed=self.params.get('reinforcement_removed', False),
                has_backing_ring=self.params.get('has_backing_ring', False),
                backing_ring_thickness_mm=self.params.get('backing_ring_thickness_mm') or None,
            )
            joint_info = get_joint_info(joint_code)
            self.params['joint_standard'] = 'ГОСТ Р 59023.2-2020'

        self.params['weld_bead_width_mm'] = zone.get('bead_width_mm', '')
        self.params['weld_bead_width_inner_mm'] = zone.get('bead_width_inner_mm', '')
        self.params['e_display'] = zone.get('e_display', '')
        self.params['e1_display'] = zone.get('e1_display', '')
        self.params['g_display'] = zone.get('g_display', '')
        self.params['g_label'] = zone.get(
            'g_label', '4.2.3. Высота валика усиления (g)',
        )
        self.params['weld_bead_height_mm'] = zone.get('bead_height_mm', '')
        self.params['reinforcement_status'] = (
            'снят' if self.params.get('reinforcement_removed') else 'не снят'
        )
        self.params['g_min_mm'] = zone.get('g_min_mm', 0.5)
        self.params['g_max_mm'] = zone.get('g_max_mm', 3.5)
        self.params['backing_thickness_mm'] = zone.get('backing_thickness_mm', 0.0)
        self.params['has_backing'] = zone.get('has_backing', False)
        self.params['haz_width_mm'] = zone.get('haz_width_mm', 5.0)
        self.params['zone_width_mm'] = zone.get('zone_width_mm', '')
        self.params['film_width_min_mm'] = zone.get('film_width_min_mm', '')
        self.params['zone_note'] = zone.get('weld_note', '')
        self.params['joint_groove'] = joint_info.get('groove', '')
        self.params['joint_name'] = joint_info.get('name', joint_code)
        self.params['joint_sketch'] = joint_info.get('sketch', '')
        if use_16037:
            self.params['joint_image'] = get_image_16037(joint_code)
            self.params['joint_image_static_root'] = True
        else:
            from normative.gost_59023_2 import get_joint_image_path
            self.params['joint_image'] = get_joint_image_path(joint_code)
            self.params['joint_image_static_root'] = False

        # S / S1 по ГОСТ Р 59023.2-2020 (S_eff для радиационной толщины)
        self.params['s_mm'] = zone.get('s_mm', S)
        self.params['s1_mm'] = zone.get('s1_mm', S)
        self.params['s_eff_mm'] = zone.get('s_eff_mm', S)
        self.params['dp_mm'] = zone.get('dp_mm')
        self.params['s_equals_s1'] = zone.get('s_equals_s1', True)
        self.params['s_equals_s1_actual'] = zone.get('s_equals_s1_actual', True)
        self.params['has_internal_boring'] = zone.get('has_internal_boring', False)
        self.params['wall_thickness_mode'] = zone.get('wall_thickness_mode', 's_equals_s1')
        self.params['wall_summary'] = zone.get('wall_summary', f'S = {S} мм')
        self.params['wall_note'] = zone.get('wall_note', '')

    def _calc_sensitivity_value(self):
        """
        Требуемая чувствительность K по выбранному нормативному документу.

        НП-105-18: табл. 4.8–4.11, S_K с учётом g_min / двух стенок.
        СНиП 3.05.05-84: класс 2/3 по ГОСТ 7512-82 (п. 4.13) → абсолютное K.
        """
        from normative.calculations import calc_radiation_thickness
        from normative.snip_3_05_05_84 import (
            is_snip_quality_norm,
            get_required_sensitivity_mm as snip_get_k,
            get_sensitivity_class,
            get_pipeline_category_info,
        )

        # Для расточенных соединений (S ≠ S1) в путь излучения идёт S1
        S_nom = self.params['wall_thickness']
        S = float(self.params.get('s_eff_mm') or S_nom)
        weld_cat = self.params['weld_category']
        scheme = self.data.get('scheme_type', '4_6')
        g_min = self.params.get('g_min_mm', 0.5)
        g_max = self.params.get('g_max_mm', 3.5)
        s_pk = self.params.get('backing_thickness_mm', 0.0)
        qn = self.params.get('quality_norm_code') or NP105_CODE

        rad = calc_radiation_thickness(S, g_min, g_max, scheme, s_pk)
        self.params['rad_thickness'] = rad

        s_k = rad['s_rad_k_mm']
        s_label = 'S1' if not self.params.get('s_equals_s1_actual', True) else 'S'
        if rad['wall_count'] == 2:
            sk_desc = f'S_K = {s_label} + {s_label} = {S} + {S} = {s_k} мм'
        elif s_pk > 0:
            sk_desc = (
                f'S_K = {s_label} + g_min + Sпк = {S} + {g_min} + {s_pk} = {s_k} мм'
            )
        else:
            sk_desc = f'S_K = {s_label} + g_min = {S} + {g_min} = {s_k} мм'
        if s_label == 'S1':
            sk_desc += f' (номинальная S = {S_nom} мм; {self.params.get("wall_summary", "")})'

        if is_snip_quality_norm(qn):
            from normative.snip_3_05_05_84 import DOCUMENT_CODE as SNIP_CODE

            mm_val = snip_get_k(s_k, weld_cat)
            sens_cls = get_sensitivity_class(weld_cat)
            cat_info = get_pipeline_category_info(weld_cat)
            self.params['required_sensitivity_pct'] = None
            self.params['required_sensitivity_mm'] = mm_val
            self.params['required_sensitivity_norm_mm'] = mm_val
            self.params['s_k_mm'] = s_k
            self.params['s_rad_f_mm'] = rad['s_rad_f_mm']
            self.params['sk_desc'] = sk_desc
            self.params['sensitivity_k_display_mm'] = mm_val
            self.params['sensitivity_class_gost7512'] = sens_cls
            self.params['sensitivity_desc'] = (
                f'K ≤ {mm_val:g} мм (класс {sens_cls} по ГОСТ 7512-82; '
                f'{SNIP_CODE}, п. 4.13; {cat_info["name"]}; {sk_desc})'
            )
            return

        K = get_sensitivity(s_k, weld_cat)
        mm_val = get_sensitivity_mm(s_k, weld_cat)

        self.params['required_sensitivity_pct'] = K
        self.params['required_sensitivity_mm'] = mm_val
        self.params['required_sensitivity_norm_mm'] = mm_val
        self.params['s_k_mm'] = s_k
        self.params['s_rad_f_mm'] = rad['s_rad_f_mm']
        self.params['sk_desc'] = sk_desc
        self.params['sensitivity_k_display_mm'] = mm_val
        self.params['sensitivity_desc'] = _format_sensitivity_desc(
            mm_val, sk_desc, self.params.get('weld_category', ''),
        )

    def _apply_iqi_side_to_sensitivity(self):
        """
        ГОСТ Р 50.05.07-2018, п. 6.1.11.

        ИКИ по умолчанию — со стороны источника. Если ИКИ со стороны плёнки
        (двухстеночное / панорамное просвечивание), чувствительность РК
        смещается на одну ступень ИКИ жёстче. Ступень = диаметр проволоки
        проволочного эталона по ГОСТ 7512-82, табл. 2.

        Ужесточённый K (диаметр проволоки) используется для C, f, N, L.
        Нормативное K по НП-105 сохраняется в required_sensitivity_norm_mm.
        """
        from normative.gost_7512 import get_wire_iqi, resolve_iqi_placement

        iqi_side = self.params.get('iqi_side', 'source') or 'source'
        scheme = self.params.get('scheme_type', self.data.get('scheme_type', '4_6'))
        placement = resolve_iqi_placement(
            scheme, wall_count=0, iqi_side=iqi_side,
            doc_code=self.methodology.code,
        )
        self.params['iqi_placement'] = placement
        self.params['iqi_side'] = iqi_side

        norm_k = float(self.params.get('required_sensitivity_norm_mm')
                       or self.params.get('required_sensitivity_mm')
                       or 0)
        if not placement.get('shift_steps'):
            self.params['required_sensitivity_mm'] = norm_k
            self.params['sensitivity_k_display_mm'] = norm_k
            return

        rad_f = self.params.get('s_rad_f_mm', self.params['wall_thickness'])
        iqi_wire = get_wire_iqi(
            rad_f,
            norm_k,
            material_type=self.params.get('material_type', 'steel'),
            shift_steps=int(placement['shift_steps']),
        )
        k_eff = float(iqi_wire['wire_diameter_mm'])
        self.params['required_sensitivity_mm'] = k_eff
        self.params['sensitivity_k_display_mm'] = k_eff
        self.params['iqi_wire_preview'] = iqi_wire
        sk_desc = self.params.get('sk_desc', '')
        weld_cat = self.params.get('weld_category', '')
        self.params['sensitivity_desc'] = _format_sensitivity_desc(
            k_eff, sk_desc, weld_cat,
            norm_k_mm=norm_k, film_side=True,
        )

    def _select_sources(self):
        S = self.params['wall_thickness']
        material_type = self.params.get('material_type', 'steel')
        rad = self.params.get('rad_thickness') or {}
        table_b_thickness = rad.get('s_rad_f_mm', S)
        suitable = get_suitable_sources(table_b_thickness, material_type)
        self.params['suitable_sources'] = suitable
        chosen_code = self.params['source_code']
        if chosen_code:
            match = next((s for s in suitable if s['code'] == chosen_code), None)
            if match:
                self.params['selected_source'] = match
            else:
                # Явный выбор пользователя сохраняем в карте (как в промышленных бланках),
                # даже если табл. Б рекомендует другой источник — с предупреждением.
                catalog = next(
                    (s for s in RADIATION_SOURCES if s['code'] == chosen_code),
                    None,
                )
                self.warnings.append(
                    f'Источник {chosen_code} не допускается табл. '
                    f'Б.{ {"steel": "1", "aluminum": "2", "titanium": "3"}.get(material_type, "1") } '
                    f'для радиационной толщины {table_b_thickness} мм и выбранного материала.'
                )
                if catalog:
                    selected = dict(catalog)
                    selected['table_ref'] = 'выбран пользователем'
                    self.params['selected_source'] = selected
                else:
                    self.params['selected_source'] = suitable[0] if suitable else {}
        else:
            self.params['selected_source'] = suitable[0] if suitable else {}

    def _calc_geometric_params(self):
        """
        Рассчитывает геометрическую нерезкость.
        Если f (SFD) не задан пользователем — использует f_min из расчёта схемы.
        """
        focal = self.params['source_focal_spot_mm']
        ofd = self.params['ofd_mm']
        weld_cat = self.params['weld_category']
        sensitivity_mm = self.params.get('required_sensitivity_mm', 0)

        # Если SFD не задан вручную — используем рассчитанное f_min
        sfd = self.params.get('sfd_mm') or 0
        scheme_result = self.params.get('exposure_scheme') or {}
        f_min = scheme_result.get('f_min_mm') or 0

        # Итоговое SFD — максимум из введённого и рассчитанного минимума
        sfd_used = max(sfd, f_min) if f_min else (sfd or 700)
        self.params['sfd_used_mm'] = round(sfd_used, 1)

        # Расчёт Ug через модуль calculations (с проверкой ГОСТ)
        ug_result = calc_geometric_unsharpness_full(focal, ofd, sfd_used, sensitivity_mm)

        if ug_result.get('error'):
            self.errors.append(ug_result['error'])
            ug = 0
            ug_ok = False
        else:
            ug = ug_result['ug_mm']
            ug_ok = ug_result.get('gost_ok', True)
            if ug_ok is False:
                max_allowed = ug_result.get('max_allowed_mm', '')
                self.warnings.append(
                    f'Геометрическая нерезкость Ug = {ug:.3f} мм '
                    f'превышает допустимую {max_allowed} мм для K = {sensitivity_mm:.3f} мм.'
                )

        max_ug = ug_result.get('max_allowed_mm') or MAX_GEOMETRIC_UNSHARPNESS.get(weld_cat, 0.5)
        self.params['geometric_unsharpness_mm'] = ug
        self.params['max_geometric_unsharpness_mm'] = max_ug
        self.params['geometric_unsharpness_ok'] = ug <= max_ug
        self.params['ug_calculation'] = ug_result.get('formula', f'Ug = {ug:.3f} мм')

    def _calc_exposure_scheme(self):
        """
        Вычисляет точные параметры просвечивания (f, N, L) по выбранной схеме.
        Использует алгоритмы из normative.calculations (порт из KalkulatorRK2).
        """
        S_nom = self.params['wall_thickness']
        # Для расточки: путь излучения и d по S1 / Dр (табл. 9.30)
        S = float(self.params.get('s_eff_mm') or S_nom)
        D_nom = self.params['outer_diameter']
        dp = self.params.get('dp_mm')
        if dp not in (None, ''):
            d_inner = float(dp)
        else:
            d_inner = D_nom - 2 * S if D_nom else 0
        self.params['d_inner_mm'] = round(d_inner, 1)

        focal = self.params['source_focal_spot_mm']
        sensitivity_mm = self.params.get('required_sensitivity_mm', 0.5)
        g_max = self.params.get('g_max_mm', 3.5)

        # Схема просвечивания — выбор пользователя или авторекомендация
        scheme = self.data.get('scheme_type', '').strip()
        if not scheme:
            recommended = recommend_scheme(D_nom, d_inner)
            scheme = recommended[0] if recommended else '4_6'
            self.params['scheme_auto_selected'] = True

        d_outer = effective_outer_diameter_mm(D_nom, g_max, scheme)
        self.params['d_outer_effective_mm'] = d_outer

        sens = sensitivity_mm if sensitivity_mm > 0 else 0.5
        calc_kwargs = dict(
            scheme=scheme,
            focal_spot_mm=focal,
            sensitivity_mm=sens,
            thickness_mm=S,
            d_outer_mm=d_outer or 0,
            d_inner_mm=d_inner,
        )

        if scheme == '5b':
            # Длина плёнки — вход 5б; сначала оценка L с максимальным типовым размером
            initial_length = float(parse_film_size('480x100')['length_mm'])
            calc_result = calc_exposure_parameters(
                **calc_kwargs, film_length_mm=initial_length,
            )
            film_size = select_film_size_for_length(calc_result.get('L_mm'))
            film_length = float(film_size['length_mm'])
            if film_length != initial_length:
                calc_result = calc_exposure_parameters(
                    **calc_kwargs, film_length_mm=film_length,
                )
                film_size = select_film_size_for_length(calc_result.get('L_mm'))
                film_length = float(film_size['length_mm'])
        else:
            calc_result = calc_exposure_parameters(**calc_kwargs)
            film_size = select_film_size_for_length(calc_result.get('L_mm'))

        film_length = float(film_size['length_mm'])
        film_width = float(film_size['width_mm'])

        if calc_result.get('error'):
            self.warnings.append(f'Схема {scheme}: {calc_result["error"]}')

        if scheme == '4_6' and calc_result.get('L_mm') is None:
            zone_width = self.params.get('zone_width_mm')
            if zone_width:
                calc_result['L_mm'] = round(float(zone_width), 0)

        # Для прямолинейных швов (чертёж 2): N при 100 % охвате по длине шва
        if scheme == '4_6':
            seam_length = self._get_seam_length_mm()
            segment_l = calc_result.get('L_mm')
            if seam_length and segment_l:
                n_full = calc_straight_seam_full_coverage(seam_length, segment_l)
                calc_result['N'] = n_full
                calc_result['N_segments'] = n_full

        # Информация о схеме: имена чертежей по выбранной методике
        from techcards.scheme_display import resolve_scheme_info_for_display
        scheme_info = resolve_scheme_info_for_display(
            scheme, self.params.get('doc_code') or self.methodology.code,
        )

        # Флаг «опытным путём» (50.05.07 п. Г.5 / аналогичные случаи 7512 прил. 4)
        is_empirical = calc_result.get('is_empirical', False)
        empirical_reason = calc_result.get('empirical_reason', '')

        self.params['scheme_type'] = scheme
        self.params['exposure_scheme'] = calc_result
        self.params['scheme_info'] = scheme_info
        self.params['is_empirical'] = is_empirical
        self.params['empirical_reason'] = empirical_reason

        if is_empirical:
            # Расчётные значения для справки (или None)
            self.params['f_calculated_mm'] = clamp_f_mm(calc_result.get('f_min_mm'))
            self.params['N_calculated'] = calc_result.get('N', '')
            self.params['N_segments'] = calc_result.get('N_segments')
            self.params['L_calculated_mm'] = calc_result.get('L_mm')
        else:
            self.params['f_calculated_mm'] = clamp_f_mm(calc_result.get('f_min_mm'))
            self.params['N_calculated'] = calc_result.get('N', '')
            self.params['N_segments'] = calc_result.get('N_segments')
            self.params['L_calculated_mm'] = calc_result.get('L_mm')

        self.params['scheme_formula'] = calc_result.get('formula', '')
        self.params['scheme_notes'] = calc_result.get('notes', '')
        self.params['scheme_image'] = scheme_info.get('image', '')
        self.params['C_coeff'] = calc_result.get('C', '')
        self.params['film_size_code'] = film_size['code']
        self.params['film_length_mm'] = film_length
        self.params['film_width_mm'] = film_width
        self.params['film_size_label'] = film_size['label']

    def _select_film(self):
        weld_cat = self.params['weld_category']
        recommended = [f for f in FILM_CLASSES if weld_cat in f['allowed_for']]
        self.params['recommended_film_classes'] = recommended
        self.params['film_class_info'] = recommended[0] if recommended else {}
        od = OPTICAL_DENSITY.get(weld_cat, OPTICAL_DENSITY['III'])
        self.params['optical_density_min'] = od['min']
        self.params['optical_density_max'] = od['max']

    def _calc_screens(self):
        source_code = (self.params.get('selected_source') or {}).get('code', '')
        screens = SCREEN_REQUIREMENTS.get(source_code, {
            'front_mm': '0,10', 'back_mm': '0,20',
            'material': 'Свинцовые (Pb)', 'note': '',
        })
        self.params['screens'] = screens

    def _calc_iqi(self):
        from normative.gost_7512 import get_wire_iqi, resolve_iqi_placement, WIRE_IQI_TYPE

        self.params['recommended_iqi'] = dict(WIRE_IQI_TYPE)

        # Нормативное K (НП-105) — база для выбора проволоки до сдвига ступени.
        norm_k = float(
            self.params.get('required_sensitivity_norm_mm')
            or self.params.get('required_sensitivity_mm')
            or 0
        )
        scheme = self.params.get('scheme_type', self.data.get('scheme_type', '4_6'))
        material_type = self.params.get('material_type', 'steel')
        iqi_side = self.params.get('iqi_side', 'source')
        rad_f = self.params.get('s_rad_f_mm', self.params['wall_thickness'])

        placement = self.params.get('iqi_placement') or resolve_iqi_placement(
            scheme, wall_count=0, iqi_side=iqi_side,
            doc_code=self.methodology.code,
        )
        self.params['iqi_placement'] = placement
        self.params['iqi_side'] = iqi_side

        iqi_wire = self.params.get('iqi_wire_preview') or get_wire_iqi(
            rad_f, norm_k,
            material_type=material_type,
            shift_steps=placement['shift_steps'],
        )
        self.params['iqi_wire'] = iqi_wire
        self.params['iqi_marking'] = iqi_wire['marking']
        self.params['iqi_material_code'] = iqi_wire['material_code']
        self.params['iqi_set_number'] = iqi_wire['set_number']
        self.params['iqi_wire_number'] = iqi_wire['wire_number']
        self.params['iqi_wire_diameter_mm'] = iqi_wire['wire_diameter_mm']
        self.params['iqi_label'] = iqi_wire['label']

        sk_desc = self.params.get('sk_desc', '')
        weld_cat = self.params.get('weld_category', '')
        film_side = bool(placement.get('shift_steps'))

        if film_side:
            display_k = float(iqi_wire['wire_diameter_mm'])
            self.params['sensitivity_k_display_mm'] = display_k
            # Геометрия уже посчитана по ужесточённому K в _apply_iqi_side_to_sensitivity.
            self.params['required_sensitivity_mm'] = display_k
            self.params['sensitivity_desc'] = _format_sensitivity_desc(
                display_k, sk_desc, weld_cat,
                norm_k_mm=norm_k, film_side=True,
            )
        else:
            self.params['sensitivity_k_display_mm'] = norm_k
            self.params['required_sensitivity_mm'] = norm_k
            self.params['sensitivity_desc'] = _format_sensitivity_desc(
                norm_k, sk_desc, weld_cat,
            )
    def _fill_processing(self):
        self.params['film_processing'] = FILM_PROCESSING

    def _fill_personnel(self):
        weld_cat = self.params['weld_category']
        self.params['personnel_requirements'] = PERSONNEL_REQUIREMENTS.get(
            weld_cat, PERSONNEL_REQUIREMENTS['III']
        )

    def _fill_safety(self):
        self.params['safety_requirements'] = SAFETY_REQUIREMENTS

    def _fill_acceptance_criteria(self):
        from normative.snip_3_05_05_84 import (
            is_snip_quality_norm,
            build_acceptance_criteria_docx_data as snip_acceptance_docx,
            build_quality_criteria_summary as snip_quality_summary,
            undercut_limit_mm,
            surface_ndt_before_rt_required,
            DOCUMENT_CODE as SNIP_CODE,
        )

        weld_cat = self.params['weld_category']
        material_type = self.params.get('material_type', 'steel')
        wall = float(self.params.get('wall_thickness', 0) or 0)
        assess_raw = self.params.get('assessment_thickness_mm')
        thickness = float(assess_raw) if assess_raw not in (None, '') else wall
        self.params['assessment_thickness_used_mm'] = thickness
        qn = self.params.get('quality_norm_code') or NP105_CODE

        if is_snip_quality_norm(qn):
            docx_data = snip_acceptance_docx(material_type, weld_cat, thickness)
            self.params['acceptance_table_ref'] = docx_data.get('table_ref', 'прил. 4')
            self.params['acceptance_criteria_docx'] = docx_data
            self.params['quality_normative'] = SNIP_CODE
            self.params['quality_criteria_summary'] = snip_quality_summary(
                weld_cat, thickness,
            )
            self.params['root_acceptance_limits'] = None
            uc = undercut_limit_mm(weld_cat)
            uc_txt = (
                'подрезы не допускаются'
                if uc <= 0
                else f'подрезы не более {uc:g} мм'
            )
            root_txt = (
                f'По {SNIP_CODE}, п. 4.10: трещины, прожоги, незаваренные кратеры, '
                f'грубая чешуйчатость — не допускаются; {uc_txt}.'
            )
            if surface_ndt_before_rt_required(weld_cat):
                root_txt += (
                    ' До РК/УЗК — МПД или КК поверхности шва и зоны +20 мм '
                    'с каждой стороны.'
                )
            self.params['root_acceptance_docx'] = root_txt
            return

        table_ref = resolve_acceptance_table(material_type, weld_cat)
        docx_data = build_acceptance_criteria_docx_data(
            material_type, weld_cat, thickness,
        )
        self.params['acceptance_table_ref'] = table_ref
        self.params['acceptance_criteria_docx'] = docx_data
        self.params['quality_normative'] = NP105_CODE
        self.params['quality_criteria_summary'] = (
            f'По {NP105_CODE}, категория {weld_cat}. '
            f'Трещины, несплавления, непровары — не допускаются '
            f'(п. 23 приложения 4). '
            f'Поры и шлаковые включения — по таблице N {table_ref}.'
        )
        root_limits = lookup_root_acceptance_limits(
            wall,
            float(self.params.get('outer_diameter') or 0),
            joint_mobility=self.params.get('joint_mobility', 'non_rotating'),
            has_backing=bool(self.params.get('has_backing')),
        )
        self.params['root_acceptance_limits'] = root_limits
        self.params['root_acceptance_docx'] = build_root_acceptance_docx_text(root_limits)

    def _get_seam_length_mm(self) -> float:
        """Длина прямолинейного шва для схемы 4.6 (мм)."""
        object_type = self.params.get('object_type', 'pipe')
        if object_type == 'flat':
            return float(self.params.get('flat_length_mm') or 0)
        if object_type == 'vessel':
            flat = float(self.params.get('flat_length_mm') or 0)
            if flat > 0:
                return flat
            d = float(self.params.get('outer_diameter') or 0)
            return math.pi * d if d > 0 else 0.0
        return 0.0

    def _apply_control_volume(self):
        """
        Корректирует N и N_segments по объёму выборочного контроля (НП-105-18, п. 70–72).
        """
        volume_pct = self.params.get('control_volume_pct', 100)
        object_type = self.params.get('object_type', 'pipe')
        d_nom = float(
            self.params.get('outer_diameter')
            or self.params.get('d_outer_effective_mm')
            or 0
        )
        scheme = self.params.get('scheme_type', '')

        n_full = self.params.get('N_calculated')
        n_seg_full = self.params.get('N_segments') or n_full
        full_length_ring = requires_full_length_ring_control(object_type, d_nom)

        seam_length = self._get_seam_length_mm() if scheme == '4_6' else None
        segment_l = self.params.get('L_calculated_mm') if scheme == '4_6' else None

        n_adj, n_seg_adj, controlled_length = apply_control_volume_adjustment(
            N_full=n_full,
            N_segments_full=n_seg_full,
            volume_pct=volume_pct,
            apply_sample_scaling=not full_length_ring,
            seam_length_mm=seam_length,
            segment_length_mm=segment_l,
        )

        self.params['N_calculated_full'] = n_full
        self.params['N_segments_full'] = n_seg_full
        self.params['N_calculated'] = n_adj
        self.params['N_segments'] = n_seg_adj
        self.params['control_volume_applied'] = not full_length_ring and volume_pct < 100
        self.params['control_volume_mode'] = (
            'full_length_ring' if full_length_ring else 'sample'
        )
        if controlled_length is not None:
            self.params['controlled_length_mm'] = round(controlled_length, 1)

        exposure = self.params.get('exposure_scheme')
        if exposure:
            exposure['N'] = n_adj
            exposure['N_segments'] = n_seg_adj

    @staticmethod
    def _normalize_volume_for_quality_norm(raw, quality_norm_code) -> int:
        """НП-105: 100/50/25/10/5; СНиП 3.05.05-84: 100/20/10/2/1 (п. 4.11)."""
        from normative.snip_3_05_05_84 import (
            is_snip_quality_norm,
            normalize_control_volume_pct as snip_normalize_volume,
        )
        if is_snip_quality_norm(quality_norm_code):
            return snip_normalize_volume(raw)
        return normalize_control_volume_pct(raw)

    def _calc_control_volume(self):
        """Объём контроля — значение, выбранное пользователем (п. 3.2)."""
        qn = self.params.get('quality_norm_code') or NP105_CODE
        self.params['control_volume_pct'] = self._normalize_volume_for_quality_norm(
            self.params.get('control_volume_pct', 100),
            qn,
        )


# ---------------------------------------------------------------
# Заполнение шаблона
# ---------------------------------------------------------------

_WELD_TYPE_NAMES = {
    'butt': 'Стыковое', 'corner': 'Угловое (У)',
    'tee': 'Тавровое (Т)', 'lap': 'Нахлёсточное (Н)',
}

_JOINT_MOBILITY_LABELS = {
    'rotating': 'поворотное',
    'non_rotating': 'неповоротное',
}

_SURFACE_QUALITY_BASE = (
    'подлежащие контролю сварные соединения должны быть очищены от окалины, '
    'шлака, брызг металла и других загрязнений. При этом также должны быть '
    'устранены все обнаруженные при внешнем осмотре наружные дефекты, а также '
    'неровности, изображения которых на снимке могут помешать выявлению и '
    'расшифровке изображений внутренних несплошностей и включений в сварном '
    'соединении.'
)


def _fmt_mm(value, decimals: int = 1) -> str:
    """Форматирует размер в мм с запятой для техкарты."""
    if value is None or value == '' or value == '—':
        return '—'
    try:
        return f'{float(value):.{decimals}f}'.replace('.', ',')
    except (TypeError, ValueError):
        return str(value)


def _format_signature_block(
    position: str,
    name: str,
    certificate: str,
    date_str: str,
    default_position: str = '',
) -> str:
    """Блок подписи для нижнего колонтитула."""
    lines = [position or default_position]
    if name:
        lines.append(name)
    if certificate:
        lines.append(f'Удостоверение № {certificate}')
    lines.append(f'«{date_str}» ___________')
    return '\n'.join(filter(None, lines))

_OBJECT_TYPE_NAMES = {
    'pipe': 'Трубопровод (кольцевой сварной шов)',
    'flat': 'Плоская деталь / пластина',
    'vessel': 'Сосуд давления / обечайка',
}


def _build_value_map(params: dict) -> dict:
    """
    Строит словарь: фрагмент метки → точное значение для заполнения ячейки шаблона.

    Использует рассчитанные значения из normative.calculations:
    f (расстояние), N (число экспозиций), L (длина участка), C (коэффициент).
    """
    src = params.get('selected_source') or {}
    scheme_result = params.get('exposure_scheme') or {}
    scheme_info = params.get('scheme_info') or {}
    iqi = params.get('recommended_iqi') or {}
    placement = params.get('iqi_placement') or {}
    screens = params.get('screens') or {}
    film_info = params.get('film_class_info') or {}
    pers = params.get('personnel_requirements') or {}
    fp = params.get('film_processing') or {}
    dev_opts = (fp.get('developer') or {}).get('options') or [{}]
    fix_opts = (fp.get('fixer') or {}).get('options') or [{}]
    d_opt = dev_opts[0]
    f_opt = fix_opts[0]

    S = params.get('wall_thickness', 0)
    D = params.get('outer_diameter', 0)
    d_inner = params.get('d_inner_mm', 0)

    # ---- Рассчитанные параметры просвечивания ----
    f_val = params.get('f_calculated_mm', '')
    N_val = params.get('N_calculated', '')
    L_val = params.get('L_calculated_mm', '')
    C_val = params.get('C_coeff', '')
    sfd_used = params.get('sfd_used_mm', params.get('sfd_mm', ''))
    is_empirical = params.get('is_empirical', False)
    empirical_reason = params.get('empirical_reason', '')

    # --- Формирование полей 6.5, 6.6, 6.7, 6.8 ---
    method_cite = params.get('method_doc_cite') or DOCUMENT_CODE
    if '7512' in method_cite:
        EMPIRICAL_TEXT = (
            'Определяется опытным путём в соответствии с требованиями '
            'ГОСТ 7512-82, приложение 4'
        )
        empirical_short = 'Определяется опытным путём (ГОСТ 7512-82, прил. 4)'
    else:
        EMPIRICAL_TEXT = (
            'Определяется опытным путём в соответствии с требованиями '
            'ГОСТ Р 50.05.07-2018, п. Г.5'
        )
        empirical_short = 'Определяется опытным путём (ГОСТ Р 50.05.07-2018, п. Г.5)'

    if is_empirical:
        # Схемы 5б/3б (l < d_вн) и панорамные: f и N — опытным путём
        f_field = EMPIRICAL_TEXT
        N_str = empirical_short
        l_field = empirical_short
        segments_str = N_str
        if f_val is not None and f_val != '':
            f_field += f'\n(справочно: f_расч ≥ {f_val} мм)'
        if N_val:
            N_str += f'\n(справочно: N_расч ≥ {N_val})'
    else:
        # Обычные схемы: показываем рассчитанные значения
        if f_val is not None and f_val != '':
            f_field = f'f = {f_val} мм'
        else:
            f_field = f'{sfd_used} мм'

        N_str = str(N_val) if N_val else '—'
        N_segments_val = params.get('N_segments')
        segments_str = str(N_segments_val) if N_segments_val else N_str

        scheme_type = params.get('scheme_type', '')
        L_formula = (scheme_result or {}).get('L_formula', '')

        if L_val:
            if L_formula:
                l_field = L_formula
            else:
                from normative.calculations import SCHEME_WALL_COUNT
                d_for_l = (
                    params.get('d_outer_effective_mm')
                    if SCHEME_WALL_COUNT.get(scheme_type, 1) == 2
                    else D
                ) or D
                l_field = (
                    f'{L_val} мм = π × {d_for_l} / {N_val}'
                    if d_for_l and N_val else f'{L_val} мм'
                )
        else:
            l_field = '350 × (длина шва / N) мм'

    # Поле 6.9: схема просвечивания
    scheme_code = params.get('scheme_type', '')
    scheme_name = scheme_info.get('name') or scheme_code
    scheme_desc = scheme_info.get('description', '')
    scheme_notes = params.get('scheme_notes', '')
    if is_empirical and empirical_reason:
        scheme_notes = empirical_reason
    scheme_formula = params.get('scheme_formula', '') if not is_empirical else ''
    scheme_field = '\n'.join(filter(None, [scheme_name, scheme_desc, scheme_formula, scheme_notes]))

    # Угол просвечивания (для поля 6.4)
    angle = '0° (перпендикуляр к поверхности контроля)'
    if params.get('scheme_type') in ('5a', '5b', '5v', '5g', '5d', '5zh', '5z'):
        angle = 'По схеме просвечивания (90° к оси шва)'

    # Коэффициент C для поля 5.2
    focal_field = _fmt_mm(params.get('source_focal_spot_mm', ''))
    if C_val:
        focal_field += f' (C = {C_val:.2f})'

    object_type = params.get('object_type', 'pipe')
    joint_code = params.get('joint_designation', '')
    joint_std = params.get('joint_standard') or 'ГОСТ Р 59023.2-2020'
    from normative.gost_16037_80 import is_gost_16037_joint, get_joint_info as get_info_16037
    from normative.gost_59023_2 import get_joint_info
    if joint_code and (is_gost_16037_joint(joint_code) or '16037' in joint_std):
        joint_info = get_info_16037(joint_code)
    else:
        joint_info = get_joint_info(joint_code) if joint_code else {}
    joint_type = joint_info.get('joint_type', 'butt')
    mobility = _JOINT_MOBILITY_LABELS.get(
        params.get('joint_mobility', 'non_rotating'), 'неповоротное',
    )
    weld_cat = params.get('weld_category', '')

    if object_type == 'pipe':
        outer_d_display = _fmt_mm(D) if D else '—'
        flat_length_display = '—'
    elif object_type == 'flat':
        outer_d_display = '—'
        flat_length_display = _fmt_mm(params.get('flat_length_mm')) if params.get('flat_length_mm') else '—'
    else:
        outer_d_display = _fmt_mm(D) if D else '—'
        flat_length_display = '—'

    backing_label = 'Есть' if params.get('has_backing') else 'Нет'
    backing_thickness_display = (
        f'{_fmt_mm(params.get("backing_thickness_mm"))} мм (учитывать в K, f)'
        if params.get('has_backing') and params.get('backing_thickness_mm')
        else '—'
    )

    _weld_proc = params.get('welding_process', '')
    _weld_proc_names = {
        '10': 'АДФ под флюсом', '11': 'АДФ с подваркой корня', '20': 'ЭШС',
        '30': 'РДС', '31': 'РДС с подваркой корня', '32': 'РДС на подкладке',
        '40': 'Комбинированная (корень АДС)', '42': 'Комбинированная на подкладке',
        '51': 'АДС без присадки', '52': 'АДС с присадкой', '53': 'АДС плавящимся',
        '60': 'ЭЛС',
        'ЗП': 'в защитном газе плавящимся электродом',
        'ЗН': 'в защитном газе неплавящимся электродом',
        'Р': 'ручная дуговая',
        'Ф': 'под флюсом',
        'Г': 'газовая',
    }
    _weld_proc_label = _weld_proc_names.get(_weld_proc, '')

    return {
        # ---- Раздел 1: Объект контроля ----
        '1.1': params.get('organization', ''),
        '1.2': params.get('object_name', ''),
        '1.3': params.get('drawing_number', ''),
        '1.4': params.get('weld_number', ''),
        '1.5': params.get('drawing_number', ''),
        '1.6': f'{_WELD_TYPE_NAMES.get(joint_type, "Стыковое")} ({mobility})',
        '1.7': (
            (
                f'{joint_code}, по {joint_std}'
                + (
                    f'; {params.get("wall_summary")}'
                    if params.get('wall_summary') else ''
                )
            )
            if joint_code else ''
        ),
        '1.8': (
            f'{_weld_proc} — {_weld_proc_label}'.rstrip(' —')
            + f', по {joint_std}'
            if _weld_proc else ''
        ),
        '1.9': params.get('material_display', params.get('material', '')),
        '1.10': params.get('weld_material', ''),

        # ---- Раздел 2: Документация ----
        '2.1': method_cite,
        '2.2': _format_quality_norm_cite(params),

        # ---- Раздел 3: Требования ----
        '3.1': _format_weld_category_field(params, method_cite),
        '3.2': f'{params.get("control_volume_pct", 100)} %',

        # ---- Раздел 4: Тип и размеры ----
        '4.1': _OBJECT_TYPE_NAMES.get(object_type, ''),
        '4.2.1': outer_d_display,
        '4.2.1 S': (
            params.get('wall_summary')
            or _fmt_mm(S)
        ),
        '4.2.2 длин': flat_length_display,
        '4.2.2': params.get('e_display', ''),
        '4.2.2 e1': params.get('e1_display', ''),
        '4.2.3': params.get('g_display', params.get('reinforcement_status', 'не снят')),
        '4.2.3_label': params.get(
            'g_label', '4.2.3. Высота валика усиления (g)',
        ),
        '4.2.4': f'{_fmt_mm(params.get("haz_width_mm", 5.0))} мм (с каждой стороны от краёв шва)',
        '4.2.5': _fmt_mm(params.get('zone_width_mm', '')),
        '4.2.6': backing_label,
        '4.2.7': backing_thickness_display,

        # ---- Раздел 5: Средства контроля (нумерация по бланку TC_) ----
        '5.1': (
            params.get('source_name_override')
            or src.get('name', '')
        ),
        '5.2': focal_field,
        '5.3': (
            f"Проволочный / №{params.get('iqi_marking', '')} по ГОСТ 7512 "
            f"({placement.get('side_label', 'со стороны источника')})"
        ),
        '5.4': _format_film_with_size(params, film_info),
        '5.6': 'свинцовые буквы и цифры по ГОСТ',
        '5.7': 'лупа 10× измерительная, линейка металлическая',
        '5.8': (
            'негатоскоп с яркостью ≥ 10 000 кд/м² (ГОСТ Р 8.763); '
            'денситометр фотометрический'
        ),
        '5.9': 'маркер по металлу (несмываемый)',
        '5.10': _format_vpk_vgk_sample(params),
        '5.11': (
            f'Проявитель: {d_opt.get("name","стандартный")}, '
            f't = {d_opt.get("temp_c","20±1")}°C, '
            f'τ = {d_opt.get("time_min","5–8")} мин; '
            f'Закрепитель: {f_opt.get("name","")}, '
            f'τ = {f_opt.get("time_min","10–15")} мин'
        ),

        # ---- Раздел 6: Параметры и схема контроля (РАССЧИТАННЫЕ) ----
        '6.1': _format_tube_voltage_or_energy(params, src),
        '6.2': _fmt_mm(params.get('s_k_mm', '')),
        '6.3': (
            f'K ≤ {params.get("sensitivity_k_display_mm", params.get("required_sensitivity_mm", "—")):.3f} мм'
            if params.get('sensitivity_k_display_mm') is not None
            else params.get('sensitivity_desc', '')
        ),
        '6.4': angle,
        '6.5': f_field,    # ← РАССЧИТАННОЕ расстояние f
        '6.6': N_str,      # ← РАССЧИТАННОЕ число экспозиций N
        '6.7': segments_str,  # ← Число контролируемых участков
        '6.8': l_field,    # ← РАССЧИТАННАЯ длина участка L
        '6.9': scheme_field,   # ← Схема просвечивания с описанием

        # ---- Разделы 7–8: Подготовка и условия ----
        '7.1': (
            f'Dн = {D} мм, Dвн = {d_inner:.1f} мм, S = {S} мм'
            if D else f'S = {S} мм'
        ),
        '7.2': _build_surface_quality_text(params),
        '8.1': (
            params.get('control_location')
            or 'участок радиографического контроля'
        ),
        '8.2': _build_radiation_safety_text(params),
        '8.3': _format_personnel_text(pers),
        '8.4': params.get('temperature_range') or '+5 ÷ +40',
    }


def _format_film_with_size(params: dict, film_info: dict) -> str:
    film = params.get('film_name') or film_info.get('examples', '') or ''
    length = params.get('film_length_mm')
    width = params.get('film_width_mm')
    if length and width:
        size = f'{float(length):.0f} × {float(width):.0f} мм'
        return f'{film}; формат {size}' if film else size
    return film


def _format_vpk_vgk_sample(params: dict) -> str:
    """П. 5.10 — образец-имитатор ВПК/ВГК (при D > 30 мм)."""
    d = float(params.get('outer_diameter') or 0)
    if d > 30:
        return (
            '№2 по ГОСТ Р 50.05.07-2018 (автоматизированный подбор по прил. В); '
            'применяется при Dн > 30 мм'
        )
    return 'Не требуется (Dн ≤ 30 мм)'


def _format_tube_voltage_or_energy(params: dict, src: dict) -> str:
    """П. 6.1: кВ для рентгена; прочерк для ИИИ."""
    src_type = (src or {}).get('type', '')
    kv = params.get('tube_voltage_kv')
    if src_type == 'xray':
        if kv not in (None, ''):
            return f'{float(kv):.0f} кВ (не более)'
        return (src or {}).get('energy_display', '') or '—'
    if src_type == 'isotope':
        return '—'
    return (src or {}).get('energy_display', '') or '—'


def _format_personnel_text(pers: dict) -> str:
    if not pers:
        return ''
    parts = [pers.get('level', '')]
    if pers.get('standard'):
        parts.append(pers['standard'])
    if pers.get('method'):
        parts.append(f"метод {pers['method']}")
    if pers.get('additional'):
        parts.append(pers['additional'])
    return ', '.join(p for p in parts if p)


def _build_radiation_safety_text(params: dict) -> str:
    """П. 8.2 — краткий блок РБ (как у промышленных техкарт)."""
    items = params.get('safety_requirements') or SAFETY_REQUIREMENTS
    core = [
        'Перед проведением контроля оградить радиационно-опасную зону '
        'сигнальной оградительной лентой и предупреждающими знаками.',
        'Контроль мощности дозы — дозиметром; работы выполнять в соответствии с '
        'ОСПОРБ-99/2010, НРБ-99/2009, СанПиН 2.6.1.3164-14.',
    ]
    extra = []
    for item in items:
        low = item.lower()
        if 'оспорб' in low or 'нрб' in low or 'дозиметр' in low or 'зон' in low:
            continue
        extra.append(item)
        if len(extra) >= 2:
            break
    return ' '.join(core + extra)


def _build_surface_quality_text(params: dict) -> str:
    """П. 7.2 — требования к качеству поверхности перед контролем."""
    text = _SURFACE_QUALITY_BASE
    if params.get('material_type') == 'titanium':
        text += (
            '\n'
            + build_titanium_edge_cleaning_requirement(
                params.get('welding_process', ''),
            )
        )
    return text


def _set_paragraph_text(
    para,
    text: str,
    *,
    font_size: int | None = None,
    bold: bool | None = None,
):
    """Заменяет текст параграфа, сохраняя оформление шаблона."""
    ref = _reference_run_in_paragraph(para)
    if para.runs:
        run = para.runs[0]
        run.text = str(text) if text is not None else ''
        for extra in para.runs[1:]:
            extra.text = ''
    else:
        run = para.add_run(str(text) if text is not None else '')
    if bold is not None:
        run.bold = bold
    elif ref is not None:
        run.bold = ref.bold
    if font_size is not None:
        run.font.size = Pt(font_size)
    elif ref is not None and ref.font.size:
        run.font.size = ref.font.size
    else:
        run.font.size = Pt(12)
    if ref is not None:
        _copy_run_font(ref, run)
        if bold is not None:
            run.bold = bold
        if font_size is not None:
            run.font.size = Pt(font_size)


def _paragraph_has_drawing(para) -> bool:
    wp_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    el = para._element
    return bool(el.findall(f'.//{wp_ns}inline') or el.findall(f'.//{wp_ns}anchor'))


def _clear_paragraph_content(para):
    """Удаляет все run-элементы параграфа (в т.ч. встроенные рисунки)."""
    p_el = para._element
    for child in list(p_el):
        if child.tag.endswith('}r'):
            p_el.remove(child)


def _resolve_static_image_path(static_root: str, image_rel: str) -> str | None:
    """Ищет файл изображения относительно каталога static/."""
    if not image_rel or not static_root:
        return None
    rel = image_rel.replace('\\', '/').lstrip('/')
    if rel.startswith('img/'):
        candidates = [os.path.join(static_root, rel)]
    else:
        candidates = [
            os.path.join(static_root, 'img', 'welds', rel),
            os.path.join(static_root, rel),
            os.path.join(static_root, 'img', rel),
        ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def _picture_source_for_docx(image_path: str):
    """Возвращает путь или BytesIO для add_picture (в т.ч. SVG → PNG)."""
    if image_path.lower().endswith('.svg'):
        try:
            import cairosvg
            return io.BytesIO(cairosvg.svg2png(url=image_path))
        except Exception:
            logger.warning('Не удалось конвертировать SVG для DOCX: %s', image_path)
            return None
    return image_path


def _insert_picture_in_paragraph(para, image_path: str, width_mm: float = 40):
    """Вставляет изображение в параграф DOCX."""
    source = _picture_source_for_docx(image_path)
    if source is None:
        return False
    _clear_paragraph_content(para)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(source, width=Mm(width_mm))
    return True


def _find_paragraph_index(doc: Document, fragment: str) -> int:
    needle = fragment.lower()
    for idx, para in enumerate(doc.paragraphs):
        if needle in para.text.lower():
            return idx
    return -1


def _is_empty_body_paragraph(el) -> bool:
    """True для пустого абзаца без рисунков и без разрыва страницы."""
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    if el.tag.split('}')[-1] != 'p':
        return False
    if _paragraph_has_page_break(el):
        return False
    wp_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    if el.findall(f'.//{wp_ns}inline') or el.findall(f'.//{wp_ns}anchor'):
        return False
    text = ''.join(t.text or '' for t in el.findall(f'.//{ns}t')).strip()
    return not text


def _paragraph_has_page_break(element) -> bool:
    """True, если XML-элемент абзаца содержит разрыв страницы."""
    if element is None:
        return False
    tag = element.tag.split('}')[-1]
    if tag != 'p':
        return False
    return 'w:br' in element.xml and 'page' in element.xml


def _insert_page_break_before_element(element):
    """Вставляет разрыв страницы непосредственно перед XML-элементом тела документа."""
    parent = element.getparent()
    if parent is None:
        return
    p_el = OxmlElement('w:p')
    r_el = OxmlElement('w:r')
    br_el = OxmlElement('w:br')
    br_el.set(qn('w:type'), 'page')
    r_el.append(br_el)
    p_el.append(r_el)
    parent.insert(list(parent).index(element), p_el)


def _insert_empty_body_paragraph_before(element):
    """Вставляет пустой абзац-зазор (1 строка) перед XML-элементом."""
    parent = element.getparent()
    if parent is None:
        return None
    p_el = OxmlElement('w:p')
    p_pr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    # Одна строка ~12 pt (240 twips)
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    p_pr.append(spacing)
    p_el.append(p_pr)
    r_el = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12 pt
    r_pr.append(sz)
    r_el.append(r_pr)
    t_el = OxmlElement('w:t')
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_el.text = ' '
    r_el.append(t_el)
    p_el.append(r_el)
    parent.insert(list(parent).index(element), p_el)
    return p_el


def _count_empty_paragraphs_before(target_el) -> int:
    """Число подряд идущих пустых абзацев непосредственно перед элементом."""
    count = 0
    prev = target_el.getprevious()
    while prev is not None and _is_empty_body_paragraph(prev):
        count += 1
        prev = prev.getprevious()
    return count


def _ensure_empty_paragraphs_before_element(target_el, keep: int = 1):
    """Оставляет ровно keep пустых абзацев перед элементом (добавляет/удаляет)."""
    parent = target_el.getparent()
    if parent is None:
        return
    while _count_empty_paragraphs_before(target_el) > keep:
        prev = target_el.getprevious()
        if prev is None or not _is_empty_body_paragraph(prev):
            break
        parent.remove(prev)
    while _count_empty_paragraphs_before(target_el) < keep:
        _insert_empty_body_paragraph_before(target_el)


def _previous_skipping_empty(element):
    """Предыдущий элемент тела, пропуская пустые абзацы (без page break)."""
    prev = element.getprevious()
    while prev is not None and _is_empty_body_paragraph(prev):
        prev = prev.getprevious()
    return prev


def _remove_duplicate_page_breaks_before(element):
    """Оставляет не более одного разрыва страницы перед элементом; убирает пустые между ними."""
    parent = element.getparent()
    if parent is None:
        return
    to_remove = []
    seen_breaks = 0
    probe = element.getprevious()
    while probe is not None:
        if _paragraph_has_page_break(probe):
            seen_breaks += 1
            if seen_breaks > 1:
                to_remove.append(probe)
            probe = probe.getprevious()
            continue
        if _is_empty_body_paragraph(probe):
            to_remove.append(probe)
            probe = probe.getprevious()
            continue
        break
    for el in to_remove:
        p = el.getparent()
        if p is not None:
            p.remove(el)


def _ensure_paragraph_on_new_page(doc: Document, *needles: str, keep: int = 1):
    """
    Размещает целевой абзац в начале новой страницы с keep пустыми строками
    перед ним.

    Зазор от колонтитула на всех страницах (в т.ч. при переносе таблиц)
    задаётся отдельно через _apply_body_clearance_below_header.
    keep=0 — без пустой строки после разрыва (зазор уже в top_margin).
    """
    para = None
    for p in doc.paragraphs:
        lower = p.text.lower()
        if all(n.lower() in lower for n in needles):
            para = p
            break
    if para is None:
        return

    target_el = para._element
    _remove_duplicate_page_breaks_before(target_el)
    _ensure_empty_paragraphs_before_element(target_el, keep=keep)

    insert_before = target_el
    for _ in range(keep):
        prev = insert_before.getprevious()
        if prev is not None and _is_empty_body_paragraph(prev):
            insert_before = prev
        else:
            break

    before = _previous_skipping_empty(insert_before)
    if before is None or not _paragraph_has_page_break(before):
        _insert_page_break_before_element(insert_before)
    _remove_duplicate_page_breaks_before(target_el)
    if keep == 0:
        # Не оставлять пустую строку между PB и заголовком (двойной зазор с margin)
        parent = target_el.getparent()
        if parent is not None:
            while True:
                prev = target_el.getprevious()
                if prev is None or not _is_empty_body_paragraph(prev):
                    break
                parent.remove(prev)


# Высота одной строки тела (~12 pt) и оценка шапки колонтитула
_BODY_LINE_MM = 5.0
_TWIPS_PER_MM = 1440 / 25.4


def _row_height_twips(row) -> int | None:
    """Явная высота строки таблицы в twips, если задана в trPr."""
    tr_pr = row._tr.find(qn('w:trPr'))
    if tr_pr is None:
        return None
    tr_height = tr_pr.find(qn('w:trHeight'))
    if tr_height is None:
        return None
    val = tr_height.get(qn('w:val'))
    return int(val) if val else None


def _estimate_header_table_height_mm(table) -> float:
    """Оценка высоты таблицы верхнего колонтитула в мм (с учётом переносов)."""
    if table is None:
        return 0.0
    total_twips = 0
    for row in table.rows:
        declared = _row_height_twips(row) or 0
        max_lines = 1
        for cell in row.cells:
            text = (cell.text or '').strip()
            if not text:
                continue
            for line in text.splitlines() or ['']:
                # ~42 символа на строку в ячейке шапки
                max_lines = max(max_lines, max(1, (len(line) + 41) // 42))
        # минимум ~240 twips на строку текста
        estimated = max_lines * 240
        h_rule_exact = False
        tr_pr = row._tr.find(qn('w:trPr'))
        if tr_pr is not None:
            tr_height = tr_pr.find(qn('w:trHeight'))
            if tr_height is not None:
                h_rule_exact = tr_height.get(qn('w:hRule')) == 'exact'
        if h_rule_exact and declared:
            total_twips += declared
        else:
            total_twips += max(declared, estimated)
    return total_twips / _TWIPS_PER_MM


def _estimate_header_part_height_mm(header_part) -> float:
    """Высота содержимого верхнего колонтитула (таблица или абзацы)."""
    try:
        tables = header_part.tables
    except Exception:
        return 0.0
    if tables:
        return max(_estimate_header_table_height_mm(t) for t in tables)
    try:
        n = sum(1 for p in header_part.paragraphs if (p.text or '').strip())
    except Exception:
        n = 0
    return max(n, 1) * _BODY_LINE_MM if n else 0.0


# Минимальные отступы колонтитулов от края листа (мм)
_MIN_HEADER_DISTANCE_MM = 12.0
_MIN_FOOTER_DISTANCE_MM = 12.0
# Оценка высоты подвала для bottom_margin (без раздувания от многострочного текста)
_FOOTER_CLEARANCE_CAP_MM = 16.0


def _apply_body_clearance_below_header(doc: Document, gap_lines: int = 1):
    """
    Гарантирует зазор тела от верхнего и нижнего колонтитулов.

    Word рисует продолжение таблицы с top/bottom_margin; пустой абзац туда
    вставить нельзя. Поэтому:
    top_margin = header_distance + высота шапки + gap;
    bottom_margin = footer_distance + высота подвала (с потолком) + gap.
    """
    gap_mm = max(0, gap_lines) * _BODY_LINE_MM
    for section in doc.sections:
        header_dist = (
            section.header_distance.mm
            if section.header_distance is not None
            else _MIN_HEADER_DISTANCE_MM
        )
        footer_dist = (
            section.footer_distance.mm
            if section.footer_distance is not None
            else _MIN_FOOTER_DISTANCE_MM
        )
        if header_dist < _MIN_HEADER_DISTANCE_MM - 0.05:
            section.header_distance = Mm(_MIN_HEADER_DISTANCE_MM)
            header_dist = _MIN_HEADER_DISTANCE_MM
        if footer_dist < _MIN_FOOTER_DISTANCE_MM - 0.05:
            section.footer_distance = Mm(_MIN_FOOTER_DISTANCE_MM)
            footer_dist = _MIN_FOOTER_DISTANCE_MM

        header_heights = [_estimate_header_part_height_mm(section.header)]
        footer_heights = [_estimate_header_part_height_mm(section.footer)]
        try:
            if section.different_first_page_header_footer:
                header_heights.append(
                    _estimate_header_part_height_mm(section.first_page_header)
                )
                footer_heights.append(
                    _estimate_header_part_height_mm(section.first_page_footer)
                )
        except Exception:
            pass
        header_h = max(header_heights) if header_heights else 12.0
        footer_h = min(
            max(footer_heights) if footer_heights else 12.0,
            _FOOTER_CLEARANCE_CAP_MM,
        )
        needed_top = header_dist + header_h + gap_mm
        needed_bottom = footer_dist + footer_h + gap_mm
        current_top = section.top_margin.mm if section.top_margin else 0.0
        current_bottom = section.bottom_margin.mm if section.bottom_margin else 0.0
        if current_top < needed_top - 0.05:
            section.top_margin = Mm(round(needed_top, 1))
        if current_bottom < needed_bottom - 0.05:
            section.bottom_margin = Mm(round(needed_bottom, 1))


def _ensure_section_43_on_page_three(doc: Document):
    """
    П. 4.3 «Эскиз сварного соединения» — всегда в начале новой страницы.
    Зазор 1 строки от колонтитула — через top_margin (_apply_body_clearance).
    """
    _ensure_paragraph_on_new_page(doc, '4.3', 'эскиз', keep=0)


def _place_section_69_after_68(doc: Document):
    """
    П. 6.9 «Схема просвечивания» — сразу после п. 6.8, без разрыва страницы.

    Принудительный page break давал страницу только с 6.8 и пустотой.
    Оставляем ровно одну пустую строку перед заголовком 6.9.
    """
    para = None
    for p in doc.paragraphs:
        lower = p.text.lower()
        if '6.9' in lower and 'схема' in lower:
            para = p
            break
    if para is None:
        return

    target_el = para._element
    parent = target_el.getparent()
    if parent is None:
        return

    # Убрать все разрывы страницы и пустые абзацы перед 6.9
    while True:
        prev = target_el.getprevious()
        if prev is None:
            break
        if _paragraph_has_page_break(prev) or _is_empty_body_paragraph(prev):
            parent.remove(prev)
            continue
        break

    _ensure_empty_paragraphs_before_element(target_el, keep=1)


def _clear_cell_content(cell):
    """Удаляет всё содержимое ячейки таблицы DOCX."""
    tc = cell._tc
    for child in list(tc):
        if child.tag.split('}')[-1] != 'tcPr':
            tc.remove(child)


def _fill_section_102_criteria_table(doc: Document, params: dict):
    """
    Заполняет п. 10.2 техкарты: вводный текст и вложенная таблица
    с нормами из табл. 4.8–4.11 НП-105-18 для выбранных параметров.
    """
    docx_data = params.get('acceptance_criteria_docx')
    if not docx_data:
        from normative.snip_3_05_05_84 import (
            is_snip_quality_norm,
            build_acceptance_criteria_docx_data as snip_acceptance_docx,
        )
        material_type = params.get('material_type', 'steel')
        cat = params.get('weld_category', 'II')
        th = float(params.get('wall_thickness', 0) or 0)
        qn = params.get('quality_norm_code') or NP105_CODE
        if is_snip_quality_norm(qn):
            docx_data = snip_acceptance_docx(material_type, cat, th)
        else:
            docx_data = build_acceptance_criteria_docx_data(material_type, cat, th)

    headers = docx_data.get('headers') or []
    row_values = docx_data.get('row_values') or []
    if not headers:
        return

    # НП-105: одна строка значений; СНиП (баллы): список строк
    if row_values and isinstance(row_values[0], (list, tuple)):
        data_rows = list(row_values)
    else:
        data_rows = [row_values]

    section_table = None
    row_idx = -1
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            label = row.cells[0].text.strip().lower()
            if label.startswith('10.2'):
                section_table = table
                row_idx = ri
                break
        if section_table is not None:
            break
    if section_table is None or row_idx < 0:
        return

    cell = _unique_cells(section_table.rows[row_idx])[0]
    _clear_cell_content(cell)

    intro_para = cell.add_paragraph(docx_data.get('intro', ''))
    intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_para.paragraph_format.space_before = Pt(0)
    intro_para.paragraph_format.space_after = Pt(0)
    if intro_para.runs:
        intro_para.runs[0].font.size = Pt(12)

    nested = cell.add_table(rows=1 + len(data_rows), cols=len(headers))
    nested.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_borders(nested)

    for ci, header in enumerate(headers):
        _set_cell_text(nested.rows[0].cells[ci], header, bold=True, font_size=9)
    for ri, values in enumerate(data_rows, start=1):
        for ci in range(len(headers)):
            val = values[ci] if ci < len(values) else '—'
            _set_cell_text(nested.rows[ri].cells[ci], val, font_size=9)

    _trim_empty_paragraphs_after_last_table(cell)
    _clear_table_row_height(section_table.rows[row_idx])


def _find_section_row(doc: Document, prefix: str):
    """Находит таблицу и индекс строки по префиксу метки (например '10.5')."""
    needle = prefix.lower()
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            label = row.cells[0].text.strip().lower().replace('\xa0', ' ')
            if label.startswith(needle):
                return table, ri
    return None, -1


def _fill_section_105_root_limits(doc: Document, params: dict):
    """
    Заполняет п. 10.5: допустимая выпуклость (табл. 4.5) и вогнутость
    корня (табл. 4.3 / 4.4) по НП-105-18.
    """
    text = params.get('root_acceptance_docx')
    if not text:
        root = params.get('root_acceptance_limits')
        if not root:
            root = lookup_root_acceptance_limits(
                float(params.get('wall_thickness') or 0),
                float(params.get('outer_diameter') or 0),
                joint_mobility=params.get('joint_mobility', 'non_rotating'),
                has_backing=bool(params.get('has_backing')),
            )
        text = build_root_acceptance_docx_text(root)
    if not text:
        return

    table, row_idx = _find_section_row(doc, '10.5')
    if table is None:
        return
    cell = _unique_cells(table.rows[row_idx])[0]
    _clear_cell_content(cell)
    para = cell.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if para.runs:
        para.runs[0].font.size = Pt(12)
    _clear_table_row_height(table.rows[row_idx])


def _update_section_92_optical_density(doc: Document, params: dict):
    """Подставляет в п. 9.2 расчётный диапазон оптической плотности."""
    od_min = params.get('optical_density_min')
    od_max = params.get('optical_density_max')
    if od_min is None or od_max is None:
        return
    table, row_idx = _find_section_row(doc, '9.2')
    if table is None:
        return
    cell = _unique_cells(table.rows[row_idx])[0]
    text = cell.text or ''
    if not text:
        return
    replacement = (
        f'{_fmt_mm(od_min)}–{_fmt_mm(od_max)} е.о.п.'
        if hasattr(od_min, 'real') or isinstance(od_min, (int, float))
        else f'{od_min}–{od_max} е.о.п.'
    )
    # Частые шаблонные диапазоны в бланке
    new_text = text
    for old in ('1,5–3,5 е.о.п.', '1.5–3.5 е.о.п.', '1,5-3,5 е.о.п.',
                '1,5–4,5 е.о.п.', '2,0–4,5 е.о.п.'):
        if old in new_text:
            new_text = new_text.replace(old, replacement)
            break
    else:
        # Если точного совпадения нет — заменить любой «X–Y е.о.п.»
        import re as _re
        new_text, n = _re.subn(
            r'\d+[.,]\d+\s*[–-]\s*\d+[.,]\d+\s*е\.о\.п\.',
            replacement,
            new_text,
            count=1,
        )
        if n == 0:
            return
    _clear_cell_content(cell)
    para = cell.add_paragraph(new_text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if para.runs:
        para.runs[0].font.size = Pt(12)


def _clear_table_row_height(row):
    """Снимает фиксированную высоту строки таблицы (убирает лишний зазор в ячейке)."""
    tr = row._tr
    tr_pr = tr.find(qn('w:trPr'))
    if tr_pr is None:
        return
    row_height = tr_pr.find(qn('w:trHeight'))
    if row_height is not None:
        tr_pr.remove(row_height)


def _trim_empty_paragraphs_after_last_table(cell):
    """Удаляет пустые абзацы после последней вложенной таблицы в ячейке."""
    tc = cell._tc
    children = [child for child in tc if child.tag.split('}')[-1] != 'tcPr']
    last_tbl_idx = None
    for idx, child in enumerate(children):
        if child.tag.split('}')[-1] == 'tbl':
            last_tbl_idx = idx
    if last_tbl_idx is None:
        return
    for child in children[last_tbl_idx + 1:]:
        if child.tag.split('}')[-1] == 'p' and _is_empty_body_paragraph(child):
            tc.remove(child)


def _collapse_excessive_empty_paragraphs(doc: Document, max_consecutive: int = 1):
    """
    Схлопывает серии пустых абзацев в теле документа до max_consecutive.
    Убирает большие отступы от верхнего колонтитула между блоками.
    """
    body = doc.element.body
    to_remove = []
    run_len = 0
    for child in list(body):
        tag = child.tag.split('}')[-1]
        if tag == 'p' and _is_empty_body_paragraph(child):
            run_len += 1
            if run_len > max_consecutive:
                to_remove.append(child)
        else:
            run_len = 0
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _clear_cell_paragraphs(cell):
    """Удаляет все абзацы ячейки таблицы (для полной пересборки содержимого)."""
    tc = cell._tc
    for child in list(tc):
        if child.tag.split('}')[-1] == 'p':
            tc.remove(child)


def _trim_empty_paragraphs_before(doc: Document, *needles: str, keep: int = 1):
    """
    Оставляет не более keep пустых абзацев перед целевым параграфом.
    Используется для п. 4.3 — зазор в одну строку от верхнего колонтитула.
    """
    para = None
    for p in doc.paragraphs:
        lower = p.text.lower()
        if all(n.lower() in lower for n in needles):
            para = p
            break
    if para is None:
        return

    target_el = para._element
    parent = target_el.getparent()
    if parent is None:
        return

    children = list(parent)
    idx = children.index(target_el)
    empty_before = []
    for i in range(idx - 1, -1, -1):
        child = children[i]
        if not _is_empty_body_paragraph(child):
            break
        empty_before.append(child)

    for el in empty_before[keep:]:
        parent.remove(el)


def _prepare_gost16037_sketch_for_docx(image_path: str) -> str:
    """
    Подготавливает эскиз ГОСТ 16037 для п. 4.3 техкарты:
    тот же принцип, что в UI — оба чертежа без текста шапки таблицы.
    Возвращает путь к временному PNG (или исходный при ошибке).
    """
    try:
        from PIL import Image
        from normative.gost_16037_sketches import refine_dual_sketch

        with Image.open(image_path) as im:
            im = refine_dual_sketch(im.convert('RGB'))
            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp_path = tmp.name
            tmp.close()
            im.save(tmp_path, format='PNG', optimize=True)
            return tmp_path
    except Exception:
        logger.exception('Не удалось подготовить эскиз ГОСТ 16037: %s', image_path)
        return image_path


def _insert_joint_sketch_into_docx(doc: Document, params: dict, static_root: str):
    """
    Вставляет эскиз сварного соединения в п. 4.3 шаблона техкарты.

    Используется то же изображение, что на шаге 3 (get_joint_image_path).
    Для ГОСТ 16037 — крупнее (84 мм ≈ ×2 к прежним 42 мм) и с точным кропом.
    """
    from normative.gost_59023_2 import get_joint_image_path
    from normative.gost_16037_80 import (
        is_gost_16037_joint,
        get_joint_image_path as get_image_16037,
        get_joint_image_abs_path,
    )

    idx = _find_paragraph_index(doc, '4.3')
    if idx < 0 or 'эскиз' not in doc.paragraphs[idx].text.lower():
        return

    joint_code = params.get('joint_designation', '')
    use_16037 = bool(
        params.get('joint_image_static_root') or is_gost_16037_joint(joint_code)
    )
    if use_16037:
        image_rel = params.get('joint_image') or get_image_16037(joint_code)
        abs16037 = get_joint_image_abs_path(joint_code)
        image_path = str(abs16037) if abs16037 else _resolve_static_image_path(
            static_root, image_rel,
        )
    else:
        image_rel = params.get('joint_image') or get_joint_image_path(joint_code)
        image_path = _resolve_static_image_path(static_root, image_rel)
    if not image_path:
        logger.warning('Эскиз шва не найден: %s', image_rel)
        return

    # ГОСТ 16037: эскизы шире (два чертежа в ряд) — 84 мм; 59023 — 42 мм
    width_mm = 84 if use_16037 else 42
    insert_path = image_path
    tmp_path = None
    if use_16037:
        prepared = _prepare_gost16037_sketch_for_docx(image_path)
        if prepared != image_path:
            tmp_path = prepared
            insert_path = prepared

    try:
        if idx + 1 < len(doc.paragraphs):
            _insert_picture_in_paragraph(
                doc.paragraphs[idx + 1], insert_path, width_mm=width_mm,
            )

        if idx + 2 < len(doc.paragraphs):
            caption = doc.paragraphs[idx + 2]
            if joint_code:
                _set_paragraph_text(
                    caption,
                    f'Сварное соединение {joint_code} по '
                    f'{params.get("joint_standard") or "ГОСТ Р 59023.2-2020"}'
                    + (
                        f' ({params.get("wall_summary")})'
                        if params.get('wall_summary') else ''
                    ),
                    font_size=9,
                )
    finally:
        if tmp_path and os.path.isfile(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _insert_scheme_section_into_docx(doc: Document, params: dict, static_root: str):
    """
    Встраивает PNG-схему просвечивания напрямую в DOCX (п. 6.9).

    Изображение записывается в word/media/ через run.add_picture().
    """
    from techcards.scheme_display import (
        get_scheme_docx_caption,
        get_scheme_docx_image_rel,
        get_scheme_docx_image_width,
    )

    scheme_code = params.get('scheme_type', '')
    scheme_info = params.get('scheme_info') or {}
    image_rel = get_scheme_docx_image_rel(scheme_code, scheme_info)
    if not image_rel:
        return

    image_path = _resolve_static_image_path(static_root, image_rel)
    if not image_path:
        logger.warning('Схема просвечивания не найдена: %s', image_rel)
        return

    idx = _find_paragraph_index(doc, '6.9')
    if idx < 0 or 'схема' not in doc.paragraphs[idx].text.lower():
        return

    width_mm = get_scheme_docx_image_width(scheme_code, doc)
    caption_text = get_scheme_docx_caption(
        scheme_code, scheme_info, doc_code=params.get('doc_code'),
    )

    image_idx = None
    caption_idx = None
    for j in range(idx + 1, min(idx + 8, len(doc.paragraphs))):
        if _paragraph_has_drawing(doc.paragraphs[j]):
            image_idx = j
            break
    if image_idx is None:
        for j in range(idx + 1, min(idx + 8, len(doc.paragraphs))):
            if not doc.paragraphs[j].text.strip():
                image_idx = j
                break
    for j in range(idx + 1, min(idx + 8, len(doc.paragraphs))):
        if j == image_idx:
            continue
        para = doc.paragraphs[j]
        if caption_idx is None and para.text.strip() and not _paragraph_has_drawing(para):
            caption_idx = j

    if image_idx is None and idx + 1 < len(doc.paragraphs):
        image_idx = idx + 1

    if image_idx is None:
        return

    image_para = doc.paragraphs[image_idx]

    if not _insert_picture_in_paragraph(image_para, image_path, width_mm=width_mm):
        logger.warning('Не удалось встроить схему в DOCX: %s', image_path)
        return

    for j in range(idx + 1, min(idx + 8, len(doc.paragraphs))):
        if j != image_idx and _paragraph_has_drawing(doc.paragraphs[j]):
            _clear_paragraph_content(doc.paragraphs[j])

    if caption_idx is not None and caption_text:
        _set_paragraph_text(doc.paragraphs[caption_idx], caption_text, font_size=9)
    elif caption_text and idx + 3 < len(doc.paragraphs):
        _set_paragraph_text(doc.paragraphs[idx + 3], caption_text, font_size=9)


def _fill_dimension_rows(doc: Document, value_map: dict):
    """Заполняет многоячеечные строки раздела 4.2 шаблона техкарты."""
    for table in doc.tables:
        for row in table.rows:
            ucells = _unique_cells(row)
            if not ucells:
                continue
            label = ucells[0].text.strip().lower()

            if '4.2.1' in label and 'внешний' in label and len(ucells) >= 4:
                _set_cell_text(ucells[1], value_map.get('4.2.1', '—'))
                _set_cell_text(ucells[3], value_map.get('4.2.1 S', ''))
                continue

            if ('длинна' in label or 'длина' in label) and '4.2.2' in label:
                if len(ucells) >= 2:
                    _set_cell_text(ucells[1], value_map.get('4.2.2 длин', '—'))
                continue

            if 'наружной поверхности' in label and len(ucells) >= 4:
                _set_cell_text(ucells[1], value_map.get('4.2.2', ''))
                _set_cell_text(ucells[3], value_map.get('4.2.2 e1', ''))
                continue

            if '4.2.3' in label and 'высота' in label:
                label_txt = value_map.get('4.2.3_label')
                if label_txt:
                    _set_cell_text(ucells[0], label_txt)
                if len(ucells) >= 2:
                    _set_cell_text(ucells[-1], value_map.get('4.2.3', ''))
                continue

            if '4.2.4' in label and 'околошовной' in label and len(ucells) >= 4:
                _set_cell_text(ucells[1], value_map.get('4.2.4', ''))
                _set_cell_text(ucells[3], value_map.get('4.2.5', ''))
                continue

            if '4.2.6' in label and len(ucells) >= 4:
                _set_cell_text(ucells[1], value_map.get('4.2.6', ''))
                _set_cell_text(ucells[3], value_map.get('4.2.7', ''))


def generate_from_template(params: dict, template_path: str, output_path: str,
                           static_root: str = '') -> str:
    """
    Заполняет шаблон DOCX технологической карты расчётными данными.

    Стратегия:
    1. Открывает оригинальный шаблон как копию.
    2. Для каждой таблицы, для каждой строки находит ячейку-метку
       и записывает значение в ячейку-значение (последнюю уникальную ячейку строки).
    3. Заменяет плейсхолдер организации в шапке.
    4. Сохраняет результат.

    :param params: рассчитанные параметры техкарты
    :param template_path: путь к файлу шаблона
    :param output_path: путь для сохранения результата
    :return: путь к созданному файлу
    """
    doc = Document(template_path)
    _remove_docx_comments(doc)
    value_map = _build_value_map(params)

    # --- Замена в таблицах ---
    for table in doc.tables:
        for row in table.rows:
            ucells = _unique_cells(row)
            if not ucells:
                continue

            label_text = ucells[0].text.strip()
            label_lower = label_text.lower()
            # Раздел 4.2 заполняется отдельно (_fill_dimension_rows) по ГОСТ Р 59023.2-2020
            if '4.2.' in label_lower or label_lower.startswith('4.2 '):
                continue

            value_cell = ucells[-1] if len(ucells) >= 2 else None

            matched_value = _match_value_for_label(label_text, value_map)

            if matched_value is not None and value_cell is not None:
                # Не перезаписываем, если ячейка та же что и метка
                if value_cell._tc is not ucells[0]._tc:
                    _set_cell_text(value_cell, matched_value)

    _fill_dimension_rows(doc, value_map)

    _fill_section_102_criteria_table(doc, params)
    _fill_section_105_root_limits(doc, params)
    _update_section_92_optical_density(doc, params)

    # Титульный лист и колонтитулы — по структуре шаблона normative_docs
    _fill_body_title_page(doc, params)
    _fill_template_headers_footers(doc, params)

    # Служебные заголовки-образцы (без удаления пустых абзацев-отступов)
    _compact_document(doc)

    # П. 4.3 — новая страница; п. 6.9 — сразу после 6.8 (без page break)
    _ensure_section_43_on_page_three(doc)
    _place_section_69_after_68(doc)

    # Между блоками на одной странице — не более 1 пустой строки подряд
    _collapse_excessive_empty_paragraphs(doc, max_consecutive=1)
    _ensure_section_43_on_page_three(doc)
    _place_section_69_after_68(doc)

    # Зазор тела от колонтитулов — после правок структуры (в т.ч. перенос таблиц)
    _apply_body_clearance_below_header(doc, gap_lines=1)

    # Изображения встраиваются в финальную структуру документа (после compact/page breaks)
    _insert_joint_sketch_into_docx(doc, params, static_root)
    _insert_scheme_section_into_docx(doc, params, static_root)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _insert_scheme_image_into_docx(doc: Document, params: dict, static_root: str):
    """
    Вставляет изображение схемы просвечивания в поле 6.9 техкарты.

    Добавляет картинку и текстовое описание прямо в найденную ячейку
    (строка с меткой '6.9'). Поддерживает как однояченные (слитые),
    так и многоячеечные строки.

    :param doc: объект DOCX документа
    :param params: параметры техкарты (должен содержать scheme_info)
    :param static_root: путь к статическим файлам Django
    """
    from techcards.scheme_display import get_scheme_docx_image_rel

    scheme_code = params.get('scheme_type', '')
    scheme_info = params.get('scheme_info') or {}
    image_rel = get_scheme_docx_image_rel(scheme_code, scheme_info)

    if not image_rel:
        return

    # Схемы хранятся в static/img/ (не img/welds/)
    image_path = os.path.join(static_root, image_rel)
    if not os.path.exists(image_path):
        # Попробуем без поддиректории
        alt_path = os.path.join(static_root, os.path.basename(image_rel))
        if os.path.exists(alt_path):
            image_path = alt_path
        else:
            return   # Файл не найден — пропускаем без ошибки

    scheme_result = params.get('exposure_scheme') or {}
    scheme_name = scheme_info.get('name', '')
    scheme_desc = scheme_info.get('description', '')
    n_exp = scheme_result.get('n_exposures_min', '')
    rad_note = params.get('rad_thickness', {}).get('wall_desc', '')

    # Ищем ячейку с меткой "6.9"
    for table in doc.tables:
        for row in table.rows:
            ucells = _unique_cells(row)
            if not ucells:
                continue
            label_text = ucells[0].text.strip()
            if '6.9' in label_text:
                # Всегда вставляем в последнюю уникальную ячейку
                # (при слитой строке — это та же единственная ячейка)
                target_cell = ucells[-1]
                try:
                    # Добавляем изображение в новом параграфе ячейки
                    img_para = target_cell.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = img_para.add_run()
                    run_img.add_picture(image_path, width=Mm(75))

                    # Добавляем подпись к схеме
                    desc_para = target_cell.add_paragraph()
                    desc_run = desc_para.add_run(
                        f'{scheme_name}. {scheme_desc}'
                    )
                    desc_run.font.size = Pt(8)
                    desc_run.italic = True

                    # Добавляем информацию о радиационной толщине
                    if rad_note:
                        rad_para = target_cell.add_paragraph()
                        s_rad_f = params.get('s_rad_f_mm', '')
                        rad_run = rad_para.add_run(
                            f'{rad_note}. '
                            f'S_K = {params.get("s_k_mm", "")} мм (для K); '
                            f'Sрад(f) = {s_rad_f} мм'
                        )
                        rad_run.font.size = Pt(8)

                    return   # Успешно вставили — выходим
                except Exception as exc:
                    # Логируем ошибку, но не падаем
                    import traceback
                    traceback.print_exc()
                    return


# ---------------------------------------------------------------
# Генерация PDF
# ---------------------------------------------------------------

def _insert_page_number_field(para):
    """
    Вставляет в параграф нумерацию «страница N страниц M».
    Используется для колонтитулов Word (поля PAGE / NUMPAGES).
    """
    def _field_run(instr_text: str):
        run = para.add_run()
        fc_begin = OxmlElement('w:fldChar')
        fc_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = instr_text
        fc_end = OxmlElement('w:fldChar')
        fc_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fc_begin)
        run._r.append(instr)
        run._r.append(fc_end)

    para.add_run('страница ')
    _field_run(' PAGE ')
    para.add_run(' страниц ')
    _field_run(' NUMPAGES ')


def _remove_docx_comments(doc: Document):
    """Удаляет комментарии Word из документа (служебные пометки шаблона)."""
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    comment_tags = (
        f'{{{w_ns}commentRangeStart',
        f'{{{w_ns}commentRangeEnd',
        f'{{{w_ns}commentReference',
    )
    roots = [doc.element.body]
    for section in doc.sections:
        roots.extend([
            section.header._element,
            section.footer._element,
            section.first_page_header._element,
            section.first_page_footer._element,
        ])
    for root in roots:
        for tag in comment_tags:
            for el in list(root.iter(tag)):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)


def _clear_header_footer_part(part):
    """Удаляет все параграфы и таблицы из верхнего/нижнего колонтитула Word."""
    element = part._element
    for child in list(element):
        tag = child.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            element.remove(child)


def _clear_section_headers_footers(section):
    """
    Очищает все колонтитулы секции, включая отдельные для первой страницы.
    Шаблон DOCX использует different_first_page — без очистки first_page_*
    остаются старые «ФГУП МАРКС» / «Иванов» поверх новых колонтитулов.
    """
    for attr in (
        'header', 'footer',
        'first_page_header', 'first_page_footer',
    ):
        part = getattr(section, attr, None)
        if part is not None:
            _clear_header_footer_part(part)


def _replace_in_paragraph_runs(para, mapping: dict[str, str]):
    """Точечная замена текста в run-элементах без сброса вёрстки параграфа."""
    for run in para.runs:
        text = run.text
        for old, new in mapping.items():
            if old and old in text:
                text = text.replace(old, new)
        run.text = text


def _replace_in_paragraph(para, mapping: dict[str, str]):
    """
    Замена по всему тексту параграфа (в т.ч. если фраза разбита на run-элементы).
    Сохраняет оформление первого run.
    """
    if not mapping:
        return
    full = para.text
    new_full = full
    for old, new in mapping.items():
        if old and old in new_full:
            new_full = new_full.replace(old, new)
    if new_full == full:
        return
    ref = _reference_run_in_paragraph(para)
    if para.runs:
        run = para.runs[0]
        run.text = new_full
        for extra in para.runs[1:]:
            extra.text = ''
    else:
        run = para.add_run(new_full)
    if ref is not None:
        _copy_run_font(ref, run)


def _replace_in_cell(cell, mapping: dict[str, str]):
    for para in cell.paragraphs:
        _replace_in_paragraph(para, mapping)


def _replace_in_table(table, mapping: dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell(cell, mapping)


def _replace_in_document_part(part, mapping: dict[str, str]):
    """Заменяет плейсхолдеры в колонтитуле, сохраняя форматирование шаблона."""
    for para in part.paragraphs:
        _replace_in_paragraph_runs(para, mapping)
    for table in part.tables:
        _replace_in_table(table, mapping)


def _replace_certificate_in_cell(cell, certificate: str):
    """Подставляет номер удостоверения, сохраняя строку шаблона с точками."""
    if not certificate:
        return
    for para in cell.paragraphs:
        if 'Удостоверение' not in para.text:
            continue
        new_text = re.sub(
            r'(Удостоверение №)\s*[^\n]*',
            rf'\1 {certificate}',
            para.text,
            count=1,
        )
        if new_text != para.text:
            _replace_in_paragraph(para, {para.text: new_text})


def _fill_header_card_number_cell(cell, card_num: str):
    """Заполняет ячейку «№» в шапке одной строкой без дублирования."""
    display = f'№ {card_num}'.strip() if card_num else '№'
    _clear_cell_paragraphs(cell)
    para = cell.add_paragraph()
    run = para.add_run(display)
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fill_body_title_page(doc: Document, params: dict):
    """
    Титульный лист: строка «№ …» в одном экземпляре (без дублирования run-фрагментов).
    """
    card_num = (params.get('card_number') or '').strip()
    if not card_num:
        return
    display = f'№ {card_num}'
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if not stripped.startswith('№'):
            continue
        ref = _reference_run_in_paragraph(para)
        _clear_paragraph_content(para)
        run = para.add_run(display)
        if ref is not None:
            _copy_run_font(ref, run)
        elif not run.font.size:
            run.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return


def _short_position_for_title_page(position: str) -> str:
    """Краткая должность для нижнего колонтитула титульного листа."""
    lower = (position or '').lower()
    if 'инженер' in lower:
        return 'Ведущий инженер'
    if 'начальник' in lower and 'лаборатор' in lower:
        return 'Начальник лаборатории'
    return position or ''


def _add_table_borders(table):
    """Добавляет границы таблице через XML (работает в headers/footers)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _fill_default_footer_table(table, params: dict):
    """Нижний колонтитул со 2-й и далее страниц (3×2, без дат)."""
    dev_pos = params.get('developed_by_position', '') or 'Ведущий инженер технолог'
    chk_pos = params.get('checked_by_position', '') or 'Начальник лаборатории НК'
    gmo_pos = params.get('gmo_checked_by_position', '')
    dev_name = params.get('developed_by_name', '') or params.get('inspector_name', '')
    chk_name = params.get('checked_by_name', '')
    gmo_name = params.get('gmo_checked_by_name', '')
    dev_cert = params.get('developed_by_certificate', '')
    chk_cert = params.get('checked_by_certificate', '')

    if gmo_pos or gmo_name:
        gmo_line = ' / '.join(p for p in (gmo_pos, gmo_name) if p)
        if chk_pos and gmo_line:
            chk_pos = f'{chk_pos}\nПроверил от ГМО: {gmo_line}'
        elif gmo_line:
            chk_pos = f'Проверил от ГМО: {gmo_line}'

    mapping = {
        'Ведущий инженер технолог': dev_pos,
        'Инженер-технолог': dev_pos,
        'Начальник лаборатории НК': chk_pos,
        'Иванов Н.Н.': dev_name,
        'Коровушкин Андрей Витальевич': dev_name,
        'Сидоров И.Н.': chk_name,
        'Andrey Kovlech': chk_name,
    }
    _replace_in_table(table, mapping)
    if dev_cert and len(table.rows) >= 3:
        _replace_certificate_in_cell(table.rows[2].cells[0], dev_cert)
    if chk_cert and len(table.rows) >= 3 and len(table.rows[2].cells) > 1:
        _replace_certificate_in_cell(table.rows[2].cells[1], chk_cert)


def _fill_first_page_footer_table(table, params: dict):
    """Нижний колонтитул титульного листа — точечная замена в вёрстке шаблона."""
    dev_date = params.get('develop_date', datetime.now().strftime('%d.%m.%Y'))
    check_date = params.get('check_date', dev_date)
    dev_pos = _short_position_for_title_page(
        params.get('developed_by_position', '') or 'Ведущий инженер технолог',
    )
    chk_pos = _short_position_for_title_page(
        params.get('checked_by_position', '') or 'Начальник лаборатории НК',
    )
    dev_name = params.get('developed_by_name', '') or params.get('inspector_name', '')
    chk_name = params.get('checked_by_name', '')
    dev_cert = params.get('developed_by_certificate', '')

    if len(table.rows) < 3:
        return

    _replace_in_cell(table.rows[1].cells[0], {
        'Ведущий инженер': dev_pos,
        'Инженер-технолог': dev_pos,
    })
    _replace_in_cell(table.rows[1].cells[1], {'Начальник лаборатории': chk_pos})
    _replace_in_cell(table.rows[2].cells[0], {
        'Иванов Н.Н.': dev_name,
        'Коровушкин Андрей Витальевич': dev_name,
        '21.06.2026': dev_date,
        '29.06.2026': dev_date,
    })
    _replace_in_cell(table.rows[2].cells[1], {
        'Сидоров И.Н.': chk_name,
        'Andrey Kovlech': chk_name,
        '21.06.2026': check_date,
        '29.06.2026': check_date,
    })
    if dev_cert:
        _replace_certificate_in_cell(table.rows[2].cells[0], dev_cert)


def _fill_template_headers_footers(doc: Document, params: dict):
    """
    Заполняет колонтитулы шаблона техкарты данными пользователя.

    Шаблон normative_docs использует:
    - одинаковую шапку на всех страницах (header + first_page_header);
    - разные нижние колонтитулы: титул (3×3, даты/подписи) и остальные (3×2).
    """
    for section in doc.sections:
        section.different_first_page_header_footer = True
        sect_pr = section._sectPr
        if sect_pr.find(qn('w:titlePg')) is None:
            sect_pr.insert(0, OxmlElement('w:titlePg'))

        for header_part in (section.header, section.first_page_header):
            header_table = _find_header_table(header_part)
            if header_table is not None:
                _fill_template_header_table(header_table, params)

        default_footer = _find_footer_table(section.footer)
        if default_footer is not None:
            _fill_default_footer_table(default_footer, params)

        first_footer = _find_footer_table(section.first_page_footer)
        if first_footer is not None:
            _fill_first_page_footer_table(first_footer, params)


def _find_header_table(part):
    """Находит таблицу шапки в колонтитуле (2×3, «Технологическая карта»)."""
    for table in part.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 3:
            text = ' '.join(cell.text for row in table.rows for cell in row.cells)
            if 'Технологическая карта' in text:
                return table
    return part.tables[0] if part.tables else None


def _find_footer_table(part):
    """Находит таблицу подписей в нижнем колонтитуле."""
    for table in part.tables:
        text = ' '.join(cell.text for row in table.rows for cell in row.cells)
        if 'Разработал' in text and 'Проверил' in text:
            return table
    return part.tables[0] if part.tables else None


def _fill_page_number_cell(cell):
    """Заменяет статичную нумерацию шаблона полями PAGE / NUMPAGES."""
    _clear_cell_paragraphs(cell)
    para = cell.add_paragraph()
    _insert_page_number_field(para)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        if not run.font.size:
            run.font.size = Pt(12)


def _insert_page_break_before_first_table(doc: Document):
    """
    Вставляет разрыв страницы перед первой таблицей тела документа.
    Титульный лист остаётся без таблиц техкарты (как обложка).
    """
    body = doc.element.body
    for child in list(body):
        if child.tag.split('}')[-1] != 'tbl':
            continue
        p_el = OxmlElement('w:p')
        r_el = OxmlElement('w:r')
        br_el = OxmlElement('w:br')
        br_el.set(qn('w:type'), 'page')
        r_el.append(br_el)
        p_el.append(r_el)
        body.insert(list(body).index(child), p_el)
        return


def _fill_template_header_table(table, params: dict):
    """Заполняет шапку, сохраняя табличную вёрстку шаблона."""
    org = params.get('organization', '') or 'Наименование организации'
    dept = params.get('department', '')
    if dept:
        org_display = f'{org}\n{dept}' if org else dept
    else:
        org_display = org
    card_num = params.get('card_number', '')
    r0 = table.rows[0]
    r1 = table.rows[1]
    _replace_in_cell(r0.cells[0], {
        'ФГУП МАРКС': org_display,
        'Наименование организации': org_display,
    })
    if len(r1.cells) > 1:
        _fill_header_card_number_cell(r1.cells[1], card_num)
    if len(r1.cells) > 2:
        _fill_page_number_cell(r1.cells[2])


def _build_headers_footers_scratch(doc: Document, params: dict):
    """
    Создаёт колонтитулы с нуля (только для fallback без шаблона DOCX).
    """
    org = params.get('organization', '')
    card_num = params.get('card_number', '___')
    dev_date = params.get('develop_date', datetime.now().strftime('%d.%m.%Y'))
    check_date = params.get('check_date', dev_date)

    section = doc.sections[0]
    section.header_distance = Mm(8)
    section.footer_distance = Mm(10)
    section.different_first_page_header_footer = False
    _clear_section_headers_footers(section)

    # ---- Верхний колонтитул ----
    header = section.header

    # Таблица 2 строки × 3 столбца
    ht = header.add_table(rows=2, cols=3, width=Mm(185))
    _add_table_borders(ht)
    for i, w in enumerate([Mm(55), Mm(100), Mm(30)]):
        for cell in ht.columns[i].cells:
            cell.width = w

    # Строка 1: организация | заголовок | пусто
    r0 = ht.rows[0]
    _set_cell_text(r0.cells[0], org or 'Наименование организации', font_size=8)
    _set_cell_text(
        r0.cells[1],
        'Технологическая карта радиографического контроля',
        bold=True, font_size=9,
    )
    r0.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Строка 2: номер карты | номер карты | Лист X / Листов Y
    r1 = ht.rows[1]
    _set_cell_text(r1.cells[0], '', font_size=8)
    _set_cell_text(r1.cells[1], f'№ {card_num}', font_size=9)
    r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Поле нумерации страниц в правой ячейке
    page_para = r1.cells[2].paragraphs[0]
    _insert_page_number_field(page_para)

    # ---- Нижний колонтитул ----
    footer = section.footer

    # Таблица 3 строки × 2 столбца (структура шаблона)
    ft = footer.add_table(rows=3, cols=2, width=Mm(185))
    _add_table_borders(ft)
    for cell in ft.columns[0].cells:
        cell.width = Mm(90)
    for cell in ft.columns[1].cells:
        cell.width = Mm(95)

    dev_pos = params.get('developed_by_position', '') or 'Ведущий инженер технолог'
    chk_pos = params.get('checked_by_position', '') or 'Начальник лаборатории НК'
    dev_name = params.get('developed_by_name', '') or params.get('inspector_name', '')
    chk_name = params.get('checked_by_name', '')
    dev_cert = params.get('developed_by_certificate', '')
    chk_cert = params.get('checked_by_certificate', '')

    dev_line3 = dev_name
    if dev_cert:
        dev_line3 += f'\nУдостоверение № {dev_cert}'
    dev_line3 += f'\n«{dev_date}» ___________'
    chk_line3 = chk_name
    if chk_cert:
        chk_line3 += f'\nУдостоверение № {chk_cert}'
    chk_line3 += f'\n«{check_date}» ___________'

    fr0 = ft.rows[0]
    _set_cell_text(fr0.cells[0], 'Разработал', bold=True, font_size=9)
    _set_cell_text(fr0.cells[1], 'Проверил', bold=True, font_size=9)

    fr1 = ft.rows[1]
    _set_cell_text(fr1.cells[0], dev_pos, font_size=9)
    _set_cell_text(fr1.cells[1], chk_pos, font_size=9)

    fr2 = ft.rows[2]
    _set_cell_text(fr2.cells[0], dev_line3.strip(), font_size=9)
    _set_cell_text(fr2.cells[1], chk_line3.strip(), font_size=9)


def _insert_page_break_before_section10(doc: Document):
    """
    Вставляет разрыв страницы перед разделом «10. Оценка качества»,
    чтобы он всегда начинался с нового листа.

    Ищет таблицу, содержащую текст '10.' или '10.Оценка', и вставляет
    перед ней параграф с разрывом страницы.
    """
    body = doc.element.body

    for table in doc.tables:
        # Ищем таблицу с разделом 10
        table_text = ' '.join(
            cell.text for row in table.rows for cell in row.cells
        )
        if '10.' in table_text and (
            'Оценка' in table_text or 'оценк' in table_text.lower()
        ):
            tbl_el = table._tbl
            parent = tbl_el.getparent()
            if parent is None:
                continue

            # Создаём параграф с разрывом страницы
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            p_el = OxmlElement('w:p')
            r_el = OxmlElement('w:r')
            br_el = OxmlElement('w:br')
            br_el.set(qn('w:type'), 'page')
            r_el.append(br_el)
            p_el.append(r_el)

            # Вставляем параграф прямо перед таблицей
            idx = list(parent).index(tbl_el)
            parent.insert(idx, p_el)
            return   # Достаточно одного разрыва


def _compact_document(doc: Document):
    """
    Удаляет только служебный заголовок-образец.
    Пустые абзацы-отступы между блоками и от колонтитулов сохраняются.
    """
    paras_to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ''
        if (
            'Пример технологической карты' in text
            or 'Heading' in style_name
        ):
            paras_to_remove.append(para._element)

    for el in paras_to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def generate_radiographic_pdf(params: dict, output_path: str) -> str:
    """
    Создаёт PDF-версию технологической карты.
    Используется как дополнение к DOCX или если шаблон недоступен.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    register_cyrillic_fonts()

    doc_pdf = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=15*mm, leftMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle('T', parent=styles['Title'], fontSize=13,
                              spaceAfter=4, alignment=TA_CENTER, fontName=FONT_BOLD)
    head_s = ParagraphStyle('H', parent=styles['Normal'], fontSize=10,
                             spaceAfter=3, spaceBefore=5,
                             backColor=colors.Color(0.84, 0.89, 0.94),
                             leftIndent=4, fontName=FONT_BOLD)
    norm_s = ParagraphStyle('N', parent=styles['Normal'], fontSize=9,
                             spaceAfter=2, fontName=FONT_REGULAR)
    label_s = ParagraphStyle('L', parent=styles['Normal'], fontSize=9,
                              textColor=colors.Color(0.2, 0.2, 0.6),
                              fontName=FONT_REGULAR)

    story = []
    card_num = params.get('card_number', '___')
    dev_date = params.get('develop_date', datetime.now().strftime('%d.%m.%Y'))

    story.append(Paragraph('ТЕХНОЛОГИЧЕСКАЯ КАРТА', title_s))
    story.append(Paragraph('РАДИОГРАФИЧЕСКОГО КОНТРОЛЯ', title_s))
    story.append(Paragraph(
        f'№ {card_num}     Дата: {dev_date}     '
        f'Методический документ: {params.get("method_doc_cite", DOCUMENT_CODE)}',
        norm_s,
    ))
    story.append(Spacer(1, 5*mm))

    src = params.get('selected_source') or {}
    scheme = params.get('exposure_scheme') or {}
    iqi = params.get('recommended_iqi') or {}
    screens = params.get('screens') or {}
    film_info = params.get('film_class_info') or {}
    pers = params.get('personnel_requirements') or {}

    def section(title, rows):
        story.append(Paragraph(title, head_s))
        tdata = []
        for lbl, val in rows:
            tdata.append([
                Paragraph(f'<b>{lbl}</b>', label_s),
                Paragraph(str(val or '—'), norm_s),
            ])
        if tdata:
            t = Table(tdata, colWidths=[80*mm, 95*mm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.95, 0.97, 1.0)),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(t)
            story.append(Spacer(1, 3*mm))

    section('1–3. ОБЪЕКТ И ТРЕБОВАНИЯ', [
        ('1.1 Предприятие-изготовитель', params.get('organization')),
        ('1.2 Наименование детали (объекта)', params.get('object_name')),
        ('1.3 № чертежа', params.get('drawing_number')),
        ('1.4 Контролируемый элемент', params.get('weld_number')),
        ('1.6 Тип сварного соединения', _WELD_TYPE_NAMES.get(params.get('weld_type', ''), '')),
        ('1.8 Способ сварки', params.get('welding_process')),
        ('1.9 Основной металл', params.get('material_display') or params.get('material')),
        ('2.1 Методическая документация', params.get('method_doc_cite', DOCUMENT_CODE)),
        ('2.2 Нормативная документация', _format_quality_norm_cite(params)),
        ('3.1 Категория сварного соединения', _format_weld_category_field(params)),
        ('3.2 Объём контроля', str(params.get('control_volume_pct', 100)) + ' %'),
    ])

    section('4. ТИП И РАЗМЕРЫ', [
        ('4.1 Тип контролируемого элемента', _OBJECT_TYPE_NAMES.get(params.get('object_type', ''), '')),
        ('4.2.1 Наружный диаметр, мм', params.get('outer_diameter') or '—'),
        ('Толщина стенки (S), мм', params.get('wall_thickness')),
        ('Требуемая чувствительность (К)', params.get('sensitivity_desc')),
    ])

    section('5. СРЕДСТВА КОНТРОЛЯ', [
        ('5.1 Источник излучения', src.get('name')),
        ('Энергия излучения', src.get('energy_display')),
        ('5.2 Размер фокусного пятна (d), мм', params.get('source_focal_spot_mm')),
        ('Активность / мощность', params.get('source_activity') or 'по паспорту источника'),
        ('5.3 Тип и номер ИКИ', (
            f"Проволочный эталон {params.get('iqi_marking', '')} — "
            f"{params.get('iqi_label', '')} "
            f"({(params.get('iqi_placement') or {}).get('side_label', '')})"
        )),
        ('Диаметр контрольной проволоки, мм', params.get('iqi_wire_diameter_mm')),
        ('5.4 Тип плёнки', params.get('film_name') or film_info.get('examples', '')),
        ('Класс плёнки (ГОСТ ИСО 11699-1)', film_info.get('class')),
        ('Мин. оптическая плотность', params.get('optical_density_min')),
        ('Передний экран, мм', screens.get('front_mm')),
        ('Задний экран, мм', screens.get('back_mm')),
    ])

    section('6. ПАРАМЕТРЫ И СХЕМА КОНТРОЛЯ', [
        ('6.1 Энергия / напряжение', src.get('energy_display')),
        ('6.2 Толщина для расчёта (Sк), мм', _fmt_mm(params.get('s_k_mm', ''))),
        ('6.3 Требуемая чувствительность (К)', params.get('sensitivity_desc')),
        ('6.5 Расстояние источник–детектор (SFD), мм', params.get('sfd_mm')),
        ('Расстояние объект–детектор (OFD), мм', params.get('ofd_mm')),
        ('Геометрическая нерезкость (Ug), мм', params.get('ug_calculation')),
        ('Доп. нерезкость (Ug max), мм', params.get('max_geometric_unsharpness_mm')),
        ('6.6 Число экспозиций, шт.', scheme.get('n_exposures_min')),
        ('6.9 Схема просвечивания', scheme.get('name')),
    ])

    section('7–10. ПОДГОТОВКА, УСЛОВИЯ, ОЦЕНКА', [
        ('8.3 Состав рабочего звена', pers.get('level')),
        ('8.4 Диапазон рабочих температур, °С', '+5 ÷ +40'),
        ('9. Нормативный документ оценки качества', params.get('quality_normative') or _format_quality_norm_cite(params)),
        ('10. Критерии качества', params.get('quality_criteria_summary')),
        ('Специалист НК', params.get('inspector_name') or '___________________'),
    ])

    # Предупреждения не включаются в готовую техкарту
    # (они остаются в логах системы)

    doc_pdf.build(story)
    return output_path


def generate_pdf_for_techcard(
    params: dict,
    pdf_abs: str,
    docx_abs: str | None = None,
    *,
    template_path: str | None = None,
    static_root: str = '',
) -> str:
    """
    Создаёт PDF техкарты: DOCX → HTML → PDF (mammoth + xhtml2pdf).

    При ошибке конвертации использует ReportLab как резервный путь.
    """
    docx_for_pdf = docx_abs
    temp_docx: str | None = None

    try:
        if not docx_for_pdf or not os.path.isfile(docx_for_pdf):
            fd, temp_docx = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            docx_for_pdf = temp_docx
            if template_path and os.path.exists(template_path):
                generate_from_template(
                    params, template_path, docx_for_pdf, static_root=static_root,
                )
            else:
                _generate_docx_fallback(params, docx_for_pdf)

        from common.docx_to_pdf import convert_docx_to_pdf
        convert_docx_to_pdf(docx_for_pdf, pdf_abs)
        logger.info('PDF создан из DOCX: %s', pdf_abs)
        return pdf_abs
    except Exception as exc:
        logger.warning(
            'DOCX→PDF не удался (%s), используем ReportLab fallback', exc,
        )
        return generate_radiographic_pdf(params, pdf_abs)
    finally:
        if temp_docx and os.path.isfile(temp_docx):
            os.remove(temp_docx)


# ---------------------------------------------------------------
# Главная функция генерации
# ---------------------------------------------------------------

def generate_tech_card(input_data: dict, media_root: str,
                       template_path: str = None) -> dict:
    """
    Главная функция генерации технологической карты.

    1. Рассчитывает параметры.
    2. Если шаблон доступен — заполняет его данными (DOCX).
    3. Если шаблона нет — создаёт DOCX программно (fallback).
    4. В обоих случаях создаёт PDF из DOCX (mammoth + xhtml2pdf, fallback ReportLab).

    :param input_data: данные из формы пользователя
    :param media_root: путь к медиа-директории Django
    :param template_path: путь к шаблону DOCX (опционально)
    :return: словарь {'params', 'docx_path', 'pdf_path', 'errors', 'warnings'}
    """
    calc = RadiographicTechCardCalculator(input_data)
    params = calc.calculate()
    params['errors'] = calc.errors
    params['warnings'] = calc.warnings

    uid = uuid.uuid4().hex[:10]
    card_num = (input_data.get('card_number', uid) or uid).replace('/', '-').replace(' ', '_')
    date_dir = datetime.now().strftime('%Y/%m')

    docx_rel = f'techcards/docx/{date_dir}/TC_{card_num}_{uid}.docx'
    pdf_rel = f'techcards/pdf/{date_dir}/TC_{card_num}_{uid}.pdf'
    docx_abs = os.path.join(media_root, docx_rel)
    pdf_abs = os.path.join(media_root, pdf_rel)

    os.makedirs(os.path.dirname(docx_abs), exist_ok=True)
    os.makedirs(os.path.dirname(pdf_abs), exist_ok=True)

    # DOCX: используем шаблон если он есть
    from django.conf import settings as django_settings
    static_root = str(getattr(django_settings, 'STATICFILES_DIRS', [''])[0]) if getattr(django_settings, 'STATICFILES_DIRS', []) else ''

    if template_path and os.path.exists(template_path):
        generate_from_template(params, template_path, docx_abs, static_root=static_root)
    else:
        _generate_docx_fallback(params, docx_abs)

    # PDF из DOCX (mammoth + xhtml2pdf), при ошибке — ReportLab
    generate_pdf_for_techcard(
        params, pdf_abs, docx_abs,
        template_path=template_path, static_root=static_root,
    )

    return {
        'params': params,
        'docx_path': docx_rel,
        'pdf_path': pdf_rel,
        'errors': calc.errors,
        'warnings': calc.warnings,
    }


def get_default_template_path() -> str | None:
    """Путь к шаблону DOCX техкарты (эталон — card_templates/)."""
    from django.conf import settings as django_settings
    base = django_settings.BASE_DIR
    card_dir = os.path.join(base, 'card_templates')

    # Откорректированный образец с отступами и разрывами блоков
    if os.path.isdir(card_dir):
        for name in sorted(os.listdir(card_dir)):
            if name.startswith('TC_') and name.endswith('.docx'):
                path = os.path.join(card_dir, name)
                if os.path.isfile(path):
                    return path

    candidates = [
        os.path.join(
            card_dir,
            'Пример технологической карты радиографического контроля.docx',
        ),
        os.path.join(
            base, 'normative_docs',
            'Пример_технологической_карты_радиографического_контроля  с комментариями.docx',
        ),
    ]
    for template_path in candidates:
        if os.path.exists(template_path):
            return template_path
    return None


def regenerate_techcard_files(
    techcard,
    media_root: str,
    template_path: str | None = None,
    *,
    docx: bool = True,
    pdf: bool = True,
) -> None:
    """
    Восстанавливает файлы техкарты из JSON в БД.

    На Render и других PaaS локальный media/ очищается при деплое — файлы
    нужно пересоздавать по запросу скачивания.
    """
    params = techcard.generated_data
    if not params:
        raise ValueError('Нет сохранённых параметров техкарты для восстановления файла.')

    if template_path is None:
        template_path = get_default_template_path()

    from django.conf import settings as django_settings
    static_root = ''
    if getattr(django_settings, 'STATICFILES_DIRS', []):
        static_root = str(django_settings.STATICFILES_DIRS[0])

    if docx:
        docx_rel = str(techcard.docx_file) if techcard.docx_file else ''
        if not docx_rel:
            uid = uuid.uuid4().hex[:10]
            card_num = (techcard.card_number or str(techcard.pk)).replace('/', '-').replace(' ', '_')
            date_dir = (
                techcard.created_at.strftime('%Y/%m')
                if getattr(techcard, 'created_at', None) else datetime.now().strftime('%Y/%m')
            )
            docx_rel = f'techcards/docx/{date_dir}/TC_{card_num}_{uid}.docx'
            techcard.docx_file = docx_rel
        docx_abs = os.path.join(media_root, docx_rel)
        os.makedirs(os.path.dirname(docx_abs), exist_ok=True)
        if template_path and os.path.exists(template_path):
            generate_from_template(params, template_path, docx_abs, static_root=static_root)
        else:
            _generate_docx_fallback(params, docx_abs)

    if pdf:
        pdf_rel = str(techcard.pdf_file) if techcard.pdf_file else ''
        if not pdf_rel:
            uid = uuid.uuid4().hex[:10]
            card_num = (techcard.card_number or str(techcard.pk)).replace('/', '-').replace(' ', '_')
            date_dir = (
                techcard.created_at.strftime('%Y/%m')
                if getattr(techcard, 'created_at', None) else datetime.now().strftime('%Y/%m')
            )
            pdf_rel = f'techcards/pdf/{date_dir}/TC_{card_num}_{uid}.pdf'
            techcard.pdf_file = pdf_rel
        pdf_abs = os.path.join(media_root, pdf_rel)
        os.makedirs(os.path.dirname(pdf_abs), exist_ok=True)
        docx_abs_for_pdf = None
        if docx:
            docx_abs_for_pdf = os.path.join(media_root, str(techcard.docx_file))
        elif techcard.docx_file:
            candidate = os.path.join(media_root, str(techcard.docx_file))
            if os.path.isfile(candidate):
                docx_abs_for_pdf = candidate
        generate_pdf_for_techcard(
            params, pdf_abs, docx_abs_for_pdf,
            template_path=template_path, static_root=static_root,
        )

    update_fields = []
    if docx and techcard.docx_file:
        update_fields.append('docx_file')
    if pdf and techcard.pdf_file:
        update_fields.append('pdf_file')
    if update_fields:
        techcard.save(update_fields=update_fields)


def _generate_docx_fallback(params: dict, output_path: str) -> str:
    """
    Резервная генерация DOCX без шаблона.
    Создаёт простой документ с двумя колонками (метка / значение).
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('ТЕХНОЛОГИЧЕСКАЯ КАРТА РАДИОГРАФИЧЕСКОГО КОНТРОЛЯ')
    r.bold = True
    r.font.size = Pt(13)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f'№ {params.get("card_number","___")}     '
        f'Дата: {params.get("develop_date","")}     '
        f'Документ: {params.get("method_doc_cite", DOCUMENT_CODE)}'
    ).font.size = Pt(9)

    doc.add_paragraph()
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    vmap = _build_value_map(params)
    label_names = {
        '1.1': '1.1 Предприятие-изготовитель',
        '1.2': '1.2 Наименование объекта',
        '1.3': '1.3 № чертежа',
        '1.4': '1.4 Контролируемый элемент',
        '1.6': '1.6 Тип сварного соединения',
        '1.8': '1.8 Способ сварки',
        '1.9': '1.9 Основной металл',
        '2.1': '2.1 Методическая документация',
        '2.2': '2.2 Нормативная документация',
        '3.1': '3.1 Категория сварного соединения',
        '3.2': '3.2 Объём контроля',
        '4.1': '4.1 Тип контролируемого элемента',
        '4.2.1': '4.2.1 Наружный диаметр, мм',
        '4.2.1 S': 'Толщина стенки (S), мм',
        '5.1': '5.1 Источник излучения',
        '5.2': '5.2 Размер фокусного пятна, мм',
        '5.3': '5.3 Тип и номер ИКИ',
        '5.4': '5.4 Тип плёнки',
        '6.1': '6.1 Энергия / напряжение',
        '6.3': '6.3 Требуемая чувствительность (К)',
        '6.5': '6.5 Расстояние источник–детектор, мм',
        '6.6': '6.6 Число экспозиций, шт.',
        '6.9': '6.9 Схема просвечивания',
    }

    for key, label in label_names.items():
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = str(vmap.get(key, ''))
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    # Настраиваем колонтитулы (fallback без шаблона)
    _build_headers_footers_scratch(doc, params)

    doc.save(output_path)
    return output_path
