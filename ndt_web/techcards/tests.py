"""
Тесты приложения «Технологические карты».

Проверяет модели, генератор карт и нормативные данные.
"""

import os

from django.test import TestCase, Client
from django.urls import reverse
from django.contrib.auth import get_user_model

from .models import TechCard, NormativeDocument
from .generator import RadiographicTechCardCalculator
from accounts.models import UserBalance
from normative.gost_50_05_07 import (
    get_sensitivity, get_sensitivity_mm, calc_geometric_unsharpness,
    calc_min_sfd, get_suitable_sources,
)
from normative.np_105_18 import assess_defect, assess_multiple_defects

User = get_user_model()


class NormativeDataTests(TestCase):
    """
    Тесты нормативных данных.

    Чувствительность K — табличные значения из НП-105-18, Таблица 4.8.
    Функция get_sensitivity() возвращает абсолютное значение K в мм,
    не процент от толщины.
    """

    def test_sensitivity_cat_i_thin(self):
        """Категория I, S=4 мм → K=0,10 мм (Табл. 4.8 строка 3,0–4,5 мм)."""
        self.assertEqual(get_sensitivity(4.0, 'I'), 0.10)

    def test_sensitivity_cat_i_medium(self):
        """Категория I, S=8 мм → K=0,20 мм (Табл. 4.8 строка 7,5–10,0 мм)."""
        self.assertEqual(get_sensitivity(8.0, 'I'), 0.20)

    def test_suitable_sources_aluminum_differs_from_steel(self):
        """Для алюминия диапазоны источников отличаются от стали."""
        steel = {s['code'] for s in get_suitable_sources(25, 'steel')}
        aluminum = {s['code'] for s in get_suitable_sources(25, 'aluminum')}
        self.assertNotEqual(steel, aluminum)

    def test_suitable_sources_titanium_includes_ir192(self):
        """Для титана Ir-192 применим в диапазоне 10–40 мм."""
        sources = get_suitable_sources(20, 'titanium')
        self.assertIn('Ir-192', [s['code'] for s in sources])

    def test_sensitivity_cat_ii_medium(self):
        """Категория II, S=40 мм → K=0,60 мм (Табл. 4.8 строка 38,0–44,0 мм)."""
        self.assertEqual(get_sensitivity(40.0, 'II'), 0.60)

    def test_sensitivity_cat_iii_thick(self):
        """Категория III, S=150 мм → K=2,50 мм (Табл. 4.8 строка 130–165 мм)."""
        self.assertEqual(get_sensitivity(150.0, 'III'), 2.50)

    def test_sensitivity_mm_equals_sensitivity(self):
        """get_sensitivity_mm() возвращает то же значение что get_sensitivity()."""
        self.assertEqual(get_sensitivity_mm(10.0, 'I'), get_sensitivity(10.0, 'I'))

    def test_geometric_unsharpness_formula(self):
        """Формула геометрической нерезкости: Ug = d*b/(f-b)."""
        ug = calc_geometric_unsharpness(2.0, 700, 5)
        expected = 2.0 * 5 / (700 - 5)
        self.assertAlmostEqual(ug, expected, places=3)

    def test_geometric_unsharpness_invalid_sfd(self):
        """SFD ≤ OFD должен вызвать ValueError."""
        with self.assertRaises(ValueError):
            calc_geometric_unsharpness(2.0, 100, 200)

    def test_min_sfd_category_i(self):
        """Минимальное SFD рассчитывается правильно для категории I."""
        # Ug_max = 0.3, OFD=5, d=2
        min_sfd = calc_min_sfd(2.0, 5.0, 'I')
        expected = 5 * (2.0 + 0.3) / 0.3
        self.assertAlmostEqual(min_sfd, round(expected, 1), places=1)

    def test_select_film_size_for_length(self):
        """Типовой размер плёнки подбирается по длине участка L."""
        from normative.gost_50_05_07 import select_film_size_for_length
        self.assertEqual(select_film_size_for_length(80)['code'], '120x100')
        self.assertEqual(select_film_size_for_length(138)['code'], '240x100')
        self.assertEqual(select_film_size_for_length(400)['code'], '480x100')
        self.assertEqual(select_film_size_for_length(None)['code'], '240x100')

    def test_suitable_sources_iridium_steel_from_5mm(self):
        """Иридий-192 по табл. Б.1 (сталь) применим с 5 мм."""
        sources = get_suitable_sources(10, 'steel')
        codes = [s['code'] for s in sources]
        self.assertIn('Ir-192', codes)

    def test_suitable_sources_iridium_range(self):
        """Иридий-192 появляется в диапазоне 20–100 мм (сталь)."""
        sources = get_suitable_sources(30, 'steel')
        codes = [s['code'] for s in sources]
        self.assertIn('Ir-192', codes)

    def test_suitable_sources_tm170_thin(self):
        """Тулий-170 появляется для тонких изделий (сталь)."""
        sources = get_suitable_sources(5, 'steel')
        codes = [s['code'] for s in sources]
        self.assertIn('Tm-170', codes)

    def test_suitable_films_by_source_steel(self):
        """Плёнки для Ir-192 при 25 мм стали — по строке Se/Ir табл. Б.1."""
        from normative.gost_50_05_07 import get_suitable_films
        all_films = get_suitable_films(25, 'steel')
        ir_films = get_suitable_films(25, 'steel', 'Ir-192')
        self.assertIn('D5 "Структурикс"', ir_films)
        self.assertNotIn('D7 "Структурикс"', ir_films)
        self.assertIn('D7 "Структурикс"', get_suitable_films(25, 'steel', 'X-300kV'))
        self.assertGreater(len(all_films), len(ir_films))

    def test_suitable_films_aluminum_thin(self):
        """Для алюминия ≤5 мм — без NDT45 в списке."""
        from normative.gost_50_05_07 import get_suitable_films
        films = get_suitable_films(4, 'aluminum', 'Yb-169')
        self.assertIn('РТ-14', films)
        self.assertNotIn('NDT45 "Дюпонт"', films)


class QualityAssessmentLogicTests(TestCase):
    """
    Тесты логики оценки качества по НП-105-18 (с изм. 2024).

    Значения критериев берутся из Таблицы N 4.8 документа.
    """

    def test_crack_always_rejected(self):
        """Трещина недопустима в любой категории (п. 14 НП-105-18)."""
        result = assess_defect('crack', 'I', 10, size_1_mm=2.0)
        self.assertFalse(result['is_acceptable'])

    def test_lack_of_fusion_rejected(self):
        """Несплавление недопустимо (п. 14 НП-105-18)."""
        result = assess_defect('lack_of_fusion', 'II', 15)
        self.assertFalse(result['is_acceptable'])

    def test_pore_acceptable_cat_i_s10(self):
        """
        Пора 0.5 мм, категория I, S=10 мм.
        По Табл. 4.8: строка 7.5–10.0 мм → макс. включение 1.2 мм.
        0.5 ≤ 1.2 → допустимо.
        """
        result = assess_defect('pore', 'I', 10, size_1_mm=0.5)
        self.assertTrue(result['is_acceptable'])

    def test_pore_rejected_cat_i_s10(self):
        """
        Пора 4.0 мм, категория I, S=10 мм.
        Превышает допустимый размер крупных включений (3,5 мм, табл. 4.8).
        """
        result = assess_defect('pore', 'I', 10, size_1_mm=4.0)
        self.assertFalse(result['is_acceptable'])
        self.assertEqual(result['inclusion_group'], 'large')

    def test_pore_acceptable_cat_iii_s20(self):
        """
        Пора 2.5 мм, категория III, S=20 мм.
        По Табл. 4.8: строка 18.0–22.0 мм → макс. включение 3.0 мм.
        2.5 ≤ 3.0 → допустимо.
        """
        result = assess_defect('pore', 'III', 20, size_1_mm=2.5)
        self.assertTrue(result['is_acceptable'])

    def test_pore_rejected_cat_iii_s20(self):
        """
        Пора 8.0 мм, категория III, S=20 мм.
        Превышает допустимый размер крупных включений (7,0 мм, табл. 4.8).
        """
        result = assess_defect('pore', 'III', 20, size_1_mm=8.0)
        self.assertFalse(result['is_acceptable'])
        self.assertEqual(result['inclusion_group'], 'large')

    def test_undercut_depth_acceptable(self):
        """
        Подрез 0.3 мм, S=10 мм, кат. I.
        max_depth = min(0.1×10, 0.5) = 0.5 мм → 0.3 ≤ 0.5 → допустимо.
        """
        result = assess_defect('undercut', 'I', 10, size_1_mm=0.3)
        self.assertTrue(result['is_acceptable'])

    def test_undercut_depth_rejected(self):
        """
        Подрез 0.8 мм, S=5 мм, кат. I.
        max_depth = min(0.1×5, 0.5) = 0.5 мм → 0.8 > 0.5 → брак.
        """
        result = assess_defect('undercut', 'I', 5, size_1_mm=0.8)
        self.assertFalse(result['is_acceptable'])

    def test_multiple_defects_one_crack(self):
        """Одна трещина в наборе → общий вердикт БРАК."""
        defects = [
            {'type': 'crack', 'size_1': 5, 'size_2': 0, 'count': 1},
            {'type': 'pore', 'size_1': 0.5, 'size_2': 0, 'count': 1},
        ]
        result = assess_multiple_defects(defects, 'II', 10)
        self.assertFalse(result['is_acceptable'])
        self.assertEqual(result['verdict'], 'БРАК')

    def test_multiple_acceptable_defects(self):
        """Все допустимые дефекты → ГОДЕН."""
        defects = [
            {'type': 'pore', 'size_1': 0.3, 'size_2': 0, 'count': 1},
            {'type': 'undercut', 'size_1': 0.1, 'size_2': 0, 'count': 1},
        ]
        result = assess_multiple_defects(defects, 'II', 10)
        self.assertTrue(result['is_acceptable'])
        self.assertEqual(result['verdict'], 'ГОДЕН')

    def test_required_sensitivity_cat_i_s10(self):
        """Категория I, S=10 мм → K=0,20 мм (Табл. 4.8)."""
        from normative.np_105_18 import get_required_sensitivity
        self.assertEqual(get_required_sensitivity('I', 10.0), 0.20)

    def test_required_sensitivity_cat_iii_s150(self):
        """Категория III, S=150 мм → K=2,50 мм (Табл. 4.8 строка 130–165 мм)."""
        from normative.np_105_18 import get_required_sensitivity
        self.assertEqual(get_required_sensitivity('III', 150.0), 2.50)

    def test_pore_count_exceeds_max_per_100mm(self):
        """Превышение числа включений на участке 100 мм (табл. 4.8)."""
        result = assess_multiple_defects(
            [{'type': 'pore', 'size_1': 0.5, 'count': 1}],
            'I', 10,
            inclusion_cluster_count_100mm=15,
        )
        self.assertFalse(result['is_acceptable'])
        self.assertTrue(result['count_exceeded'])

    def test_pore_count_within_max_per_100mm(self):
        """Число включений на участке 100 мм в пределах нормы."""
        result = assess_multiple_defects(
            [{'type': 'pore', 'size_1': 0.5, 'count': 1}],
            'I', 10,
            inclusion_cluster_count_100mm=12,
        )
        self.assertTrue(result['is_acceptable'])

    def test_aggregate_inclusion_count_exceeded(self):
        """Число включений на участке 100 мм превышает норму."""
        result = assess_multiple_defects(
            [
                {'type': 'pore', 'size_1': 0.5, 'count': 1},
                {'type': 'slag', 'size_1': 0.8, 'count': 1},
            ],
            'I', 10,
            inclusion_cluster_count_100mm=13,
        )
        self.assertFalse(result['is_acceptable'])
        self.assertTrue(result['count_exceeded'])
        self.assertEqual(result['total_inclusion_count'], 13)

    def test_aggregate_inclusion_count_ok(self):
        """Число включений на участке 100 мм в пределах нормы."""
        result = assess_multiple_defects(
            [
                {'type': 'pore', 'size_1': 0.5, 'count': 1},
                {'type': 'slag', 'size_1': 0.8, 'count': 1},
            ],
            'I', 10,
            inclusion_cluster_count_100mm=12,
        )
        self.assertTrue(result['is_acceptable'])
        self.assertFalse(result['count_exceeded'])

    def test_segment_count_not_scaled_by_weld_length(self):
        """Длина шва не увеличивает допустимое число на участке 100 мм."""
        result = assess_multiple_defects(
            [{'type': 'pore', 'size_1': 0.5, 'count': 1}],
            'III', 10,
            weld_length_mm=300,
            inclusion_cluster_count_100mm=9,
        )
        self.assertTrue(result['is_acceptable'])
        self.assertEqual(result['max_inclusion_count_allowed'], 14)

    def test_large_inclusion_classified_by_size(self):
        """Пора 2,0 мм (кат. I, S=10) относится к крупным включениям."""
        result = assess_defect('pore', 'I', 10, size_1_mm=2.0, count=1)
        self.assertEqual(result['inclusion_group'], 'large')
        self.assertEqual(result['max_allowed_count'], 1)
        self.assertTrue(result['is_acceptable'])

    def test_large_inclusion_count_exceeds_limit(self):
        """Крупные включения: не более 1 шт. на участке 100,0 мм."""
        result = assess_multiple_defects(
            [{'type': 'pore', 'size_1': 2.0, 'count': 1}],
            'I', 10,
            large_inclusion_count_100mm=2,
        )
        self.assertFalse(result['is_acceptable'])
        self.assertTrue(result['count_exceeded'])

    def test_regular_inclusion_uses_regular_count_limit(self):
        """Мелкая пора допустима по размеру; число проверяется отдельным полем."""
        result = assess_defect('pore', 'I', 10, size_1_mm=0.5, count=12)
        self.assertEqual(result['inclusion_group'], 'regular')
        self.assertTrue(result['is_acceptable'])

    def test_aggregate_large_inclusions_separate(self):
        """Крупные включения оцениваются отдельным полем ввода."""
        result = assess_multiple_defects(
            [
                {'type': 'pore', 'size_1': 0.5, 'count': 1},
                {'type': 'pore', 'size_1': 2.0, 'count': 1},
            ],
            'I', 10,
            inclusion_cluster_count_100mm=6,
            large_inclusion_count_100mm=2,
        )
        self.assertFalse(result['is_acceptable'])
        self.assertTrue(result['count_exceeded'])
        self.assertEqual(result['total_large_inclusion_count'], 2)

    def test_aggregate_regular_and_large_both_ok(self):
        """Обычные и крупные включения в пределах нормы — ГОДЕН."""
        result = assess_multiple_defects(
            [{'type': 'pore', 'size_1': 0.5, 'count': 1}],
            'I', 10,
            inclusion_cluster_count_100mm=6,
            large_inclusion_count_100mm=1,
        )
        self.assertTrue(result['is_acceptable'])
        self.assertEqual(len(result['aggregates']), 2)


class Gost7512DefectNotationTests(TestCase):
    """Условная запись дефектов по ГОСТ 7512-82, приложения 5–6."""

    def test_five_spherical_pores(self):
        from normative.gost_7512 import format_gost_7512_defect_notation
        self.assertEqual(format_gost_7512_defect_notation('pore', 3, 0, 5), '5П3')

    def test_single_spherical_pore(self):
        from normative.gost_7512 import format_gost_7512_defect_notation
        self.assertEqual(format_gost_7512_defect_notation('pore', 3, 0, 1), 'П3')

    def test_elongated_slag(self):
        from normative.gost_7512 import format_gost_7512_defect_notation
        self.assertEqual(format_gost_7512_defect_notation('slag', 15, 2, 1), 'Ш15x2')

    def test_crack_length(self):
        from normative.gost_7512 import format_gost_7512_defect_notation
        self.assertEqual(format_gost_7512_defect_notation('crack', 40, 0, 1), 'Т40')

    def test_two_incomplete_penetrations(self):
        from normative.gost_7512 import format_gost_7512_defect_notation
        self.assertEqual(
            format_gost_7512_defect_notation('incomplete_penetration', 15, 0, 2),
            '2Н15',
        )

    def test_chain_notation_example(self):
        from normative.gost_7512 import format_gost_7512_defect_notation
        self.assertEqual(
            format_gost_7512_defect_notation(
                'pore', 30, 0, 1,
                morphology='chain',
                max_inclusion_l=5,
                max_inclusion_w=3,
            ),
            'Ц30П5x3',
        )

    def test_cluster_notation_example(self):
        from normative.gost_7512 import format_gost_7512_defect_notation
        self.assertEqual(
            format_gost_7512_defect_notation('pore', 10, 0.5, 2, morphology='cluster'),
            '2С10П0,5',
        )

    def test_assess_defect_includes_notation(self):
        result = assess_defect('pore', 'I', 10, size_1_mm=0.5, count=5)
        self.assertEqual(result.get('gost_notation'), '5П0,5')

    def test_combined_notation_in_multiple_defects(self):
        defects = [
            {'type': 'pore', 'size_1': 3, 'size_2': 0, 'count': 5},
            {'type': 'slag', 'size_1': 15, 'size_2': 2, 'count': 1},
        ]
        result = assess_multiple_defects(defects, 'I', 10)
        self.assertIn('5П3', result['combined_gost_notation'])
        self.assertIn('Ш15x2', result['combined_gost_notation'])


class TechCardCalculatorTests(TestCase):
    """Тесты вычислительного ядра генератора техкарт."""

    def setUp(self):
        self.base_input = {
            'organization': 'ТестОрг',
            'object_name': 'Трубопровод',
            'drawing_number': 'ТП-001',
            'weld_number': 'Ш-01',
            'card_number': 'ТК-001',
            'object_type': 'pipe',
            'material': '08Х18Н10Т',
            'wall_thickness': '10',
            'outer_diameter': '219.1',
            'weld_type': 'butt',
            'weld_category': 'I',
            'source_code': 'Ir-192',
            'focal_spot_mm': '2.0',
            'sfd_mm': '700',
            'ofd_mm': '5',
            'film_name': '',
            'inspector_name': '',
        }

    def test_weld_category_preserved(self):
        """Категория сварного соединения сохраняется в параметрах."""
        calc = RadiographicTechCardCalculator(self.base_input)
        params = calc.calculate()
        self.assertEqual(params['weld_category'], 'I')

    def test_film_size_auto_selected(self):
        """Размер плёнки подбирается автоматически по длине участка L."""
        calc = RadiographicTechCardCalculator(self.base_input)
        params = calc.calculate()
        L = params.get('L_calculated_mm')
        self.assertIsNotNone(L)
        # Для схемы 5a, D=219.1: L≈138 мм → 240×100
        self.assertEqual(params['film_length_mm'], 240)
        self.assertEqual(params['film_width_mm'], 100)
        self.assertEqual(params['film_size_label'], '240 × 100')

    def test_material_type_resolved(self):
        """Тип материала определяется для выбора источника."""
        data = dict(self.base_input, material='__titanium__', material_custom='ВТ6')
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        self.assertEqual(params['material_type'], 'titanium')
        self.assertIn('ВТ6', params['material_display'])

    def test_suitable_sources_have_table_b_metadata(self):
        """Источники содержат ссылку на табл. Б, без устаревших диапазонов."""
        from normative.gost_50_05_07 import get_suitable_sources
        sources = get_suitable_sources(10, 'steel')
        self.assertTrue(sources)
        self.assertIn('table_ref', sources[0])
        self.assertNotIn('thickness_min', sources[0])
        self.assertNotIn('is_optimal', sources[0])

    def test_single_xray_source_by_nomogram(self):
        """Рекомендуется один рентгеновский аппарат по рис. 6, п. 6.3.2."""
        from normative.gost_50_05_07 import (
            get_suitable_sources, get_max_xray_voltage_kv,
        )
        sources = get_suitable_sources(25, 'steel')
        xray = [s for s in sources if s.get('type') == 'xray']
        self.assertEqual(len(xray), 1)
        nomogram = get_max_xray_voltage_kv(25, 'steel')
        self.assertGreater(nomogram['max_voltage_kv'], 200)
        self.assertLessEqual(xray[0]['recommended_max_kv'], nomogram['max_voltage_kv'] + 1)
        self.assertIn('рис. 6', xray[0]['energy_display'])

    def test_iqi_wire_marking_two_digits(self):
        """Маркировка проволочного ИКИ: материал + номер эталона (ГОСТ 7512, п. 2.13)."""
        from normative.gost_7512 import get_wire_iqi, format_iqi_marking
        self.assertEqual(format_iqi_marking(1, 2), '12')
        iqi = get_wire_iqi(11.0, 0.20, material_type='steel')
        self.assertEqual(iqi['material_code'], 1)
        self.assertRegex(iqi['marking'], r'^1[1-4]$')
        self.assertIn('проволочный эталон', iqi['label'])

    def test_generator_uses_wire_iqi_marking(self):
        """Генератор заполняет маркировку проволочного ИКИ."""
        calc = RadiographicTechCardCalculator(self.base_input)
        params = calc.calculate()
        self.assertRegex(params.get('iqi_marking', ''), r'^\d{2,3}$')
        self.assertEqual(params['recommended_iqi']['code'], 'wire')

    def test_sensitivity_value_filled(self):
        """Значение чувствительности в % рассчитывается."""
        calc = RadiographicTechCardCalculator(self.base_input)
        params = calc.calculate()
        self.assertGreater(params['required_sensitivity_pct'], 0)
        self.assertGreater(params['required_sensitivity_mm'], 0)

    def test_geometric_unsharpness_computed(self):
        """Геометрическая нерезкость рассчитывается."""
        calc = RadiographicTechCardCalculator(self.base_input)
        params = calc.calculate()
        self.assertIn('geometric_unsharpness_mm', params)
        self.assertGreater(params['geometric_unsharpness_mm'], 0)

    def test_source_selected(self):
        """Источник излучения определяется."""
        calc = RadiographicTechCardCalculator(self.base_input)
        params = calc.calculate()
        self.assertIn('selected_source', params)
        self.assertTrue(params['selected_source'])

    def test_exposure_scheme_set(self):
        """Схема просвечивания определяется."""
        calc = RadiographicTechCardCalculator(self.base_input)
        params = calc.calculate()
        self.assertIn('exposure_scheme', params)
        self.assertTrue(params['exposure_scheme'])

    def test_no_errors_for_valid_input(self):
        """Корректные данные не генерируют ошибок."""
        calc = RadiographicTechCardCalculator(self.base_input)
        calc.calculate()
        self.assertEqual(len(calc.errors), 0)

    def test_warning_for_large_ug(self):
        """При большой нерезкости генерируется предупреждение."""
        data = dict(self.base_input, focal_spot_mm='5.0', sfd_mm='100', ofd_mm='50')
        calc = RadiographicTechCardCalculator(data)
        calc.calculate()
        # При SFD=100, OFD=50, d=5: Ug = 5*50/(100-50) = 5 >> 0.3 мм
        self.assertTrue(len(calc.warnings) > 0)

    def test_generator_iqi_film_side_in_sensitivity_desc(self):
        """Выбор стороны плёнки снижает K в п. 6.3 на одну ступень."""
        from techcards.generator import _build_value_map

        data = dict(
            self.base_input,
            scheme_type='5v',
            wall_thickness='12',
            outer_diameter='344',
            weld_category='III',
            iqi_side='film',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        self.assertEqual(params['iqi_placement']['side'], 'film')
        self.assertEqual(params['required_sensitivity_norm_mm'], 0.5)
        self.assertEqual(params['sensitivity_k_display_mm'], 0.4)
        self.assertIn('стороны плёнки', params['sensitivity_desc'])
        vm = _build_value_map(params)
        self.assertIn('0,400', vm['6.3'].replace('.', ','))
        self.assertIn('0,500', params['sensitivity_desc'].replace('.', ','))

    def test_field_6_3_k_unchanged_on_source_side(self):
        """При ИКИ со стороны источника п. 6.3 содержит нормативное K."""
        from techcards.generator import _build_value_map

        data = dict(
            self.base_input,
            scheme_type='5v',
            wall_thickness='12',
            outer_diameter='344',
            weld_category='III',
            iqi_side='source',
        )
        params = RadiographicTechCardCalculator(data).calculate()
        vm = _build_value_map(params)
        self.assertEqual(params['sensitivity_k_display_mm'], params['required_sensitivity_norm_mm'])
        self.assertIn('0,500', vm['6.3'].replace('.', ','))


class SchemeDisplayTests(TestCase):
    """Тесты пользовательских названий схем и расчётных правил."""

    def test_scheme_info_has_no_internal_codes(self):
        """В пользовательских названиях схем нет внутренних кодов 5a, 5б и т.д."""
        from normative.calculations import SCHEME_INFO
        for info in SCHEME_INFO.values():
            name = info.get('name', '')
            desc = info.get('description', '')
            self.assertNotIn('схема 5', name.lower())
            self.assertNotIn('источник снаружи', desc.lower())
            self.assertNotIn('источник внутри', desc.lower())

    def test_sk_two_walls_is_sum_of_thicknesses(self):
        """При двух стенках S_K = S + S (НП-105-18, п. 46)."""
        from normative.calculations import calc_radiation_thickness, resolve_table_b_thickness_mm
        rad = calc_radiation_thickness(10, 0.5, 3.5, '5v')
        self.assertEqual(rad['s_rad_k_mm'], 20.0)
        self.assertEqual(rad['formula_k'], '10 + 10 = 20.0 мм')
        self.assertEqual(rad['s_rad_f_mm'], 27.0)

        rad_b = resolve_table_b_thickness_mm(10, '5a', 'С-3', '30')
        self.assertEqual(rad_b['table_b_thickness_mm'], rad_b['s_rad_f_mm'])
        self.assertLess(rad_b['table_b_thickness_mm'], rad['s_rad_f_mm'])

    def test_sk_one_wall_with_backing(self):
        """При одной стенке с подкладкой S_K = S + g_min + Sпк."""
        from normative.calculations import calc_radiation_thickness
        rad = calc_radiation_thickness(4, 0.5, 2.0, '5a', backing_thickness_mm=4)
        self.assertEqual(rad['s_rad_k_mm'], 8.5)
        self.assertIn('4.0', rad['formula_k'])

    def test_example_s4_scheme_5g_cat_ii(self):
        """Пример из ТС: S=4 мм, схема 3г → S_K=8 мм, K=0,20 мм (кат. II)."""
        from normative.calculations import calc_radiation_thickness
        from normative.gost_50_05_07 import get_sensitivity
        rad = calc_radiation_thickness(4, 0.5, 2.0, '5g')
        self.assertEqual(rad['s_rad_k_mm'], 8.0)
        self.assertEqual(get_sensitivity(rad['s_rad_k_mm'], 'II'), 0.20)

    def test_iqi_film_side_user_choice(self):
        """При выборе «со стороны плёнки» ИКИ сдвигается на 1 ступень жёстче."""
        from normative.gost_7512 import get_wire_iqi, resolve_iqi_placement
        placement = resolve_iqi_placement(iqi_side='film')
        self.assertEqual(placement['side'], 'film')
        self.assertEqual(placement['shift_steps'], 1)
        base = get_wire_iqi(11.0, 0.20, shift_steps=0)
        shifted = get_wire_iqi(11.0, 0.20, shift_steps=1)
        self.assertLess(shifted['wire_diameter_mm'], base['wire_diameter_mm'])

    def test_negative_f_clamped_to_zero(self):
        """Отрицательное f в техкарте приводится к 0 мм."""
        from normative.calculations import clamp_f_mm, calc_exposure_parameters
        self.assertEqual(clamp_f_mm(-12.3), 0.0)
        result = calc_exposure_parameters(
            scheme='5g',
            focal_spot_mm=2.0,
            sensitivity_mm=0.5,
            d_outer_mm=50,
            d_inner_mm=30,
        )
        self.assertGreaterEqual(result.get('f_min_mm', 0), 0)

    def test_effective_outer_diameter_two_walls_includes_gmax(self):
        from normative.calculations import effective_outer_diameter_mm, calc_exposure_parameters

        self.assertEqual(effective_outer_diameter_mm(219.0, 3.5, '5v'), 226.0)
        self.assertEqual(effective_outer_diameter_mm(219.0, 3.5, '5a'), 219.0)

        nominal = calc_exposure_parameters(
            scheme='5v', focal_spot_mm=2.0, sensitivity_mm=0.5,
            d_outer_mm=219.0, d_inner_mm=200.0,
        )
        effective = calc_exposure_parameters(
            scheme='5v', focal_spot_mm=2.0, sensitivity_mm=0.5,
            d_outer_mm=226.0, d_inner_mm=200.0,
        )
        self.assertGreater(effective['f_min_mm'], nominal['f_min_mm'])

    def test_joint_c22_1_in_choices_and_weld_width(self):
        from normative.gost_59023_2 import (
            get_joint_type_choices, get_weld_width, get_joint_image_path,
        )

        codes = [c for c, _ in get_joint_type_choices()]
        self.assertIn('С-22-1', codes)
        self.assertEqual(get_joint_image_path('С-22-1'), 'gost/С_22_1.gif')

        weld = get_weld_width('С-22-1', 2.0)
        self.assertEqual(weld['e_mm'], 7.0)
        self.assertEqual(weld['g_nom'], 1.0)
        self.assertEqual(weld['g_min_mm'], 0.5)
        self.assertEqual(weld['g_max_mm'], 1.5)
        self.assertEqual(weld['e_tol_mm'], 2.0)
        self.assertEqual(weld['e_display'], '7,0 ±2,0')
        self.assertIn('9.29', weld['note'])

        weld_15 = get_weld_width('С-22-1', 1.5)
        self.assertEqual(weld_15['e_mm'], 6.0)
        self.assertEqual(weld_15['e_tol_mm'], 2.0)

        weld_35 = get_weld_width('С-22-1', 3.5)
        self.assertEqual(weld_35['e_mm'], 10.0)
        self.assertEqual(weld_35['g_nom'], 1.0)

    def test_calculator_uses_effective_diameter_for_two_wall_scheme(self):
        from techcards.generator import RadiographicTechCardCalculator

        data = {
            'object_type': 'pipe',
            'material': '12Х18Н10Т',
            'wall_thickness': 10.0,
            'outer_diameter': 219.0,
            'joint_designation': 'С-4',
            'welding_process': '30',
            'weld_category': 'II',
            'scheme_type': '5v',
            'source_code': 'X-300kV',
            'focal_spot_mm': 2.0,
        }
        params = RadiographicTechCardCalculator(data).calculate()
        g_max = params['g_max_mm']
        self.assertEqual(params['d_outer_effective_mm'], 219.0 + 2 * g_max)
        self.assertGreater(params.get('f_calculated_mm', 0), 0)


class NormativeDocumentModelTests(TestCase):
    """Тесты модели NormativeDocument."""

    def setUp(self):
        self.doc = NormativeDocument.objects.create(
            code='ГОСТ Р 50.05.07-2018',
            full_name='Радиографический контроль',
            control_method='RT',
            is_implemented=True,
            is_active=True,
        )

    def test_str_representation(self):
        """__str__ возвращает код документа."""
        self.assertEqual(str(self.doc), 'ГОСТ Р 50.05.07-2018')

    def test_method_display(self):
        """Метод контроля отображается корректно."""
        self.assertEqual(self.doc.get_control_method_display(), 'Радиографический контроль (РГК)')


class CalculationReferenceTests(TestCase):
    """Тесты генерации технической справки по расчётам."""

    def setUp(self):
        self.user = User.objects.create_user(
            'refuser', email='ref@test.com', password='pass123', email_verified=True,
        )
        self.doc = NormativeDocument.objects.create(
            code='ГОСТ Р 50.05.07-2018',
            full_name='Радиографический контроль',
            control_method='RT',
            is_implemented=True,
        )
        input_data = {
            'organization': 'АО Тест',
            'object_name': 'Трубопровод',
            'wall_thickness': 10,
            'outer_diameter': 108,
            'joint_designation': 'C1',
            'weld_category': 'II',
            'welding_process': '30',
            'source_code': 'Ir-192',
            'focal_spot_mm': 2.0,
            'ofd_mm': 5,
            'scheme_type': '5a',
            'card_number': 'TC-001',
        }
        calc = RadiographicTechCardCalculator(input_data)
        params = calc.calculate()
        self.techcard = TechCard.objects.create(
            user=self.user,
            normative_doc=self.doc,
            title='Трубопровод',
            card_number='TC-001',
            status=TechCard.STATUS_DONE,
            input_data=input_data,
            generated_data=params,
        )

    def test_build_calculation_log_has_sections(self):
        """Лог расчётов содержит основные секции."""
        from .calculation_reference import build_calculation_log
        sections = build_calculation_log(self.techcard.input_data, self.techcard.generated_data)
        titles = [s['title'] for s in sections]
        self.assertIn('4. Требуемая чувствительность K', titles)
        self.assertIn('7. Параметры схемы просвечивания', titles)

    def test_generate_reference_docx(self):
        """DOCX технической справки генерируется без ошибок."""
        from .calculation_reference import generate_calculation_reference_docx
        buffer = generate_calculation_reference_docx(
            self.techcard.input_data,
            self.techcard.generated_data,
            card_number=self.techcard.card_number,
        )
        self.assertGreater(len(buffer.getvalue()), 1000)

    def test_download_reference_endpoint(self):
        """Эндпоинт скачивания технической справки доступен владельцу."""
        self.client = Client()
        self.client.login(username='refuser', password='pass123')
        url = reverse('download_file', args=[self.techcard.pk, 'reference'])
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertIn(
            'wordprocessingml',
            response['Content-Type'],
        )

    def test_download_docx_regenerates_missing_file(self):
        """DOCX восстанавливается из JSON, если файл пропал с диска (Render)."""
        from django.conf import settings
        docx_rel = 'techcards/docx/2026/06/TC_regen_test.docx'
        self.techcard.docx_file = docx_rel
        self.techcard.save(update_fields=['docx_file'])
        docx_abs = os.path.join(settings.MEDIA_ROOT, docx_rel)
        if os.path.exists(docx_abs):
            os.remove(docx_abs)

        self.client.login(username='refuser', password='pass123')
        url = reverse('download_file', args=[self.techcard.pk, 'docx'])
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertTrue(os.path.exists(docx_abs))
        self.assertGreater(os.path.getsize(docx_abs), 1000)


class Step3AjaxTests(TestCase):
    """AJAX и предрасчёт данных шага 3."""

    def test_sources_ajax_requires_scheme(self):
        response = self.client.get('/ajax/sources/?thickness=10&material=08Х18Н10Т')
        data = response.json()
        self.assertEqual(data.get('error'), 'scheme_required')

    def test_sources_ajax_returns_radiation_thickness(self):
        response = self.client.get(
            '/ajax/sources/?thickness=10&material=08Х18Н10Т'
            '&scheme=5a&joint=C1&welding_process=30',
        )
        data = response.json()
        self.assertIn('sources', data)
        self.assertGreater(data.get('radiation_thickness_mm', 0), 10)
        self.assertIn('table_ref', data)

    def test_build_step3_scheme_data_has_sources_per_scheme(self):
        from techcards.views import build_step3_scheme_data
        data = build_step3_scheme_data(10, 'steel', 'C1', '30')
        self.assertIn('5a', data['sources_by_scheme'])
        self.assertTrue(data['sources_by_scheme']['5a'])
        self.assertGreater(
            data['radiation_by_scheme']['5v']['radiation_thickness_mm'],
            data['radiation_by_scheme']['5a']['radiation_thickness_mm'],
        )
        self.assertIn('Ir-192', data['films_by_scheme']['5a'])

    def test_build_step3_scheme_data_filtered_by_object_type(self):
        from techcards.views import build_step3_scheme_data
        flat_data = build_step3_scheme_data(10, 'steel', 'C1', '30', object_type='flat')
        self.assertIn('4_6', flat_data['sources_by_scheme'])
        self.assertNotIn('5a', flat_data['sources_by_scheme'])
        pipe_data = build_step3_scheme_data(10, 'steel', 'C1', '30', object_type='pipe')
        self.assertIn('5a', pipe_data['sources_by_scheme'])
        self.assertNotIn('4_6', pipe_data['sources_by_scheme'])

    def test_step3_form_scheme_limited_by_object_type(self):
        from techcards.forms import TechCardStep3Form
        flat_form = TechCardStep3Form(object_type='flat', wall_thickness=10)
        flat_codes = {c for c, _ in flat_form.fields['scheme_type'].choices if c}
        self.assertEqual(flat_codes, {'4_6'})
        pipe_form = TechCardStep3Form(object_type='pipe', wall_thickness=10)
        pipe_codes = {c for c, _ in pipe_form.fields['scheme_type'].choices if c}
        self.assertIn('5a', pipe_codes)
        self.assertNotIn('4_6', pipe_codes)

    def test_step3_form_rejects_pipe_scheme_for_flat_object(self):
        from techcards.forms import TechCardStep3Form
        form = TechCardStep3Form(
            {
                'scheme_type': '5a',
                'source_code': 'Ir-192',
                'focal_spot_mm': '2',
                'ofd_mm': '5',
                'iqi_side': 'source',
            },
            object_type='flat',
            wall_thickness=10,
        )
        self.assertFalse(form.is_valid())
        self.assertIn('scheme_type', form.errors)

    def test_scheme_5b_label_has_no_ellipse_method(self):
        from techcards.scheme_display import SCHEME_CHOICES, SCHEME_DESCRIPTIONS
        label_5b = dict(SCHEME_CHOICES).get('5b', '')
        self.assertNotIn('эллипс', label_5b.lower())
        self.assertNotIn('эллипс', SCHEME_DESCRIPTIONS.get('5b', '').lower())

    def test_scheme_preview_requires_login(self):
        response = self.client.get('/ajax/scheme-preview/?scheme_type=5a&source_code=Ir-192')
        self.assertEqual(response.status_code, 302)

    def test_scheme_preview_without_scheme(self):
        user = User.objects.create_user(
            'preview1', email='p1@test.com', password='pass123', email_verified=True,
        )
        self.client.login(username='preview1', password='pass123')
        response = self.client.get('/ajax/scheme-preview/')
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Выберите схему просвечивания')

    def test_scheme_preview_without_source_shows_scheme_image(self):
        user = User.objects.create_user(
            'preview2', email='p2@test.com', password='pass123', email_verified=True,
        )
        self.client.login(username='preview2', password='pass123')
        session = self.client.session
        session['techcard_data'] = {
            'wall_thickness': 10,
            'outer_diameter': 219.1,
            'material': '08Х18Н10Т',
            'joint_designation': 'C1',
            'welding_process': '30',
            'weld_category': 'II',
            'object_type': 'pipe',
        }
        session.save()
        response = self.client.get('/ajax/scheme-preview/?scheme_type=5a')
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Выберите источник излучения')
        self.assertContains(response, 'scheme_5a')

    def test_scheme_preview_returns_calculated_values(self):
        user = User.objects.create_user(
            'preview3', email='p3@test.com', password='pass123', email_verified=True,
        )
        self.client.login(username='preview3', password='pass123')
        session = self.client.session
        session['techcard_data'] = {
            'wall_thickness': 10,
            'outer_diameter': 219.1,
            'material': '08Х18Н10Т',
            'joint_designation': 'C1',
            'welding_process': '30',
            'weld_category': 'II',
            'object_type': 'pipe',
        }
        session.save()
        response = self.client.get(
            '/ajax/scheme-preview/?scheme_type=5a&source_code=Ir-192'
            '&focal_spot_mm=2.0&ofd_mm=5&iqi_side=source',
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Предпросмотр расчёта')
        self.assertContains(response, 'preview-metric-label')
        self.assertContains(response, 'Ir-192')

    def test_build_scheme_preview_context_ready(self):
        from techcards.views import build_scheme_preview_context
        from django.test import RequestFactory

        factory = RequestFactory()
        request = factory.get(
            '/ajax/scheme-preview/?scheme_type=5a&source_code=Ir-192'
            '&focal_spot_mm=2.0&ofd_mm=5',
        )
        request.session = self.client.session
        request.session['techcard_data'] = {
            'wall_thickness': 10,
            'outer_diameter': 219.1,
            'material': '08Х18Н10Т',
            'joint_designation': 'C1',
            'welding_process': '30',
            'weld_category': 'II',
            'object_type': 'pipe',
        }
        context = build_scheme_preview_context(request)
        self.assertTrue(context['ready'])
        self.assertIsNotNone(context.get('f_mm'))
        self.assertIsNotNone(context.get('N'))
        self.assertIsNotNone(context.get('L_mm'))
        self.assertIsNotNone(context.get('k_mm'))

    def test_scheme_preview_k_changes_with_iqi_film_side(self):
        """K в предпросмотре отражает диаметр проволоки ИКИ при стороне плёнки."""
        user = User.objects.create_user(
            'preview4', email='p4@test.com', password='pass123', email_verified=True,
        )
        self.client.login(username='preview4', password='pass123')
        session = self.client.session
        session['techcard_data'] = {
            'wall_thickness': 12,
            'outer_diameter': 344,
            'material': '08Х17Н15М3Т',
            'joint_designation': 'C1',
            'welding_process': '30',
            'weld_category': 'III',
            'object_type': 'pipe',
        }
        session.save()
        base_qs = (
            'scheme_type=5v&source_code=Ir-192'
            '&focal_spot_mm=3.0&ofd_mm=5'
        )
        src = self.client.get(f'/ajax/scheme-preview/?{base_qs}&iqi_side=source')
        film = self.client.get(f'/ajax/scheme-preview/?{base_qs}&iqi_side=film')
        self.assertEqual(src.status_code, 200)
        self.assertEqual(film.status_code, 200)
        self.assertContains(src, '0,500')
        self.assertContains(film, '0,400')
        self.assertContains(film, 'жёстче на 1 ступень')


class DocxToPdfTests(TestCase):
    """Тесты конвертации DOCX → PDF (mammoth + xhtml2pdf)."""

    def setUp(self):
        self.input_data = {
            'organization': 'АО Тест',
            'object_name': 'Трубопровод',
            'wall_thickness': 10,
            'outer_diameter': 219.1,
            'joint_designation': 'C1',
            'weld_category': 'II',
            'welding_process': '30',
            'source_code': 'Ir-192',
            'focal_spot_mm': 2.0,
            'ofd_mm': 5,
            'scheme_type': '5a',
            'card_number': 'TC-PDF-TEST',
            'material': '08Х18Н10Т',
        }
        calc = RadiographicTechCardCalculator(self.input_data)
        self.params = calc.calculate()

    def test_convert_docx_to_pdf_creates_valid_pdf(self):
        """DOCX конвертируется в валидный PDF с кириллицей."""
        import tempfile
        from common.docx_to_pdf import convert_docx_to_pdf
        from techcards.generator import _generate_docx_fallback

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'card.docx')
            pdf_path = os.path.join(tmpdir, 'card.pdf')
            _generate_docx_fallback(self.params, docx_path)
            convert_docx_to_pdf(docx_path, pdf_path)
            self.assertTrue(os.path.isfile(pdf_path))
            self.assertGreater(os.path.getsize(pdf_path), 1000)
            with open(pdf_path, 'rb') as pdf_file:
                self.assertTrue(pdf_file.read(4).startswith(b'%PDF'))

    def test_generate_pdf_for_techcard_from_docx(self):
        """generate_pdf_for_techcard создаёт PDF из готового DOCX."""
        import tempfile
        from techcards.generator import _generate_docx_fallback, generate_pdf_for_techcard

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'card.docx')
            pdf_path = os.path.join(tmpdir, 'card.pdf')
            _generate_docx_fallback(self.params, docx_path)
            generate_pdf_for_techcard(self.params, pdf_path, docx_path)
            self.assertTrue(os.path.isfile(pdf_path))
            self.assertGreater(os.path.getsize(pdf_path), 1000)
            with open(pdf_path, 'rb') as pdf_file:
                self.assertTrue(pdf_file.read(4).startswith(b'%PDF'))

    def test_generate_pdf_fallback_on_corrupt_docx(self):
        """При битом DOCX используется ReportLab fallback."""
        import tempfile
        from techcards.generator import generate_pdf_for_techcard

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'bad.docx')
            pdf_path = os.path.join(tmpdir, 'card.pdf')
            with open(docx_path, 'wb') as bad_file:
                bad_file.write(b'not a docx')
            generate_pdf_for_techcard(self.params, pdf_path, docx_path)
            self.assertTrue(os.path.isfile(pdf_path))
            with open(pdf_path, 'rb') as pdf_file:
                self.assertTrue(pdf_file.read(4).startswith(b'%PDF'))


class WeldingProcessFilterTests(TestCase):
    """Фильтрация способов сварки по типу соединения (ГОСТ Р 59023.2-2020)."""

    def test_choices_limited_to_joint_methods(self):
        from normative.gost_59023_2 import get_welding_process_choices_for_joint
        codes = [c for c, _ in get_welding_process_choices_for_joint('С-1') if c]
        self.assertEqual(set(codes), {'53', '10'})

    def test_form_rejects_invalid_welding_process(self):
        from techcards.forms import TechCardStep2Form
        form = TechCardStep2Form({
            'object_type': 'flat',
            'material': '08Х18Н10Т',
            'wall_thickness': '10',
            'joint_designation': 'С-1',
            'welding_process': '30',
            'weld_category': 'II',
        })
        self.assertFalse(form.is_valid())
        self.assertIn('welding_process', form.errors)

    def test_weld_category_label_np105(self):
        from techcards.forms import TechCardStep2Form
        field = TechCardStep2Form().fields['weld_category']
        self.assertIn('НП-105-18', field.label)
        self.assertFalse(field.help_text)

    def test_weld_category_in_reload_after_session(self):
        """Iн/IIн доступны при возврате на шаг 2 из сессии."""
        from techcards.forms import TechCardStep2Form
        from techcards.views import _step2_form_initial
        from django.test import RequestFactory

        session = {
            'object_type': 'pipe',
            'material': '12Х18Н10Т',
            'wall_thickness': 20,
            'outer_diameter': 219,
            'joint_designation': 'С-4',
            'welding_process': '30',
            'weld_category': 'Iн',
            'joint_mobility': 'non_rotating',
        }
        form = TechCardStep2Form(initial=session)
        codes = [c[0] for c in form.fields['weld_category'].choices]
        self.assertIn('Iн', codes)
        self.assertIn('IIн', codes)

        factory = RequestFactory()
        request = factory.get('/')
        request.session = {}
        request.session['techcard_data'] = session
        initial = _step2_form_initial(request)
        form2 = TechCardStep2Form(initial=initial)
        codes2 = [c[0] for c in form2.fields['weld_category'].choices]
        self.assertIn('Iн', codes2)

    def test_weld_category_in_with_material_custom_only(self):
        from techcards.forms import TechCardStep2Form
        form = TechCardStep2Form(initial={
            'object_type': 'pipe',
            'material_custom': '12Х18Н10Т',
            'material': '12Х18Н10Т',
            'weld_category': 'IIн',
            'wall_thickness': 10,
            'joint_designation': 'С-4',
            'welding_process': '30',
        })
        codes = [c[0] for c in form.fields['weld_category'].choices]
        self.assertIn('IIн', codes)

    def test_weld_category_titanium_excludes_in(self):
        """Для титана доступны только I, II, III (табл. 4.11)."""
        from techcards.forms import TechCardStep2Form
        from normative.gost_59023_2 import MATERIAL_TITANIUM

        form = TechCardStep2Form(initial={
            'object_type': 'pipe',
            'material': MATERIAL_TITANIUM,
            'material_custom': 'ВТ6',
            'weld_category': 'III',
            'wall_thickness': 12,
            'joint_designation': 'С-4',
            'welding_process': '30',
        })
        codes = [c[0] for c in form.fields['weld_category'].choices]
        self.assertEqual(codes, ['I', 'II', 'III'])

    def test_weld_category_aluminum_excludes_in(self):
        from techcards.forms import TechCardStep2Form
        from normative.gost_59023_2 import MATERIAL_ALUMINUM

        form = TechCardStep2Form(initial={
            'object_type': 'flat',
            'material': MATERIAL_ALUMINUM,
            'material_custom': 'АМг6',
            'weld_category': 'II',
            'wall_thickness': 10,
            'joint_designation': 'С-1',
            'welding_process': '53',
        })
        codes = [c[0] for c in form.fields['weld_category'].choices]
        self.assertNotIn('Iн', codes)
        self.assertNotIn('IIн', codes)


class DocumentKindTests(TestCase):
    """Вид документа: методический vs нормативный."""

    def test_gost_is_methodological(self):
        doc = NormativeDocument.objects.create(
            code='ГОСТ Р 50.05.07-2018',
            full_name='Радиографический контроль',
            control_method='RT',
            document_kind=NormativeDocument.KIND_METHODOLOGICAL,
        )
        self.assertEqual(doc.get_document_kind_display(), 'Методический документ')

    def test_joint_gost_image_is_weld_cross_section(self):
        """Изображение шва — из столбца «шва сварного соединения», не кромок."""
        from normative.gost_59023_2 import get_joint_image_path
        path = get_joint_image_path('С-1')
        self.assertEqual(path, 'gost/С_1.gif')
        full = os.path.join(
            os.path.dirname(__file__), '..', 'static', 'img', 'welds', path,
        )
        self.assertTrue(os.path.isfile(full), msg=full)
        # Размер соответствует извлечению из правого столбца таблицы 9.1
        self.assertGreater(os.path.getsize(full), 1900)


class SessionSerializationTests(TestCase):
    """Сессия техкарты должна сохранять только JSON-совместимые типы."""

    def test_session_safe_data_converts_dates(self):
        from techcards.views import _session_safe_data
        from datetime import date, datetime
        result = _session_safe_data({
            'develop_date': date(2026, 6, 21),
            'checked_at': datetime(2026, 6, 21, 12, 30),
        })
        self.assertEqual(result['develop_date'], '2026-06-21')
        self.assertEqual(result['checked_at'], '2026-06-21T12:30:00')

    def test_step1_post_stores_dates_in_session(self):
        user = User.objects.create_user(
            'step1user', email='step1@test.com', password='pass123', email_verified=True,
        )
        UserBalance.objects.create(user=user, techcard_credits=10)
        NormativeDocument.objects.create(
            code='ГОСТ Р 50.05.07-2018',
            full_name='Радиографический контроль',
            control_method='RT',
            is_implemented=True,
        )
        self.client.login(username='step1user', password='pass123')
        response = self.client.post(
            '/techcards/create/ГОСТ Р 50.05.07-2018/step1/',
            {
                'object_name': 'Трубопровод',
                'weld_number': 'Ш-01',
                'develop_date': '2026-06-21',
                'check_date': '2026-06-21',
            },
        )
        self.assertEqual(response.status_code, 302)
        session_data = self.client.session.get('techcard_data', {})
        self.assertEqual(session_data.get('develop_date'), '2026-06-21')
        self.assertEqual(session_data.get('check_date'), '2026-06-21')


class TemplateCommentsTests(TestCase):
    """Требования из комментариев к шаблону техкарты DOCX."""

    def setUp(self):
        self.pipe_input = {
            'organization': 'ТестОрг',
            'object_name': 'Трубопровод',
            'drawing_number': 'ТП-001',
            'weld_number': 'Ш-01',
            'card_number': 'ТК-001',
            'object_type': 'pipe',
            'material': '12Х18Н10Т',
            'wall_thickness': '20',
            'outer_diameter': '800',
            'joint_designation': 'С-4',
            'welding_process': '30',
            'weld_category': 'II',
            'joint_mobility': 'non_rotating',
            'source_code': 'Se-75',
            'focal_spot_mm': '',
            'scheme_type': '5a',
            'iqi_side': 'source',
        }

    def test_welding_material_defaults_to_base_metal(self):
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        self.assertEqual(params['weld_material'], '12Х18Н10Т')

    def test_welding_material_same_as_base_sentinel(self):
        data = dict(self.pipe_input, welding_material='same_as_base')
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        self.assertEqual(params['weld_material'], '12Х18Н10Т')

    def test_dates_formatted_ddmmyyyy_from_iso(self):
        data = dict(
            self.pipe_input,
            develop_date='2026-06-21',
            check_date='2026-12-31',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        self.assertEqual(params['develop_date'], '21.06.2026')
        self.assertEqual(params['check_date'], '31.12.2026')

    def test_label_110_does_not_match_11(self):
        from techcards.generator import _label_matches_value_key, _match_value_for_label
        self.assertTrue(_label_matches_value_key('1.1 Предприятие', '1.1'))
        self.assertFalse(_label_matches_value_key('1.10. Марка сварочного материала', '1.1'))
        self.assertTrue(_label_matches_value_key('1.10. Марка сварочного материала', '1.10'))
        vmap = {'1.1': 'ORG', '1.10': 'WELD', '1.9': 'BASE'}
        self.assertEqual(
            _match_value_for_label('1.10. Марка сварочного материала', vmap),
            'WELD',
        )

    def test_label_22_does_not_match_422(self):
        from techcards.generator import _label_matches_value_key, _match_value_for_label
        self.assertFalse(_label_matches_value_key('4.2.2 Длинна (листы/пластины)', '2.2'))
        vmap = {'2.2': 'НП-105-18', '4.2.2': '15,0 ±4,0', '4.2.2 длин': '—'}
        self.assertIsNone(
            _match_value_for_label('4.2.2 Длинна (листы/пластины)', vmap),
        )
        self.assertIsNone(
            _match_value_for_label(
                '4.2.2. Ширина валиков усиления на наружной поверхности (е):',
                vmap,
            ),
        )

    def test_section_42_gost_dimensions_in_value_map(self):
        from normative.gost_59023_2 import get_inspection_zone
        from techcards.generator import _build_value_map, _fmt_mm

        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        zone = get_inspection_zone('С-4', 20.0, '30')

        self.assertEqual(params['e_display'], zone['e_display'])
        self.assertEqual(params['e1_display'], zone['e1_display'])
        self.assertEqual(params['g_display'], zone['g_display'])
        self.assertEqual(params['haz_width_mm'], zone['haz_width_mm'])
        self.assertEqual(params['zone_width_mm'], zone['zone_width_mm'])

        vmap = _build_value_map(params)
        self.assertEqual(vmap['4.2.2'], zone['e_display'])
        self.assertEqual(vmap['4.2.2 e1'], zone['e1_display'])
        self.assertEqual(vmap['4.2.3'], zone['g_display'])
        self.assertIn(_fmt_mm(zone['haz_width_mm']), vmap['4.2.4'])
        self.assertEqual(vmap['4.2.5'], _fmt_mm(zone['zone_width_mm']))

    def test_section_42_gost_dimensions_in_docx(self):
        from techcards.generator import (
            generate_from_template, get_default_template_path, _unique_cells, _fmt_mm,
        )
        from normative.gost_59023_2 import get_inspection_zone
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')

        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        zone = get_inspection_zone('С-4', 20.0, '30')

        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            rows = {}
            for row in doc.tables[0].rows:
                ucells = _unique_cells(row)
                if ucells:
                    rows[ucells[0].text.strip().lower()] = [c.text.strip() for c in ucells]

            length_row = next(
                v for k, v in rows.items() if '4.2.2' in k and 'длинна' in k
            )
            self.assertEqual(length_row[1], '—')
            # В 4-й колонке шаблона — ссылка на НП-105-18 (не подменяется генератором)
            self.assertEqual(length_row[-1], 'НП-105-18')

            width_row = next(
                v for k, v in rows.items() if 'наружной поверхности' in k
            )
            self.assertEqual(width_row[1], zone['e_display'])
            self.assertEqual(width_row[3], zone['e1_display'])

            height_row = next(v for k, v in rows.items() if '4.2.3' in k and 'высота' in k)
            self.assertEqual(height_row[-1], zone['g_display'])

            haz_row = next(v for k, v in rows.items() if '4.2.4' in k and 'околошовной' in k)
            self.assertIn(_fmt_mm(zone['haz_width_mm']), haz_row[1])
            self.assertEqual(haz_row[3], _fmt_mm(zone['zone_width_mm']))

    def test_titanium_haz_width_in_calculator(self):
        from normative.gost_59023_2 import get_inspection_zone

        data = dict(
            self.pipe_input,
            material='__titanium__',
            material_custom='ПТ-3В',
            wall_thickness=3,
            outer_diameter=58,
            joint_designation='С-1',
            welding_process='53',
            weld_category='III',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        zone = get_inspection_zone(
            'С-1', 3.0, '53', material_type='titanium',
        )
        self.assertEqual(params['haz_width_mm'], 20.0)
        self.assertEqual(params['zone_width_mm'], zone['zone_width_mm'])

    def test_isotope_focal_spot_defaults_to_3mm(self):
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        self.assertEqual(params['source_focal_spot_mm'], 3.0)

    def test_pipe_length_field_is_dash(self):
        from techcards.generator import _build_value_map
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        vmap = _build_value_map(params)
        self.assertEqual(vmap['4.2.2 длин'], '—')
        self.assertEqual(vmap['4.2.1'], '800,0')

    def test_flat_outer_diameter_is_dash(self):
        from techcards.generator import _build_value_map
        data = dict(
            self.pipe_input,
            object_type='flat',
            outer_diameter='',
            flat_length_mm='1500',
            scheme_type='4_6',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        vmap = _build_value_map(params)
        self.assertEqual(vmap['4.2.1'], '—')
        self.assertEqual(vmap['4.2.2 длин'], '1500,0')

    def test_joint_mobility_in_section_16(self):
        from techcards.generator import _build_value_map
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        vmap = _build_value_map(params)
        self.assertIn('неповоротное', vmap['1.6'])

    def test_reinforcement_removed_zeros_g(self):
        data = dict(self.pipe_input, reinforcement_removed=True)
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        self.assertEqual(params['g_min_mm'], 0.0)
        self.assertEqual(params['g_max_mm'], 0.0)
        self.assertEqual(params['reinforcement_status'], 'снят')

    def test_e_and_e1_display_present(self):
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        self.assertIn('±', params['e_display'])
        self.assertIn('±', params['e1_display'])

    def test_backing_ring_thickness_used(self):
        data = dict(self.pipe_input, has_backing_ring=True, backing_ring_thickness_mm='3')
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        self.assertTrue(params['has_backing'])
        self.assertEqual(params['backing_thickness_mm'], 3.0)

    def test_generate_from_template_fills_dimension_rows(self):
        from techcards.generator import (
            generate_from_template, get_default_template_path, _paragraph_has_drawing,
        )
        import tempfile
        from docx import Document
        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            self.assertGreater(os.path.getsize(out), 5000)
            doc = Document(out)
            table_text = ' '.join(
                cell.text for table in doc.tables for row in table.rows for cell in row.cells
            )
            self.assertIn('12Х18Н10Т', table_text)
            self.assertIn('неповоротное', table_text)
            self.assertIn('3,0', table_text)

            sketch_heading = next(
                (p for p in doc.paragraphs if '4.3' in p.text and 'Эскиз' in p.text),
                None,
            )
            self.assertIsNotNone(sketch_heading)
            caption = next(
                (p for p in doc.paragraphs if 'Сварное соединение С-4' in p.text),
                None,
            )
            self.assertIsNotNone(caption)
            self.assertTrue(any(_paragraph_has_drawing(p) for p in doc.paragraphs))

    def test_template_title_page_and_footers(self):
        """Титульный лист и разные колонтитулы по шаблону normative_docs."""
        from techcards.generator import (
            generate_from_template, get_default_template_path,
        )
        import tempfile
        import zipfile
        import re
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        data = dict(
            self.pipe_input,
            developed_by_name='Петров П.П.',
            checked_by_name='Сидоров С.С.',
            developed_by_certificate='12345',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            section = doc.sections[0]
            self.assertTrue(section.different_first_page_header_footer)

            body_title = [
                p.text.strip() for p in doc.paragraphs
                if p.text.strip().startswith('Технологическая карта')
                or p.text.strip() == 'радиографического контроля'
                or p.text.strip().startswith('№')
            ]
            self.assertIn('Технологическая карта', body_title[0])
            self.assertIn('радиографического контроля', body_title)
            num_lines = [t for t in body_title if t.startswith('№')]
            self.assertEqual(len(num_lines), 1)
            self.assertEqual(num_lines[0], '№ ТК-001')

            fp_footer_table = section.first_page_footer.tables[0]
            dev_name_para = fp_footer_table.rows[2].cells[0].paragraphs[0]
            self.assertEqual(dev_name_para.alignment, WD_ALIGN_PARAGRAPH.RIGHT)

            header_text = ' '.join(
                cell.text for table in section.header.tables
                for row in table.rows for cell in row.cells
            )
            self.assertIn('ТестОрг', header_text)
            self.assertNotIn('ФГУП МАРКС', header_text)
            self.assertNotIn('Иванов', header_text)
            self.assertIn('страница', header_text.lower())
            self.assertIn('страниц', header_text.lower())

            body = doc.element.body
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            first_tbl = next(
                c for c in body if c.tag.split('}')[-1] == 'tbl'
            )
            has_title_break = False
            for child in body:
                if child is first_tbl:
                    break
                if child.tag.split('}')[-1] != 'p':
                    continue
                br = child.find(f'.//{ns}br')
                if br is not None and br.get(qn('w:type')) == 'page':
                    has_title_break = True
                    break
            self.assertTrue(has_title_break)
            for child in list(body):
                if child is first_tbl:
                    break
                self.assertNotEqual(child.tag.split('}')[-1], 'tbl')

            with zipfile.ZipFile(out) as zf:
                header_xml = zf.read('word/header1.xml').decode('utf-8')
            self.assertIn('страница', header_xml.lower())
            self.assertIn('NUMPAGES', header_xml)
            self.assertIn('PAGE', header_xml)

            default_footer = ' '.join(
                cell.text for table in section.footer.tables
                for row in table.rows for cell in row.cells
            )
            self.assertIn('Ведущий инженер технолог', default_footer)
            self.assertIn('Петров П.П.', default_footer)
            self.assertNotIn('(дата)', default_footer)

            first_footer = ' '.join(
                cell.text for table in section.first_page_footer.tables
                for row in table.rows for cell in row.cells
            )
            self.assertIn('(дата)', first_footer)
            self.assertIn('(подпись)', first_footer)
            self.assertIn('Петров П.П.', first_footer)
            self.assertIn('Сидоров С.С.', first_footer)
            self.assertIn('Ведущий инженер', first_footer)
            self.assertNotIn('Ведущий инженер технолог', first_footer)

            with zipfile.ZipFile(out) as zf:
                sect_xml = re.search(
                    r'<w:sectPr.*?</w:sectPr>',
                    zf.read('word/document.xml').decode('utf-8'),
                    re.DOTALL,
                ).group(0)
            self.assertIn('titlePg', sect_xml)
            self.assertIn('w:type="first"', sect_xml)

    def test_default_template_is_tc_reference(self):
        from techcards.generator import get_default_template_path
        path = get_default_template_path()
        self.assertIsNotNone(path)
        self.assertIn('card_templates', path)
        self.assertTrue(os.path.basename(path).startswith('TC_'))

    def test_template_spacing_preserved(self):
        """Генератор не добавляет лишних разрывов — только отступы из шаблона."""
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        from docx import Document
        from docx.oxml.ns import qn

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')

        def _layout_stats(path):
            doc = Document(path)
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            page_breaks = 0
            empty = 0
            for child in doc.element.body:
                if child.tag.split('}')[-1] != 'p':
                    continue
                text = ''.join(t.text or '' for t in child.findall(f'.//{ns}t')).strip()
                has_br = any(
                    br.get(qn('w:type')) == 'page' for br in child.findall(f'.//{ns}br')
                )
                if has_br:
                    page_breaks += 1
                elif not text:
                    empty += 1
            return page_breaks, empty

        tpl_breaks, tpl_empty = _layout_stats(template)
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            gen_breaks, gen_empty = _layout_stats(out)
            # +2 разрыва: п. 4.3 (стр. 3) и п. 6.9 (схема просвечивания)
            self.assertEqual(gen_breaks, tpl_breaks + 2)
            # Пустые строки схлопнуты — не более одной подряд
            self.assertLess(gen_empty, tpl_empty)
            doc = Document(out)
            from techcards.generator import _is_empty_body_paragraph
            max_run = 0
            cur = 0
            for child in doc.element.body:
                if child.tag.split('}')[-1] == 'p' and _is_empty_body_paragraph(child):
                    cur += 1
                    max_run = max(max_run, cur)
                else:
                    cur = 0
            self.assertLessEqual(max_run, 1)

    def test_section_43_single_line_gap_before_heading(self):
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            para_43 = next(
                p for p in doc.paragraphs if '4.3' in p.text and 'эскиз' in p.text.lower()
            )
            empty = 0
            prev = para_43._element.getprevious()
            while prev is not None and prev.tag.endswith('p'):
                from techcards.generator import _is_empty_body_paragraph
                if not _is_empty_body_paragraph(prev):
                    break
                empty += 1
                prev = prev.getprevious()
            self.assertEqual(empty, 1)

    def test_section_65_shows_f_value_only(self):
        from techcards.generator import _build_value_map
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        vmap = _build_value_map(params)
        self.assertIn('f =', vmap['6.5'])
        self.assertNotIn('расчёт:', vmap['6.5'].lower())
        self.assertNotIn('scheme_formula', vmap['6.5'])

    def test_generated_docx_has_materials_and_no_comments(self):
        """П. 1.9/1.10 и отсутствие служебных комментариев в готовой техкарте."""
        from techcards.generator import (
            generate_from_template, get_default_template_path,
        )
        import tempfile
        import zipfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        data = dict(
            self.pipe_input,
            material='12Х18Н10Т',
            welding_material='same_as_base',
            develop_date='2026-06-21',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            def _row_value(table, fragment):
                for row in table.rows:
                    if fragment in row.cells[0].text:
                        return row.cells[-1].text.strip()
                return ''
            main = doc.tables[0]
            self.assertEqual(_row_value(main, '1.9'), '12Х18Н10Т')
            self.assertEqual(_row_value(main, '1.10'), '12Х18Н10Т')
            with zipfile.ZipFile(out) as zf:
                doc_xml = zf.read('word/document.xml').decode('utf-8')
            self.assertNotIn('commentRangeStart', doc_xml)
            self.assertNotIn('commentReference', doc_xml)
            self.assertEqual(params['develop_date'], '21.06.2026')

    def test_generated_docx_preserves_reference_layout(self):
        """Шрифт 12 pt в таблице, поля и нумерация как в эталонной техкарте."""
        from techcards.generator import (
            generate_from_template, get_default_template_path,
        )
        import tempfile
        import zipfile
        from docx import Document
        from docx.shared import Pt

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            section = doc.sections[0]
            self.assertAlmostEqual(section.top_margin.mm, 15.0, delta=0.5)
            self.assertAlmostEqual(section.left_margin.mm, 20.0, delta=0.5)

            value_cell = None
            for row in doc.tables[0].rows:
                if row.cells[0].text.strip().startswith('1.1'):
                    value_cell = row.cells[-1]
                    break
            self.assertIsNotNone(value_cell)
            run = value_cell.paragraphs[0].runs[0]
            self.assertIsNotNone(run.font.size)
            self.assertGreaterEqual(run.font.size.pt, 12.0)

            with zipfile.ZipFile(out) as zf:
                header_xml = zf.read('word/header1.xml').decode('utf-8')
            self.assertIn('страница', header_xml.lower())
            self.assertNotIn('Страница    ', header_xml)
            self.assertIn('NUMPAGES', header_xml)

    def test_section_102_contains_acceptance_table(self):
        """П. 10.2 — вложенная таблица норм по табл. 4.8 для стали кат. II."""
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            section10 = next(
                t for t in doc.tables
                if any('10.2' in c.text for row in t.rows for c in row.cells)
            )
            row_102 = next(
                row for row in section10.rows if row.cells[0].text.strip().startswith('10.2')
            )
            cell = row_102.cells[0]
            self.assertIn('таблица N 4.8', cell.text)
            self.assertIn('20,0', cell.text)
            nested = cell.tables
            self.assertEqual(len(nested), 1)
            self.assertGreater(len(nested[0].rows), 1)
            data_row = ' '.join(c.text for c in nested[0].rows[1].cells)
            self.assertIn('2,5', data_row)

    def test_section_102_aluminum_uses_table_410(self):
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        data = dict(
            self.pipe_input,
            material='__aluminum__',
            material_custom='АМг6',
            weld_category='II',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            section10 = next(
                t for t in doc.tables
                if any('10.2' in c.text for row in t.rows for c in row.cells)
            )
            row_102 = next(
                row for row in section10.rows if row.cells[0].text.strip().startswith('10.2')
            )
            self.assertIn('таблица N 4.10', row_102.cells[0].text)

    def test_section_102_row_has_no_fixed_height(self):
        """П. 10.2 — без фиксированной высоты строки после вставки таблицы норм."""
        from techcards.generator import generate_from_template, get_default_template_path
        from docx.oxml.ns import qn
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            section10 = next(
                t for t in doc.tables
                if any('10.2' in c.text for row in t.rows for c in row.cells)
            )
            row_102 = next(
                row for row in section10.rows if row.cells[0].text.strip().startswith('10.2')
            )
            tr_pr = row_102._tr.find(qn('w:trPr'))
            if tr_pr is not None:
                self.assertIsNone(tr_pr.find(qn('w:trHeight')))

    def test_section_102_cell_fits_table_without_trailing_gap(self):
        """П. 10.2 — ячейка без пустых абзацев после вложенной таблицы."""
        from techcards.generator import (
            generate_from_template, get_default_template_path, _is_empty_body_paragraph,
        )
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            section10 = next(
                t for t in doc.tables
                if any('10.2' in c.text for row in t.rows for c in row.cells)
            )
            row_102 = next(
                row for row in section10.rows if row.cells[0].text.strip().startswith('10.2')
            )
            cell = row_102.cells[0]
            children = [c for c in cell._tc if c.tag.split('}')[-1] != 'tcPr']
            tbl_idx = next(
                i for i, c in enumerate(children) if c.tag.split('}')[-1] == 'tbl'
            )
            trailing = children[tbl_idx + 1:]
            self.assertFalse(
                any(
                    c.tag.split('}')[-1] == 'p' and _is_empty_body_paragraph(c)
                    for c in trailing
                ),
                'После таблицы норм не должно быть пустых абзацев',
            )

    def test_titanium_section_72_includes_np104_point_84(self):
        """Для титана в п. 7.2 добавляется требование НП-104-18, п. 84."""
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        data = dict(
            self.pipe_input,
            material='__titanium__',
            material_custom='ВТ6',
            weld_category='III',
            welding_process='30',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            prep_table = next(
                t for t in doc.tables
                if any('7.2' in c.text for row in t.rows for c in row.cells)
            )
            row_72 = next(
                row for row in prep_table.rows if '7.2' in row.cells[0].text
            )
            value = row_72.cells[-1].text
            self.assertIn('НП-104-18', value)
            self.assertIn('п. 84', value)
            self.assertIn('цветов побежалости', value)
            self.assertIn('20,0 мм', value)

    def test_titanium_esw_section_72_uses_50mm_width(self):
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        data = dict(
            self.pipe_input,
            material='__titanium__',
            material_custom='ВТ6',
            weld_category='III',
            joint_designation='С-17',
            welding_process='20',
        )
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            prep_table = next(
                t for t in doc.tables
                if any('7.2' in c.text for row in t.rows for c in row.cells)
            )
            row_72 = next(
                row for row in prep_table.rows if '7.2' in row.cells[0].text
            )
            self.assertIn('50,0 мм', row_72.cells[-1].text)
            self.assertIn('электрошлаковую сварку', row_72.cells[-1].text)

    def test_section_43_has_page_break_before_gap(self):
        """П. 4.3 — разрыв страницы и одна пустая строка перед заголовком."""
        from techcards.generator import (
            generate_from_template, get_default_template_path,
            _is_empty_body_paragraph, _paragraph_has_page_break,
        )
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            para_43 = next(
                p for p in doc.paragraphs if '4.3' in p.text and 'эскиз' in p.text.lower()
            )
            empty = 0
            prev = para_43._element.getprevious()
            while prev is not None and _is_empty_body_paragraph(prev):
                empty += 1
                prev = prev.getprevious()
            self.assertEqual(empty, 1)
            if prev is not None:
                self.assertTrue(_paragraph_has_page_break(prev))

    def test_title_card_number_single_instance(self):
        """Номер техкарты на титуле — один раз, без дублирования фрагментов."""
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        data = dict(self.pipe_input, card_number='тк 67')
        calc = RadiographicTechCardCalculator(data)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            num_para = next(
                p for p in doc.paragraphs
                if p.text.strip().startswith('№') and 'тк 67' in p.text
            )
            self.assertEqual(num_para.text.strip(), '№ тк 67')
            self.assertEqual(len(num_para.runs), 1)

    def test_header_page_number_format(self):
        """Колонтитул: «страница N страниц M» без статичных цифр шаблона."""
        from techcards.generator import generate_from_template, get_default_template_path
        import tempfile
        import zipfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            page_cell = doc.sections[0].header.tables[0].rows[1].cells[2]
            self.assertIn('страница', page_cell.text.lower())
            self.assertIn('страниц', page_cell.text.lower())
            self.assertNotRegex(page_cell.text, r'\d')
            with zipfile.ZipFile(out) as zf:
                header_xml = zf.read('word/header1.xml').decode('utf-8')
            self.assertIn('страница', header_xml.lower())
            self.assertIn('PAGE', header_xml)
            self.assertIn('NUMPAGES', header_xml)
            self.assertNotIn('Страница 1', header_xml)

    def test_section_69_has_page_break_before_gap(self):
        """П. 6.9 — разрыв страницы и одна пустая строка перед заголовком."""
        from techcards.generator import (
            generate_from_template, get_default_template_path,
            _is_empty_body_paragraph, _paragraph_has_page_break,
        )
        import tempfile
        from docx import Document

        template = get_default_template_path()
        if not template:
            self.skipTest('Шаблон DOCX не найден')
        calc = RadiographicTechCardCalculator(self.pipe_input)
        params = calc.calculate()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, 'card.docx')
            from django.conf import settings
            static_root = str(settings.STATICFILES_DIRS[0])
            generate_from_template(params, template, out, static_root=static_root)
            doc = Document(out)
            para_69 = next(
                p for p in doc.paragraphs
                if '6.9' in p.text and 'схема' in p.text.lower()
            )
            empty = 0
            prev = para_69._element.getprevious()
            while prev is not None and _is_empty_body_paragraph(prev):
                empty += 1
                prev = prev.getprevious()
            self.assertEqual(empty, 1)
            if prev is not None:
                self.assertTrue(_paragraph_has_page_break(prev))
